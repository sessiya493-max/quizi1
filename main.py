"""
AI Quiz Bot
- Groq AI orqali avtomatik test tuzadi
- @QuizBot ga yuboradi va havola beradi
- Knopkali interfeys
- Ko'p akkaunt pool
"""

import asyncio
import json
import logging
import os
import random
import re
import sqlite3
from collections import deque
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

from groq import Groq
from telethon import TelegramClient, events, Button
from telethon.errors import SessionPasswordNeededError, PhoneCodeInvalidError, UserNotParticipantError
from telethon.tl.functions.channels import GetParticipantRequest
from telethon.tl.functions.messages import SendMediaRequest
from telethon.tl.types import InputMediaPoll, Poll, PollAnswer, TextWithEntities

# ============================================================
#  SOZLAMALAR — Railway environment variables dan o'qiladi
# ============================================================
import os as _os

BOT_TOKEN    = _os.environ["BOT_TOKEN"]
API_ID       = int(_os.environ["API_ID"])
API_HASH     = _os.environ["API_HASH"]
GROQ_API_KEY = _os.environ["GROQ_API_KEY"]
ADMIN_IDS    = [int(x) for x in _os.environ.get("ADMIN_IDS", "0").split(",") if x.strip()]
LOG_GROUP_ID = int(_os.environ.get("LOG_GROUP_ID", "0"))  # loglar yuboriladigan guruh ID
NOTIFY_PHONE = _os.environ.get("NOTIFY_PHONE", "")

# Telefon raqamlar: PHONE_NUMBERS env da vergul bilan yoziladi
# Misol: +998901234567,+998901234568
PHONE_NUMBERS = [
    p.strip() for p in _os.environ.get("PHONE_NUMBERS", "").split(",")
    if p.strip()
]

# Humo kartalar: HUMO_CARDS env da vergul bilan
# Misol: 9860 1234 5678 9001,9860 1234 5678 9002
HUMO_CARDS = [
    c.strip() for c in _os.environ.get("HUMO_CARDS", "").split(",")
    if c.strip()
]

AD_EVERY   = int(_os.environ.get("AD_EVERY", "6"))
AD_TEXT    = _os.environ.get("AD_TEXT", "📢 @quiz_import_bot — AI yordamida @QuizBot testi yarating!")
GROQ_MODEL = _os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")

ACCOUNTS_FILE = _os.environ.get("ACCOUNTS_FILE", "/data/accounts.json")
DB_FILE       = _os.environ.get("DB_FILE",       "/data/bot.db")

# /data papkasi bo'lmasa — local papkada saqlaymiz
if not _os.path.exists("/data"):
    _os.makedirs("data", exist_ok=True)
    ACCOUNTS_FILE = "data/accounts.json"
    DB_FILE       = "data/bot.db"


# ============================================================
#  TO'LOV TIZIMI
# ============================================================

# Humo kartalar ro'yxati — navbat bilan beriladi
# ⬇️ O'z karta raqamlaringizni shu yerga yozing!
HUMO_CARDS = [
   "9860 3501 4339 8906",   # karta 1 — o'zgartiring
    "9860 3566 0573 8935",   # karta 2 — o'zgartiring
    "9860 3466 0594 5705",   # kart
]
AI_PRICE           = 2000    # 1 ta AI test narxi (so'm)
FILE_PRICE_PER_25  = 1500    # har 25 savol uchun narx (fayl orqali)
PAYMENT_TIMEOUT    = 600     # sekund (10 daqiqa)
HUMOCARD_BOT    = "@humocardbot"
NOTIFY_PHONE    = "+998934897111"  # @humocardbot xabar keladigan raqam

# Kartalar band/bo'sh holati: card_num -> user_id yoki None
card_assignments: dict = {card: None for card in HUMO_CARDS}


def get_free_card(user_id: int) -> Optional[str]:
    """Bo'sh karta berish — navbat bilan"""
    busy = set(card_assignments.values())
    for card in HUMO_CARDS:
        if card_assignments.get(card) is None:
            card_assignments[card] = user_id
            return card
    return None   # hammasi band


def calc_file_price(q_count: int) -> int:
    """Fayl uchun narx: har 25 savolga 2000 so'm, qisman bo'lsa ham to'liq hisoblanadi"""
    import math
    blocks = math.ceil(q_count / 25)
    return blocks * FILE_PRICE_PER_25


def release_card(card_num: str):
    """Kartani bo'shatish"""
    card_assignments[card_num] = None


def get_db():
    return sqlite3.connect(DB_FILE)


def db_save_user(user_id: int, first_name: str = "",
                 last_name: str = "", username: str = ""):
    """Foydalanuvchini saqlash yoki yangilash"""
    con = get_db()
    con.execute("""
        INSERT INTO users (user_id, first_name, last_name, username, created_at, last_seen)
        VALUES (?, ?, ?, ?, datetime('now'), datetime('now'))
        ON CONFLICT(user_id) DO UPDATE SET
            first_name = excluded.first_name,
            last_name  = excluded.last_name,
            username   = excluded.username,
            last_seen  = datetime('now')
    """, (user_id, first_name or "", last_name or "", username or ""))
    con.commit()
    con.close()


def db_count_users() -> int:
    """Jami foydalanuvchilar soni"""
    con = get_db()
    n = con.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    con.close()
    return n


def db_get_users(limit: int = 20, offset: int = 0) -> list:
    """Foydalanuvchilar ro'yxati"""
    con = get_db()
    rows = con.execute("""
        SELECT user_id, first_name, last_name, username, created_at, last_seen
        FROM users ORDER BY last_seen DESC LIMIT ? OFFSET ?
    """, (limit, offset)).fetchall()
    con.close()
    return rows


def db_get_user(user_id: int):
    """Bitta foydalanuvchi"""
    con = get_db()
    row = con.execute(
        "SELECT * FROM users WHERE user_id=?", (user_id,)
    ).fetchone()
    con.close()
    return row


def db_get_balance(user_id: int) -> int:
    con = get_db()
    row = con.execute(
        "SELECT balance FROM users WHERE user_id=?", (user_id,)
    ).fetchone()
    con.close()
    return row[0] if row else 0



def db_add_balance(user_id: int, amount: int, reason: str = ""):
    con = get_db()
    con.execute(
        "UPDATE users SET balance = balance + ? WHERE user_id=?",
        (amount, user_id)
    )
    con.execute(
        """INSERT INTO balance_log (user_id, amount, reason, created_at)
           VALUES (?, ?, ?, datetime('now'))""",
        (user_id, amount, reason)
    )
    con.commit()
    con.close()


def db_deduct_balance(user_id: int, amount: int, reason: str = "") -> bool:
    con = get_db()
    row = con.execute(
        "SELECT balance FROM users WHERE user_id=?", (user_id,)
    ).fetchone()
    if not row or row[0] < amount:
        con.close()
        return False
    con.execute(
        "UPDATE users SET balance = balance - ? WHERE user_id=?",
        (amount, user_id)
    )
    con.execute(
        """INSERT INTO balance_log (user_id, amount, reason, created_at)
           VALUES (?, ?, ?, datetime('now'))""",
        (user_id, -amount, reason)
    )
    con.commit()
    con.close()
    return True


def db_create_payment(user_id: int, card_num: str, amount: int) -> int:
    con = get_db()
    cur = con.cursor()
    cur.execute(
        """INSERT INTO payments
           (user_id, card_num, amount, status, created_at, expires_at)
           VALUES (?, ?, ?, 'pending',
                   datetime('now'),
                   datetime('now','+3 minutes'))""",
        (user_id, card_num, amount)
    )
    pay_id = cur.lastrowid
    con.commit()
    con.close()
    return pay_id


def db_get_pending(user_id: int) -> Optional[tuple]:
    """Foydalanuvchining kutilayotgan to'lovi (id, card, amount, expires)"""
    con = get_db()
    row = con.execute(
        """SELECT id, card_num, amount, expires_at FROM payments
           WHERE user_id=? AND status='pending'
           AND expires_at > datetime('now')
           ORDER BY id DESC LIMIT 1""",
        (user_id,)
    ).fetchone()
    con.close()
    return row


def db_confirm_payment(pay_id: int) -> Optional[tuple]:
    """To'lovni tasdiqlash → (user_id, amount, card_num)"""
    con = get_db()
    row = con.execute(
        "SELECT user_id, amount, card_num FROM payments WHERE id=? AND status='pending'",
        (pay_id,)
    ).fetchone()
    if row:
        con.execute(
            """UPDATE payments SET status='confirmed',
               paid_at=datetime('now') WHERE id=?""",
            (pay_id,)
        )
        con.commit()
    con.close()

    # Hamkor komissiyasi — to'lov summасidan 20%
    if row:
        user_id, amount, _ = row
        partner_id = db_get_referred_by(user_id)
        if partner_id:
            commission = int(amount * PARTNER_PAY_PERCENT / 100)
            if commission > 0:
                db_add_partner_balance(partner_id, commission, f"To'lov komissiyasi: user {user_id}, {amount} so'm")

    return row


def db_expire_old():
    """Muddati o'tgan to'lovlarni bekor qilish"""
    con = get_db()
    expired = con.execute(
        """SELECT id, card_num FROM payments WHERE status='pending'
           AND expires_at < datetime('now')"""
    ).fetchall()
    for pay_id, card_num in expired:
        con.execute(
            "UPDATE payments SET status='expired' WHERE id=?", (pay_id,)
        )
        release_card(card_num)
    con.commit()
    con.close()
    return len(expired)


def db_payment_stats() -> dict:
    con = get_db()
    total = con.execute(
        "SELECT COALESCE(SUM(amount),0) FROM payments WHERE status='confirmed'"
    ).fetchone()[0]
    today = con.execute(
        """SELECT COALESCE(SUM(amount),0) FROM payments
           WHERE status='confirmed' AND DATE(paid_at)=DATE('now')"""
    ).fetchone()[0]
    pending = con.execute(
        "SELECT COUNT(*) FROM payments WHERE status='pending'"
    ).fetchone()[0]
    con.close()
    return {"total": total, "today": today, "pending": pending}


def db_init():
    """Barcha jadvallarni yaratish"""
    con = get_db()
    cur = con.cursor()
    cur.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            user_id     INTEGER PRIMARY KEY,
            first_name  TEXT DEFAULT '',
            last_name   TEXT DEFAULT '',
            username    TEXT DEFAULT '',
            balance     INTEGER DEFAULT 0,
            invited_by  INTEGER DEFAULT NULL,
            created_at  TEXT DEFAULT '',
            last_seen   TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS payments (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            card_num    TEXT NOT NULL,
            amount      INTEGER NOT NULL,
            status      TEXT DEFAULT 'pending',
            created_at  TEXT DEFAULT '',
            paid_at     TEXT DEFAULT '',
            expires_at  TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS balance_log (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            amount      INTEGER NOT NULL,
            reason      TEXT DEFAULT '',
            created_at  TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS quizzes (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            fan_name    TEXT DEFAULT '',
            q_count     INTEGER DEFAULT 0,
            variant_num INTEGER DEFAULT 1,
            url         TEXT DEFAULT '',
            time_choice TEXT DEFAULT '30',
            order_type  TEXT DEFAULT 'order',
            source      TEXT DEFAULT 'ai',
            created_at  TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS sessions (
            phone       TEXT PRIMARY KEY,
            session_data TEXT NOT NULL,
            updated_at  TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS referrals (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            inviter_id  INTEGER NOT NULL,
            invited_id  INTEGER NOT NULL,
            bonus       INTEGER DEFAULT 500,
            created_at  TEXT DEFAULT ''
        );
    """)
    # Eski DB lar uchun migration
    migrations = [
        "ALTER TABLE users ADD COLUMN balance INTEGER DEFAULT 0",
        "ALTER TABLE users ADD COLUMN invited_by INTEGER DEFAULT NULL",
        "ALTER TABLE users ADD COLUMN terms_agreed INTEGER DEFAULT 0",
        """CREATE TABLE IF NOT EXISTS referrals (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            inviter_id  INTEGER NOT NULL,
            invited_id  INTEGER NOT NULL,
            bonus       INTEGER DEFAULT 500,
            created_at  TEXT DEFAULT ''
        )""",
    ]
    for sql in migrations:
        try:
            cur.execute(sql)
            con.commit()
        except Exception:
            pass
    # Hamkorlik jadvallari
    db_init_partner_tables(con)

    # CLICK invoicelar jadvali
    con.execute("""
        CREATE TABLE IF NOT EXISTS click_invoices (
            id                INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id           INTEGER NOT NULL,
            amount            INTEGER NOT NULL,
            merchant_trans_id TEXT UNIQUE NOT NULL,
            status            TEXT DEFAULT 'pending',
            created_at        TEXT DEFAULT '',
            paid_at           TEXT DEFAULT ''
        )
    """)
    con.execute("""
        CREATE TABLE IF NOT EXISTS mandatory_subscription (
            id          INTEGER PRIMARY KEY CHECK (id = 1),
            enabled     INTEGER DEFAULT 0,
            chat_ref    TEXT DEFAULT '',
            invite_link TEXT DEFAULT '',
            title       TEXT DEFAULT ''
        )
    """)
    con.execute("INSERT OR IGNORE INTO mandatory_subscription (id, enabled, chat_ref, invite_link, title) VALUES (1, 0, '', '', '')")
    con.commit()
    con.close()


def db_save_session(phone: str, session_path: str):
    """Sessiya faylini DB ga saqlash (base64)"""
    import base64
    if not _os.path.exists(session_path + ".session"):
        return
    with open(session_path + ".session", "rb") as f:
        data = base64.b64encode(f.read()).decode()
    con = get_db()
    con.execute("""
        INSERT INTO sessions (phone, session_data, updated_at)
        VALUES (?, ?, datetime('now'))
        ON CONFLICT(phone) DO UPDATE SET
            session_data = excluded.session_data,
            updated_at   = excluded.updated_at
    """, (phone, data))
    con.commit()
    con.close()
    log.info(f"Sessiya DB ga saqlandi: {phone}")


def db_load_session(phone: str, session_path: str) -> bool:
    """DB dan sessiya faylini tiklash"""
    import base64
    con = get_db()
    row = con.execute(
        "SELECT session_data FROM sessions WHERE phone=?", (phone,)
    ).fetchone()
    con.close()
    if not row:
        return False
    try:
        data = base64.b64decode(row[0].encode())
        sess_dir = _os.path.dirname(session_path)
        if sess_dir:
            _os.makedirs(sess_dir, exist_ok=True)
        with open(session_path + ".session", "wb") as f:
            f.write(data)
        log.info(f"Sessiya DB dan tiklandi: {phone}")
        return True
    except Exception as e:
        log.error(f"Sessiya tiklash xato: {e}")
        return False


def db_save_quiz(user_id: int, fan_name: str, q_count: int,
                 variant_num: int, url: str, time_choice: str,
                 order_type: str, source: str = "ai") -> int:
    """Quiz ma'lumotlarini saqlash, id qaytaradi"""
    con = get_db()
    cur = con.cursor()
    cur.execute("""
        INSERT INTO quizzes
        (user_id, fan_name, q_count, variant_num, url, time_choice, order_type, source, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))
    """, (user_id, fan_name, q_count, variant_num, url, time_choice, order_type, source))
    quiz_id = cur.lastrowid
    con.commit()
    con.close()
    return quiz_id


def db_get_user_quizzes(user_id: int, limit: int = 20) -> list:
    """Foydalanuvchining quizlari"""
    con = get_db()
    rows = con.execute("""
        SELECT id, fan_name, q_count, variant_num, url, source, created_at
        FROM quizzes
        WHERE user_id = ?
        ORDER BY id DESC
        LIMIT ?
    """, (user_id, limit)).fetchall()
    con.close()
    return rows


def db_count_user_quizzes(user_id: int) -> int:
    con = get_db()
    n = con.execute(
        "SELECT COUNT(*) FROM quizzes WHERE user_id=?", (user_id,)
    ).fetchone()[0]
    con.close()
    return n


# ============================================================
#  REFERAL TIZIMI
# ============================================================
# ============================================================
#  HTML SAHIFALAR — statik fayllar o'rniga
# ============================================================
PRIVACY_HTML = """
<!DOCTYPE html>
<html lang="uz">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Maxfiylik Siyosati — AI Quiz Bot</title>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=DM+Sans:wght@300;400;500&display=swap" rel="stylesheet">
<style>
  :root {
    --bg: #0d0f14;
    --surface: #141720;
    --border: #1e2330;
    --accent: #4f7fff;
    --accent2: #7c5cfc;
    --text: #e8eaf0;
    --muted: #6b7280;
    --gold: #c9a84c;
  }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body {
    background: var(--bg);
    color: var(--text);
    font-family: 'DM Sans', sans-serif;
    font-weight: 300;
    line-height: 1.8;
    min-height: 100vh;
  }
  .noise {
    position: fixed; inset: 0; pointer-events: none; z-index: 0;
    background-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 200 200' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.9' numOctaves='4'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)' opacity='0.03'/%3E%3C/svg%3E");
    opacity: 0.4;
  }
  .glow {
    position: fixed; top: -200px; left: 50%; transform: translateX(-50%);
    width: 600px; height: 400px;
    background: radial-gradient(ellipse, rgba(79,127,255,0.08) 0%, transparent 70%);
    pointer-events: none; z-index: 0;
  }
  .container {
    position: relative; z-index: 1;
    max-width: 780px; margin: 0 auto;
    padding: 60px 24px 100px;
  }
  header {
    text-align: center; margin-bottom: 64px;
    animation: fadeUp .6s ease both;
  }
  .bot-badge {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 6px 16px; border-radius: 100px;
    border: 1px solid var(--border);
    background: var(--surface);
    font-size: 12px; color: var(--muted); letter-spacing: .08em;
    margin-bottom: 24px;
  }
  .bot-badge span { width: 6px; height: 6px; border-radius: 50%; background: var(--accent); display: inline-block; animation: pulse 2s infinite; }
  h1 {
    font-family: 'Playfair Display', serif;
    font-size: clamp(28px, 5vw, 44px);
    font-weight: 700; line-height: 1.2;
    background: linear-gradient(135deg, #e8eaf0 30%, #4f7fff);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    margin-bottom: 12px;
  }
  .subtitle { color: var(--muted); font-size: 14px; }
  .lang-switch {
    display: flex; justify-content: center; gap: 8px;
    margin-bottom: 48px; animation: fadeUp .6s .1s ease both;
  }
  .lang-btn {
    padding: 8px 20px; border-radius: 8px; border: 1px solid var(--border);
    background: transparent; color: var(--muted); cursor: pointer;
    font-family: 'DM Sans', sans-serif; font-size: 13px; font-weight: 500;
    transition: all .2s;
  }
  .lang-btn.active {
    background: var(--accent); border-color: var(--accent);
    color: #fff;
  }
  .lang-btn:hover:not(.active) { border-color: var(--accent); color: var(--text); }

  .content { display: none; }
  .content.active { display: block; animation: fadeUp .4s ease both; }

  .section {
    margin-bottom: 40px;
    border: 1px solid var(--border);
    border-radius: 16px;
    overflow: hidden;
    background: var(--surface);
  }
  .section-header {
    display: flex; align-items: center; gap: 12px;
    padding: 20px 24px;
    border-bottom: 1px solid var(--border);
  }
  .section-num {
    width: 28px; height: 28px; border-radius: 8px;
    background: linear-gradient(135deg, var(--accent), var(--accent2));
    display: flex; align-items: center; justify-content: center;
    font-size: 12px; font-weight: 500; color: #fff; flex-shrink: 0;
  }
  .section-title {
    font-family: 'Playfair Display', serif;
    font-size: 17px; font-weight: 700; color: var(--text);
  }
  .section-body { padding: 20px 24px; }
  .section-body p { color: #a8b0c0; font-size: 14px; margin-bottom: 12px; }
  .section-body p:last-child { margin-bottom: 0; }
  .list { list-style: none; padding: 0; }
  .list li {
    display: flex; gap: 10px; align-items: flex-start;
    color: #a8b0c0; font-size: 14px; margin-bottom: 10px;
  }
  .list li:last-child { margin-bottom: 0; }
  .list li::before {
    content: ''; width: 5px; height: 5px; border-radius: 50%;
    background: var(--accent); flex-shrink: 0; margin-top: 8px;
  }
  .highlight {
    display: inline-block; padding: 2px 8px; border-radius: 4px;
    background: rgba(79,127,255,0.12); color: var(--accent);
    font-size: 13px; font-family: monospace;
  }
  .table { width: 100%; border-collapse: collapse; margin-top: 4px; }
  .table th {
    text-align: left; padding: 10px 12px;
    background: rgba(255,255,255,0.03);
    color: var(--muted); font-size: 11px; letter-spacing: .06em; font-weight: 500;
    border-bottom: 1px solid var(--border);
  }
  .table td {
    padding: 10px 12px; font-size: 13px; color: #a8b0c0;
    border-bottom: 1px solid rgba(255,255,255,0.04);
  }
  .table tr:last-child td { border-bottom: none; }
  .tag {
    display: inline-block; padding: 2px 8px; border-radius: 4px;
    font-size: 11px; font-weight: 500;
  }
  .tag.yes { background: rgba(74,222,128,0.1); color: #4ade80; }
  .tag.no  { background: rgba(248,113,113,0.1); color: #f87171; }

  footer {
    text-align: center; padding-top: 48px;
    color: var(--muted); font-size: 12px;
    border-top: 1px solid var(--border);
    animation: fadeUp .6s .3s ease both;
  }
  footer a { color: var(--accent); text-decoration: none; }
  .updated {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 6px 14px; border-radius: 100px;
    background: var(--surface); border: 1px solid var(--border);
    font-size: 11px; color: var(--muted); margin-bottom: 16px;
  }

  @keyframes fadeUp {
    from { opacity: 0; transform: translateY(16px); }
    to   { opacity: 1; transform: translateY(0); }
  }
  @keyframes pulse {
    0%,100% { opacity: 1; } 50% { opacity: .4; }
  }

  @media (max-width: 600px) {
    .container { padding: 40px 16px 80px; }
    .section-body { padding: 16px; }
    .section-header { padding: 16px; }
  }
</style>
</head>
<body>
<div class="noise"></div>
<div class="glow"></div>
<div class="container">

  <header>
    <div class="bot-badge"><span></span> @quiz_import_bot</div>
    <h1>Maxfiylik Siyosati</h1>
    <p class="subtitle">Privacy Policy · Политика конфиденциальности</p>
  </header>

  <div class="lang-switch">
    <button class="lang-btn active" onclick="setLang('uz')">O'zbek</button>
    <button class="lang-btn" onclick="setLang('ru')">Русский</button>
  </div>

  <!-- ===== O'ZBEK ===== -->
  <div id="uz" class="content active">
    <div class="updated">Yangilangan: 2025-yil</div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">1</div>
        <div class="section-title">Umumiy ma'lumot</div>
      </div>
      <div class="section-body">
        <p>Ushbu maxfiylik siyosati <span class="highlight">@quiz_import_bot</span> Telegram botidan foydalanganda sizning shaxsiy ma'lumotlaringiz qanday to'planishi, ishlatilishi va saqlanishini tushuntiradi.</p>
        <p>Botdan foydalanishni boshlaganingizda, siz ushbu siyosat shartlarini qabul qilgan bo'lasiz.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">2</div>
        <div class="section-title">Qanday ma'lumotlar to'planadi</div>
      </div>
      <div class="section-body">
        <table class="table">
          <tr><th>Ma'lumot turi</th><th>Maqsad</th><th>Saqlanadi</th></tr>
          <tr><td>Telegram ID, ism</td><td>Akkaunt identifikatsiyasi</td><td><span class="tag yes">Ha</span></td></tr>
          <tr><td>Username</td><td>Muloqot uchun</td><td><span class="tag yes">Ha</span></td></tr>
          <tr><td>Yuklangan fayllar</td><td>Quiz yaratish (vaqtinchalik)</td><td><span class="tag no">Yo'q</span></td></tr>
          <tr><td>To'lov ma'lumotlari</td><td>To'lov tasdiqlash</td><td><span class="tag yes">Ha</span></td></tr>
          <tr><td>Quiz tarixi</td><td>Sizning quizlaringizni ko'rsatish</td><td><span class="tag yes">Ha</span></td></tr>
          <tr><td>Balans va tranzaksiyalar</td><td>Moliyaviy hisob</td><td><span class="tag yes">Ha</span></td></tr>
        </table>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">3</div>
        <div class="section-title">Ma'lumotlar uchinchi tomonlarga berilmasmi</div>
      </div>
      <div class="section-body">
        <p>Sizning shaxsiy ma'lumotlaringiz <strong>hech qachon</strong> sotilmaydi yoki reklama maqsadida ishlatilmaydi. Faqat quyidagi holatlarda ma'lumot uzatiladi:</p>
        <ul class="list">
          <li>To'lov tizimi (CLICK, Payme) — to'lovni tasdiqlash uchun faqat zarur ma'lumotlar</li>
          <li>Quiz yaratish — @QuizBot xizmatiga savollar yuboriladi (shaxsiy ma'lumot emas)</li>
          <li>AI xizmat (Claude API) — faqat siz yozgan savol matni, shaxsiy ma'lumot emas</li>
          <li>Qonun talabi bo'yicha — sud yoki davlat organlari rasmiy so'rovi bilan</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">4</div>
        <div class="section-title">To'lov xavfsizligi</div>
      </div>
      <div class="section-body">
        <p>Barcha to'lovlar <strong>CLICK</strong> va <strong>Payme</strong> rasmiy to'lov tizimlari orqali amalga oshiriladi. Karta ma'lumotlaringiz hech qachon botda saqlanmaydi — to'lov tizimlari tomonidan shifrlangan holda qayta ishlanadi.</p>
        <ul class="list">
          <li>Karta raqamingiz bizga ko'rinmaydi</li>
          <li>To'lov summasi o'zgartirib bo'lmaydi — tizim tomonidan qotiriladi</li>
          <li>Webhook signature tekshiruvi orqali soxta to'lovlardan himoya</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">5</div>
        <div class="section-title">Ma'lumotlarni saqlash muddati</div>
      </div>
      <div class="section-body">
        <ul class="list">
          <li>Foydalanuvchi ma'lumotlari — akkaunt faol bo'lgunga qadar</li>
          <li>Yuklangan fayllar — quiz yaratilgandan keyin darhol o'chiriladi</li>
          <li>To'lov tarixi — 3 yil (qonun talabi)</li>
          <li>Quiz tarixi — 1 yil</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">6</div>
        <div class="section-title">Sizning huquqlaringiz</div>
      </div>
      <div class="section-body">
        <ul class="list">
          <li>Ma'lumotlaringizni ko'rish va yuklab olish huquqi</li>
          <li>Ma'lumotlaringizni o'chirish talabi (bot admini orqali)</li>
          <li>To'lov tarixiga shikoyat qilish huquqi</li>
          <li>Maxfiylik siyosatidagi o'zgarishlar haqida xabardor bo'lish</li>
        </ul>
        <p style="margin-top:12px">Murojaat uchun: bot ichida <span class="highlight">/support</span> yoki admin bilan bog'laning.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">7</div>
        <div class="section-title">Cookie va kuzatuv</div>
      </div>
      <div class="section-body">
        <p>Bot hech qanday cookie yoki tashqi kuzatuv tizimidan foydalanmaydi. Faqat Telegram platformasining o'z analitikasi ishlatilishi mumkin.</p>
      </div>
    </div>
  </div>

  <!-- ===== РУССКИЙ ===== -->
  <div id="ru" class="content">
    <div class="updated">Обновлено: 2025 год</div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">1</div>
        <div class="section-title">Общая информация</div>
      </div>
      <div class="section-body">
        <p>Настоящая политика конфиденциальности объясняет, как бот <span class="highlight">@quiz_import_bot</span> собирает, использует и хранит ваши персональные данные при использовании сервиса.</p>
        <p>Начиная использовать бота, вы соглашаетесь с условиями данной политики.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">2</div>
        <div class="section-title">Какие данные собираются</div>
      </div>
      <div class="section-body">
        <table class="table">
          <tr><th>Тип данных</th><th>Цель</th><th>Хранится</th></tr>
          <tr><td>Telegram ID, имя</td><td>Идентификация аккаунта</td><td><span class="tag yes">Да</span></td></tr>
          <tr><td>Username</td><td>Для общения</td><td><span class="tag yes">Да</span></td></tr>
          <tr><td>Загруженные файлы</td><td>Создание квиза (временно)</td><td><span class="tag no">Нет</span></td></tr>
          <tr><td>Платёжные данные</td><td>Подтверждение оплаты</td><td><span class="tag yes">Да</span></td></tr>
          <tr><td>История квизов</td><td>Отображение ваших квизов</td><td><span class="tag yes">Да</span></td></tr>
          <tr><td>Баланс и транзакции</td><td>Финансовый учёт</td><td><span class="tag yes">Да</span></td></tr>
        </table>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">3</div>
        <div class="section-title">Передача данных третьим лицам</div>
      </div>
      <div class="section-body">
        <p>Ваши персональные данные <strong>никогда</strong> не продаются и не используются в рекламных целях. Передача данных осуществляется только в следующих случаях:</p>
        <ul class="list">
          <li>Платёжные системы (CLICK, Payme) — только необходимые данные для подтверждения оплаты</li>
          <li>Создание квиза — в @QuizBot передаются только вопросы (без персональных данных)</li>
          <li>AI-сервис (Claude API) — только текст вопроса, без персональных данных</li>
          <li>По требованию закона — по официальному запросу суда или государственных органов</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">4</div>
        <div class="section-title">Безопасность платежей</div>
      </div>
      <div class="section-body">
        <p>Все платежи осуществляются через официальные платёжные системы <strong>CLICK</strong> и <strong>Payme</strong>. Данные вашей карты никогда не хранятся в боте — они обрабатываются в зашифрованном виде платёжными системами.</p>
        <ul class="list">
          <li>Номер вашей карты нам недоступен</li>
          <li>Сумма платежа фиксируется системой и не может быть изменена</li>
          <li>Защита от поддельных платежей через проверку webhook-подписи</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">5</div>
        <div class="section-title">Срок хранения данных</div>
      </div>
      <div class="section-body">
        <ul class="list">
          <li>Данные пользователя — до тех пор, пока аккаунт активен</li>
          <li>Загруженные файлы — удаляются сразу после создания квиза</li>
          <li>История платежей — 3 года (по требованию закона)</li>
          <li>История квизов — 1 год</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">6</div>
        <div class="section-title">Ваши права</div>
      </div>
      <div class="section-body">
        <ul class="list">
          <li>Право на просмотр и выгрузку своих данных</li>
          <li>Право на удаление данных (через администратора бота)</li>
          <li>Право на обжалование платёжных операций</li>
          <li>Уведомление об изменениях в политике конфиденциальности</li>
        </ul>
        <p style="margin-top:12px">Для обращений: команда <span class="highlight">/support</span> внутри бота или свяжитесь с администратором.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header">
        <div class="section-num">7</div>
        <div class="section-title">Cookie и отслеживание</div>
      </div>
      <div class="section-body">
        <p>Бот не использует cookie или сторонние системы отслеживания. Может использоваться только встроенная аналитика платформы Telegram.</p>
      </div>
    </div>
  </div>

  <footer>
    <p style="margin-bottom:16px">
      <span class="updated">© 2025 AI Quiz Bot</span>
    </p>
    <p>Savollar uchun / По вопросам: <a href="https://t.me/quiz_import_bot">@quiz_import_bot</a></p>
  </footer>

</div>
<script>
function setLang(lang) {
  document.querySelectorAll('.content').forEach(c => c.classList.remove('active'));
  document.querySelectorAll('.lang-btn').forEach(b => b.classList.remove('active'));
  document.getElementById(lang).classList.add('active');
  event.target.classList.add('active');
}
</script>
</body>
</html>

"""

TERMS_HTML = """
<!DOCTYPE html>
<html lang="uz">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Foydalanish Shartlari — AI Quiz Bot</title>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=DM+Sans:wght@300;400;500&display=swap" rel="stylesheet">
<style>
  :root {
    --bg: #0d0f14;
    --surface: #141720;
    --border: #1e2330;
    --accent: #4f7fff;
    --accent2: #7c5cfc;
    --text: #e8eaf0;
    --muted: #6b7280;
  }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body {
    background: var(--bg);
    color: var(--text);
    font-family: 'DM Sans', sans-serif;
    font-weight: 300;
    line-height: 1.8;
    min-height: 100vh;
  }
  .glow {
    position: fixed; top: -200px; left: 50%; transform: translateX(-50%);
    width: 600px; height: 400px;
    background: radial-gradient(ellipse, rgba(124,92,252,0.07) 0%, transparent 70%);
    pointer-events: none; z-index: 0;
  }
  .container {
    position: relative; z-index: 1;
    max-width: 780px; margin: 0 auto;
    padding: 60px 24px 100px;
  }
  header { text-align: center; margin-bottom: 56px; animation: fadeUp .6s ease both; }
  .bot-badge {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 6px 16px; border-radius: 100px;
    border: 1px solid var(--border); background: var(--surface);
    font-size: 12px; color: var(--muted); letter-spacing: .08em; margin-bottom: 24px;
  }
  .bot-badge span { width: 6px; height: 6px; border-radius: 50%; background: var(--accent2); display: inline-block; animation: pulse 2s infinite; }
  h1 {
    font-family: 'Playfair Display', serif;
    font-size: clamp(28px, 5vw, 44px); font-weight: 700; line-height: 1.2;
    background: linear-gradient(135deg, #e8eaf0 30%, #7c5cfc);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    margin-bottom: 12px;
  }
  .subtitle { color: var(--muted); font-size: 14px; }
  .lang-switch {
    display: flex; justify-content: center; gap: 8px;
    margin-bottom: 40px; animation: fadeUp .6s .1s ease both;
  }
  .lang-btn {
    padding: 8px 20px; border-radius: 8px; border: 1px solid var(--border);
    background: transparent; color: var(--muted); cursor: pointer;
    font-family: 'DM Sans', sans-serif; font-size: 13px; font-weight: 500; transition: all .2s;
  }
  .lang-btn.active { background: var(--accent2); border-color: var(--accent2); color: #fff; }
  .lang-btn:hover:not(.active) { border-color: var(--accent2); color: var(--text); }
  .content { display: none; }
  .content.active { display: block; animation: fadeUp .4s ease both; }
  .updated {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 6px 14px; border-radius: 100px;
    background: var(--surface); border: 1px solid var(--border);
    font-size: 11px; color: var(--muted); margin-bottom: 24px;
  }
  .section {
    margin-bottom: 32px; border: 1px solid var(--border);
    border-radius: 16px; overflow: hidden; background: var(--surface);
  }
  .section-header {
    display: flex; align-items: center; gap: 12px;
    padding: 18px 24px; border-bottom: 1px solid var(--border);
  }
  .section-num {
    width: 28px; height: 28px; border-radius: 8px;
    background: linear-gradient(135deg, var(--accent2), var(--accent));
    display: flex; align-items: center; justify-content: center;
    font-size: 12px; font-weight: 500; color: #fff; flex-shrink: 0;
  }
  .section-title { font-family: 'Playfair Display', serif; font-size: 17px; font-weight: 700; }
  .section-body { padding: 20px 24px; }
  .section-body p { color: #a8b0c0; font-size: 14px; margin-bottom: 12px; }
  .section-body p:last-child { margin-bottom: 0; }
  .list { list-style: none; padding: 0; }
  .list li {
    display: flex; gap: 10px; align-items: flex-start;
    color: #a8b0c0; font-size: 14px; margin-bottom: 10px;
  }
  .list li:last-child { margin-bottom: 0; }
  .list li::before {
    content: ''; width: 5px; height: 5px; border-radius: 50%;
    background: var(--accent2); flex-shrink: 0; margin-top: 8px;
  }
  .highlight {
    display: inline-block; padding: 2px 8px; border-radius: 4px;
    background: rgba(124,92,252,0.12); color: var(--accent2);
    font-size: 13px; font-family: monospace;
  }
  .warning-box {
    padding: 14px 16px; border-radius: 10px;
    background: rgba(248,113,113,0.06);
    border: 1px solid rgba(248,113,113,0.2);
    color: #f87171; font-size: 13px; margin-top: 12px; line-height: 1.6;
  }
  .price-table { width: 100%; border-collapse: collapse; margin-top: 4px; }
  .price-table th {
    text-align: left; padding: 10px 12px;
    background: rgba(255,255,255,0.03);
    color: var(--muted); font-size: 11px; letter-spacing: .06em;
    border-bottom: 1px solid var(--border);
  }
  .price-table td {
    padding: 10px 12px; font-size: 13px; color: #a8b0c0;
    border-bottom: 1px solid rgba(255,255,255,0.04);
  }
  .price-table tr:last-child td { border-bottom: none; }
  .price { color: var(--accent2); font-weight: 500; }
  footer {
    text-align: center; padding-top: 48px;
    color: var(--muted); font-size: 12px;
    border-top: 1px solid var(--border);
  }
  footer a { color: var(--accent2); text-decoration: none; }
  @keyframes fadeUp {
    from { opacity: 0; transform: translateY(16px); }
    to   { opacity: 1; transform: translateY(0); }
  }
  @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.4} }
  @media (max-width: 600px) {
    .container { padding: 40px 16px 80px; }
    .section-body, .section-header { padding: 14px 16px; }
  }
</style>
</head>
<body>
<div class="glow"></div>
<div class="container">

  <header>
    <div class="bot-badge"><span></span> @quiz_import_bot</div>
    <h1>Foydalanish Shartlari</h1>
    <p class="subtitle">Terms of Service · Условия использования</p>
  </header>

  <div class="lang-switch">
    <button class="lang-btn active" onclick="setLang('uz', this)">O'zbek</button>
    <button class="lang-btn" onclick="setLang('ru', this)">Русский</button>
  </div>

  <!-- ===== O'ZBEK ===== -->
  <div id="uz" class="content active">
    <div class="updated">Yangilangan: 2025-yil</div>

    <div class="section">
      <div class="section-header"><div class="section-num">1</div><div class="section-title">Umumiy qoidalar</div></div>
      <div class="section-body">
        <p>Ushbu foydalanish shartlari <span class="highlight">@quiz_import_bot</span> Telegram botidan foydalanish qoidalarini belgilaydi. Botdan foydalanishni boshlaganingizda va shartlarni qabul qilganingizda, siz ushbu qoidalarga rioya qilishga rozilik bildirasiz.</p>
        <p>Bot O'zbekiston Respublikasi qonunchiligi asosida faoliyat yuritadi.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">2</div><div class="section-title">Xizmatlar va narxlar</div></div>
      <div class="section-body">
        <table class="price-table">
          <tr><th>Xizmat</th><th>Narx</th></tr>
          <tr><td>AI test tuzish</td><td class="price">2 000 so'm</td></tr>
          <tr><td>Fayldan quiz yaratish (har 25 savol)</td><td class="price">1 500 so'm</td></tr>
          <tr><td>AI taqdimot (15 slayd)</td><td class="price">5 000 so'm</td></tr>
        </table>
        <p style="margin-top:12px">Narxlar oldindan e'lon qilinmасдан o'zgarishi mumkin. O'zgarishlar bot ichida e'lon qilinadi.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">3</div><div class="section-title">To'lov va qaytarish siyosati</div></div>
      <div class="section-body">
        <ul class="list">
          <li>To'lovlar CLICK va Payme rasmiy to'lov tizimlari orqali amalga oshiriladi</li>
          <li>Balansga qo'shilgan mablag' qaytarilmaydi, faqat xizmat uchun ishlatiladi</li>
          <li>Texnik xato sababli quiz yaratilmasa — pul avtomatik qaytariladi</li>
          <li>Noto'g'ri to'lov yoki takroriy to'lov bo'lsa — admin orqali murojaat qiling</li>
        </ul>
        <div class="warning-box">⚠️ To'lov qilingan, lekin xizmat ko'rsatilmagan hollarda 24 soat ichida murojaat qilish shart. Muddati o'tgan da'volar ko'rib chiqilmaydi.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">4</div><div class="section-title">Taqiqlangan harakatlar</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Botni avtomatlashtirilgan usulda (spam) ishlatish</li>
          <li>To'lov tizimini aldash yoki chetlab o'tishga urinish</li>
          <li>Boshqa foydalanuvchilar ma'lumotlariga ruxsatsiz kirish</li>
          <li>Botdan tijorat maqsadida qayta sotish uchun foydalanish</li>
          <li>Mualliflik huquqini buzuvchi materiallar yuklash</li>
          <li>Botni buzish yoki xizmatga zarar yetkazishga urinish</li>
        </ul>
        <div class="warning-box">⚠️ Qoidabuzarlik aniqlanganda akkaunt ogohlantirmasiz bloklash huquqi saqlanadi.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">5</div><div class="section-title">Hamkorlik dasturi shartlari</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Hamkor bo'lish uchun ariza qoldirib, admin tasdig'ini olish shart</li>
          <li>Har jalb qilingan foydalanuvchi uchun: <span class="highlight">+50 so'm</span></li>
          <li>Jalb qilingan foydalanuvchining har to'lovidan: <span class="highlight">20%</span> komissiya</li>
          <li>Minimal chiqarish summasi: <span class="highlight">15 000 so'm</span></li>
          <li>Soxta akkauntlar orqali daromad olishga urinish — akkaunt bloklanadi</li>
          <li>Hamkor balansi naqd pulga o'zgartiriladi, boshqa maqsadda ishlatilmaydi</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">6</div><div class="section-title">Kontent va mualliflik huquqi</div></div>
      <div class="section-body">
        <p>Foydalanuvchi yuklagan fayllar va matnlar uning o'z mas'uliyatida. Bot mualliflik huquqini buzuvchi materiallar uchun javobgar emas.</p>
        <p>AI tomonidan yaratilgan testlar va taqdimotlar foydalanuvchiga tegishli bo'lib, tijorat maqsadida ishlatish mumkin.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">7</div><div class="section-title">Mas'uliyat chegaralari</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Bot xizmatining uzilishi yoki texnik nosozlik uchun bot to'liq mas'ul emas</li>
          <li>Telegram platformasining o'zi sabab bo'lgan muammolar uchun mas'uliyat yo'q</li>
          <li>AI tomonidan yaratilgan kontentning aniqligi kafolatlanmaydi — foydalanuvchi tekshirishi tavsiya etiladi</li>
          <li>Force majeure holatlarida majburiyatlar bajarilmasligi mumkin</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">8</div><div class="section-title">Shartlarning o'zgarishi</div></div>
      <div class="section-body">
        <p>Foydalanish shartlari oldindan xabardor qilinib o'zgartirilishi mumkin. O'zgarishlar bot ichida e'lon qilinadi. Botdan foydalanishni davom ettirish yangi shartlarni qabul qilish hisoblanadi.</p>
      </div>
    </div>
  </div>

  <!-- ===== РУССКИЙ ===== -->
  <div id="ru" class="content">
    <div class="updated">Обновлено: 2025 год</div>

    <div class="section">
      <div class="section-header"><div class="section-num">1</div><div class="section-title">Общие положения</div></div>
      <div class="section-body">
        <p>Настоящие условия использования регулируют правила работы с Telegram-ботом <span class="highlight">@quiz_import_bot</span>. Начиная пользоваться ботом и принимая условия, вы соглашаетесь соблюдать данные правила.</p>
        <p>Бот работает в соответствии с законодательством Республики Узбекистан.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">2</div><div class="section-title">Услуги и цены</div></div>
      <div class="section-body">
        <table class="price-table">
          <tr><th>Услуга</th><th>Цена</th></tr>
          <tr><td>Создание AI-теста</td><td class="price">2 000 сум</td></tr>
          <tr><td>Квиз из файла (каждые 25 вопросов)</td><td class="price">1 500 сум</td></tr>
          <tr><td>AI-презентация (15 слайдов)</td><td class="price">5 000 сум</td></tr>
        </table>
        <p style="margin-top:12px">Цены могут изменяться без предварительного уведомления. Изменения объявляются внутри бота.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">3</div><div class="section-title">Оплата и возврат средств</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Оплата производится через официальные платёжные системы CLICK и Payme</li>
          <li>Средства, добавленные на баланс, не возвращаются — используются только для услуг</li>
          <li>Если квиз не создан по технической ошибке — деньги возвращаются автоматически</li>
          <li>При ошибочной или двойной оплате — обратитесь к администратору</li>
        </ul>
        <div class="warning-box">⚠️ При оплате без получения услуги необходимо обратиться в течение 24 часов. Заявки с истёкшим сроком не рассматриваются.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">4</div><div class="section-title">Запрещённые действия</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Использование бота в автоматизированном режиме (спам)</li>
          <li>Попытки обмануть или обойти платёжную систему</li>
          <li>Несанкционированный доступ к данным других пользователей</li>
          <li>Использование бота для перепродажи в коммерческих целях</li>
          <li>Загрузка материалов, нарушающих авторские права</li>
          <li>Попытки взлома или нанесения вреда сервису</li>
        </ul>
        <div class="warning-box">⚠️ При обнаружении нарушений аккаунт может быть заблокирован без предупреждения.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">5</div><div class="section-title">Условия партнёрской программы</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Для участия необходимо подать заявку и получить подтверждение администратора</li>
          <li>За каждого привлечённого пользователя: <span class="highlight">+50 сум</span></li>
          <li>С каждого платежа привлечённого пользователя: <span class="highlight">20%</span> комиссии</li>
          <li>Минимальная сумма вывода: <span class="highlight">15 000 сум</span></li>
          <li>Попытки заработка через фиктивные аккаунты — блокировка</li>
          <li>Партнёрский баланс конвертируется только в наличные, иное использование недопустимо</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">6</div><div class="section-title">Контент и авторские права</div></div>
      <div class="section-body">
        <p>Файлы и тексты, загружаемые пользователем, являются его ответственностью. Бот не несёт ответственности за материалы, нарушающие авторские права.</p>
        <p>Тесты и презентации, созданные AI, принадлежат пользователю и могут использоваться в коммерческих целях.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">7</div><div class="section-title">Ограничение ответственности</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Бот не несёт полной ответственности за перебои в работе или технические сбои</li>
          <li>Проблемы, вызванные платформой Telegram, не входят в зону ответственности</li>
          <li>Точность контента, созданного AI, не гарантируется — рекомендуется проверка пользователем</li>
          <li>В форс-мажорных обстоятельствах обязательства могут не исполняться</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">8</div><div class="section-title">Изменение условий</div></div>
      <div class="section-body">
        <p>Условия использования могут быть изменены с предварительным уведомлением. Изменения объявляются внутри бота. Продолжение использования бота означает принятие новых условий.</p>
      </div>
    </div>
  </div>

  <footer style="margin-top:48px">
    <p style="margin-bottom:12px">© 2025 AI Quiz Bot</p>
    <p>Savollar / Вопросы: <a href="https://t.me/quiz_import_bot">@quiz_import_bot</a></p>
  </footer>

</div>
<script>
function setLang(lang, btn) {
  document.querySelectorAll('.content').forEach(c => c.classList.remove('active'));
  document.querySelectorAll('.lang-btn').forEach(b => b.classList.remove('active'));
  document.getElementById(lang).classList.add('active');
  btn.classList.add('active');
}
</script>
</body>
</html>

"""

OFERTA_HTML = """
<!DOCTYPE html>
<html lang="uz">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Oferta Shartnomasi — AI Quiz Bot</title>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=DM+Sans:wght@300;400;500&display=swap" rel="stylesheet">
<style>
  :root {
    --bg: #0d0f14;
    --surface: #141720;
    --border: #1e2330;
    --accent: #4f7fff;
    --accent2: #7c5cfc;
    --green: #1D9E75;
    --text: #e8eaf0;
    --muted: #6b7280;
  }
  * { margin: 0; padding: 0; box-sizing: border-box; }
  body {
    background: var(--bg);
    color: var(--text);
    font-family: 'DM Sans', sans-serif;
    font-weight: 300;
    line-height: 1.8;
    min-height: 100vh;
  }
  .glow {
    position: fixed; top: -200px; left: 50%; transform: translateX(-50%);
    width: 600px; height: 400px;
    background: radial-gradient(ellipse, rgba(29,158,117,0.07) 0%, transparent 70%);
    pointer-events: none; z-index: 0;
  }
  .container {
    position: relative; z-index: 1;
    max-width: 780px; margin: 0 auto;
    padding: 60px 24px 100px;
  }
  header { text-align: center; margin-bottom: 56px; animation: fadeUp .6s ease both; }
  .bot-badge {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 6px 16px; border-radius: 100px;
    border: 1px solid var(--border); background: var(--surface);
    font-size: 12px; color: var(--muted); letter-spacing: .08em; margin-bottom: 24px;
  }
  .bot-badge span { width: 6px; height: 6px; border-radius: 50%; background: var(--green); display: inline-block; animation: pulse 2s infinite; }
  h1 {
    font-family: 'Playfair Display', serif;
    font-size: clamp(26px, 5vw, 42px); font-weight: 700; line-height: 1.2;
    background: linear-gradient(135deg, #e8eaf0 30%, #1D9E75);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    margin-bottom: 12px;
  }
  .subtitle { color: var(--muted); font-size: 14px; }
  .lang-switch {
    display: flex; justify-content: center; gap: 8px;
    margin-bottom: 40px; animation: fadeUp .6s .1s ease both;
  }
  .lang-btn {
    padding: 8px 20px; border-radius: 8px; border: 1px solid var(--border);
    background: transparent; color: var(--muted); cursor: pointer;
    font-family: 'DM Sans', sans-serif; font-size: 13px; font-weight: 500; transition: all .2s;
  }
  .lang-btn.active { background: var(--green); border-color: var(--green); color: #fff; }
  .lang-btn:hover:not(.active) { border-color: var(--green); color: var(--text); }
  .content { display: none; }
  .content.active { display: block; animation: fadeUp .4s ease both; }
  .updated {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 6px 14px; border-radius: 100px;
    background: var(--surface); border: 1px solid var(--border);
    font-size: 11px; color: var(--muted); margin-bottom: 24px;
  }
  .section {
    margin-bottom: 28px; border: 1px solid var(--border);
    border-radius: 16px; overflow: hidden; background: var(--surface);
  }
  .section.yatt {
    border-color: rgba(29,158,117,0.3);
    background: rgba(29,158,117,0.04);
  }
  .section-header {
    display: flex; align-items: center; gap: 12px;
    padding: 18px 24px; border-bottom: 1px solid var(--border);
  }
  .section.yatt .section-header { border-bottom-color: rgba(29,158,117,0.2); }
  .section-num {
    width: 28px; height: 28px; border-radius: 8px;
    background: linear-gradient(135deg, var(--green), var(--accent));
    display: flex; align-items: center; justify-content: center;
    font-size: 12px; font-weight: 500; color: #fff; flex-shrink: 0;
  }
  .section-num.gold {
    background: linear-gradient(135deg, #c9a84c, #e8c870);
  }
  .section-title { font-family: 'Playfair Display', serif; font-size: 17px; font-weight: 700; }
  .section-body { padding: 20px 24px; }
  .section-body p { color: #a8b0c0; font-size: 14px; margin-bottom: 12px; line-height: 1.7; }
  .section-body p:last-child { margin-bottom: 0; }
  .list { list-style: none; padding: 0; }
  .list li {
    display: flex; gap: 10px; align-items: flex-start;
    color: #a8b0c0; font-size: 14px; margin-bottom: 10px;
  }
  .list li:last-child { margin-bottom: 0; }
  .list li::before {
    content: ''; width: 5px; height: 5px; border-radius: 50%;
    background: var(--green); flex-shrink: 0; margin-top: 8px;
  }
  .highlight {
    display: inline-block; padding: 2px 8px; border-radius: 4px;
    background: rgba(29,158,117,0.12); color: var(--green);
    font-size: 13px; font-family: monospace;
  }
  .price-table { width: 100%; border-collapse: collapse; }
  .price-table th {
    text-align: left; padding: 10px 12px;
    background: rgba(255,255,255,0.03);
    color: var(--muted); font-size: 11px; letter-spacing: .06em;
    border-bottom: 1px solid var(--border);
  }
  .price-table td {
    padding: 10px 12px; font-size: 13px; color: #a8b0c0;
    border-bottom: 1px solid rgba(255,255,255,0.04);
  }
  .price-table tr:last-child td { border-bottom: none; }
  .price { color: var(--green); font-weight: 500; }
  .yatt-grid {
    display: grid; grid-template-columns: 1fr 1fr; gap: 12px;
  }
  .yatt-item {
    padding: 12px 14px; border-radius: 10px;
    background: rgba(255,255,255,0.03);
    border: 1px solid rgba(29,158,117,0.15);
  }
  .yatt-label { font-size: 11px; color: var(--muted); letter-spacing: .06em; margin-bottom: 4px; }
  .yatt-value { font-size: 14px; color: var(--text); font-weight: 500; }
  .yatt-stir {
    font-family: monospace; font-size: 16px;
    color: var(--green); letter-spacing: .1em;
  }
  .warning-box {
    padding: 14px 16px; border-radius: 10px;
    background: rgba(248,113,113,0.06);
    border: 1px solid rgba(248,113,113,0.2);
    color: #f87171; font-size: 13px; margin-top: 12px; line-height: 1.6;
  }
  .info-box {
    padding: 14px 16px; border-radius: 10px;
    background: rgba(29,158,117,0.06);
    border: 1px solid rgba(29,158,117,0.2);
    color: #4ade80; font-size: 13px; margin-top: 12px; line-height: 1.6;
  }
  .divider {
    height: 1px; background: var(--border);
    margin: 32px 0; opacity: 0.5;
  }
  footer {
    text-align: center; padding-top: 48px;
    color: var(--muted); font-size: 12px;
    border-top: 1px solid var(--border);
  }
  footer a { color: var(--green); text-decoration: none; }
  @keyframes fadeUp {
    from { opacity: 0; transform: translateY(16px); }
    to   { opacity: 1; transform: translateY(0); }
  }
  @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.4} }
  @media (max-width: 600px) {
    .container { padding: 40px 16px 80px; }
    .section-body, .section-header { padding: 14px 16px; }
    .yatt-grid { grid-template-columns: 1fr; }
  }
</style>
</head>
<body>
<div class="glow"></div>
<div class="container">

  <header>
    <div class="bot-badge"><span></span> @quiz_import_bot</div>
    <h1>Ommaviy Oferta Shartnomasi</h1>
    <p class="subtitle">Public Offer Agreement · Публичная оферта</p>
  </header>

  <div class="lang-switch">
    <button class="lang-btn active" onclick="setLang('uz', this)">O'zbek</button>
    <button class="lang-btn" onclick="setLang('ru', this)">Русский</button>
  </div>

  <!-- ===== O'ZBEK ===== -->
  <div id="uz" class="content active">
    <div class="updated">Kuchga kirgan sana: 2025-yil 1-yanvar</div>

    <div class="section">
      <div class="section-header"><div class="section-num">1</div><div class="section-title">Oferta predmeti</div></div>
      <div class="section-body">
        <p>Ushbu ommaviy oferta shartnomasi (keyingi o'rinlarda "Shartnoma") <span class="highlight">@quiz_import_bot</span> Telegram boti orqali raqamli xizmatlar ko'rsatish shartlarini belgilaydi.</p>
        <p>Botdan foydalanish va foydalanish shartlarini qabul qilish ushbu ofertaga rozilik bildirilgan deb hisoblanadi. Shartnoma aktsept (qabul qilish) lahzasidan boshlab kuchga kiradi.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">2</div><div class="section-title">Xizmatlar ro'yxati va narxlar</div></div>
      <div class="section-body">
        <table class="price-table">
          <tr><th>Xizmat nomi</th><th>Tavsif</th><th>Narx</th></tr>
          <tr><td>AI test tuzish</td><td>Claude AI orqali istalgan fandan test</td><td class="price">2 000 so'm</td></tr>
          <tr><td>Fayldan quiz</td><td>Har 25 savolga (DOCX/PDF/TXT)</td><td class="price">1 500 so'm</td></tr>
          <tr><td>AI taqdimot</td><td>15 slaydlik PPTX, rasmlar bilan</td><td class="price">5 000 so'm</td></tr>
        </table>
        <p style="margin-top:12px">Narxlar O'zbekiston so'mida ko'rsatilgan. Xizmat ko'rsatuvchi narxlarni 3 kun oldin e'lon qilgan holda o'zgartirish huquqini saqlab qoladi.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">3</div><div class="section-title">Tomonlarning huquq va majburiyatlari</div></div>
      <div class="section-body">
        <p><strong style="color:var(--text)">Xizmat ko'rsatuvchi majburiyatlari:</strong></p>
        <ul class="list">
          <li>Buyurtma qabul qilinganidan so'ng xizmatni belgilangan muddatda ko'rsatish</li>
          <li>Texnik xato sababli xizmat ko'rsatilmagan taqdirda to'lovni avtomatik qaytarish</li>
          <li>Foydalanuvchi ma'lumotlarini maxfiy saqlash</li>
          <li>Xizmat buzilishi haqida foydalanuvchini xabardor qilish</li>
        </ul>
        <p style="margin-top:14px"><strong style="color:var(--text)">Foydalanuvchi majburiyatlari:</strong></p>
        <ul class="list">
          <li>To'lovni belgilangan miqdorda amalga oshirish</li>
          <li>Mualliflik huquqini hurmat qilgan holda materiallar yuklash</li>
          <li>Botdan faqat qonuniy maqsadlarda foydalanish</li>
          <li>Foydalanish shartlari va ushbu oferta qoidalariga rioya qilish</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">4</div><div class="section-title">To'lov tartibi</div></div>
      <div class="section-body">
        <ul class="list">
          <li>To'lovlar CLICK va Payme rasmiy to'lov tizimlari orqali amalga oshiriladi</li>
          <li>To'lov muvaffaqiyatli o'tkazilgandan so'ng balans darhol yangilanadi</li>
          <li>To'lov summasi tizim tomonidan aniq belgilanadi — o'zgartirib bo'lmaydi</li>
          <li>Har bir to'lov uchun fiskal chek avtomatik shakllantiriladi</li>
        </ul>
        <div class="info-box">✅ Barcha to'lovlar O'zbekiston Respublikasi soliq qonunchiligi talablariga muvofiq rasmiylashtiriladi.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">5</div><div class="section-title">Qaytarish siyosati</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Balansga qo'shilgan mablag' naqd pulga qaytarilmaydi</li>
          <li>Texnik xato sababli quiz yaratilmasa — mablag' balansga avtomatik qaytariladi</li>
          <li>Noto'g'ri to'lov yoki takroriy to'lov — 3 ish kuni ichida ko'rib chiqiladi</li>
          <li>Shikoyat muddati: xizmat ko'rsatilgan kundan boshlab 24 soat</li>
        </ul>
        <div class="warning-box">⚠️ Muddati o'tgan shikoyatlar va asossiz qaytarish talablari ko'rib chiqilmaydi.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">6</div><div class="section-title">Mas'uliyat chegaralari</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Xizmat ko'rsatuvchi Telegram platformasi ishlashidan kelib chiqadigan muammolar uchun mas'ul emas</li>
          <li>AI tomonidan yaratilgan kontent aniqligi kafolatlanmaydi — foydalanuvchi o'zi tekshirishi lozim</li>
          <li>Foydalanuvchi yuklagan materiallar uchun barcha mas'uliyat foydalanuvchida</li>
          <li>Force majeure (tabiiy ofat, urush, epidemiya) holatlarida majburiyatlar to'xtatiladi</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">7</div><div class="section-title">Nizolarni hal qilish</div></div>
      <div class="section-body">
        <p>Nizolar avval muzokaralar orqali hal qilinadi. Muzokaralar muvaffaqiyatsiz bo'lsa, nizo O'zbekiston Respublikasi qonunchiligi asosida Qoraqalpog'iston Respublikasi sudlarida ko'rib chiqiladi.</p>
        <p>Murojaat uchun: bot ichida <span class="highlight">/support</span> komandasi yoki <span class="highlight">+998934897111</span></p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">8</div><div class="section-title">Shartnoma muddati</div></div>
      <div class="section-body">
        <p>Shartnoma noaniq muddatga tuziladi. Foydalanuvchi botdan foydalanishni to'xtatgan taqdirda shartnoma bekor qilingan hisoblanadi. Xizmat ko'rsatuvchi 7 kun oldin ogohlantirish bilan xizmatni to'xtatish huquqiga ega.</p>
      </div>
    </div>

    <!-- YATT MA'LUMOTLARI -->
    <div class="divider"></div>

    <div class="section yatt">
      <div class="section-header">
        <div class="section-num gold">★</div>
        <div class="section-title">Xizmat ko'rsatuvchi ma'lumotlari</div>
      </div>
      <div class="section-body">
        <div class="yatt-grid">
          <div class="yatt-item" style="grid-column: 1 / -1;">
            <div class="yatt-label">TO'LIQ NOMI</div>
            <div class="yatt-value">Karimov Sherali Zokirjon O'g'li</div>
          </div>
          <div class="yatt-item" style="grid-column: 1 / -1;">
            <div class="yatt-label">SOLIQ TO'LOVCHI IDENTIFIKATSIYA RAQAMI (STIR)</div>
            <div class="yatt-stir">502 030 566 000 29</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">TASHKILOT SHAKLI</div>
            <div class="yatt-value">Yakka tartibdagi tadbirkor (YATT)</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">FAOLIYAT TURI</div>
            <div class="yatt-value">Chakana savdo / Raqamli xizmatlar</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">RO'YXATDAN O'TGAN JOY</div>
            <div class="yatt-value">Qoraqalpog'iston Respublikasi, Amudaryo tumani</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">BOG'LANISH</div>
            <div class="yatt-value">+998 93 489 71 11</div>
          </div>
        </div>
      </div>
    </div>

  </div>

  <!-- ===== РУССКИЙ ===== -->
  <div id="ru" class="content">
    <div class="updated">Дата вступления в силу: 1 января 2025 года</div>

    <div class="section">
      <div class="section-header"><div class="section-num">1</div><div class="section-title">Предмет оферты</div></div>
      <div class="section-body">
        <p>Настоящий договор публичной оферты (далее — «Договор») определяет условия оказания цифровых услуг через Telegram-бот <span class="highlight">@quiz_import_bot</span>.</p>
        <p>Использование бота и принятие условий использования считается акцептом настоящей оферты. Договор вступает в силу с момента акцепта.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">2</div><div class="section-title">Перечень услуг и цены</div></div>
      <div class="section-body">
        <table class="price-table">
          <tr><th>Услуга</th><th>Описание</th><th>Цена</th></tr>
          <tr><td>Создание AI-теста</td><td>Тест по любому предмету через Claude AI</td><td class="price">2 000 сум</td></tr>
          <tr><td>Квиз из файла</td><td>Каждые 25 вопросов (DOCX/PDF/TXT)</td><td class="price">1 500 сум</td></tr>
          <tr><td>AI-презентация</td><td>15 слайдов PPTX с изображениями</td><td class="price">5 000 сум</td></tr>
        </table>
        <p style="margin-top:12px">Цены указаны в узбекских сумах. Исполнитель вправе изменять цены с уведомлением за 3 дня.</p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">3</div><div class="section-title">Права и обязанности сторон</div></div>
      <div class="section-body">
        <p><strong style="color:var(--text)">Обязанности исполнителя:</strong></p>
        <ul class="list">
          <li>Оказывать услугу в установленные сроки после принятия заказа</li>
          <li>Автоматически возвращать оплату при технической ошибке</li>
          <li>Хранить данные пользователя в конфиденциальности</li>
          <li>Уведомлять пользователя о сбоях в работе сервиса</li>
        </ul>
        <p style="margin-top:14px"><strong style="color:var(--text)">Обязанности пользователя:</strong></p>
        <ul class="list">
          <li>Производить оплату в установленном размере</li>
          <li>Загружать материалы с соблюдением авторских прав</li>
          <li>Использовать бота исключительно в законных целях</li>
          <li>Соблюдать условия использования и настоящую оферту</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">4</div><div class="section-title">Порядок оплаты</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Оплата производится через платёжные системы CLICK и Payme</li>
          <li>После успешной оплаты баланс пополняется мгновенно</li>
          <li>Сумма платежа фиксируется системой и не может быть изменена</li>
          <li>Фискальный чек формируется автоматически при каждой оплате</li>
        </ul>
        <div class="info-box">✅ Все платежи оформляются в соответствии с налоговым законодательством Республики Узбекистан.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">5</div><div class="section-title">Политика возврата</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Средства, добавленные на баланс, наличными не возвращаются</li>
          <li>При технической ошибке средства автоматически возвращаются на баланс</li>
          <li>Ошибочные или двойные платежи рассматриваются в течение 3 рабочих дней</li>
          <li>Срок подачи жалобы: 24 часа с момента оказания услуги</li>
        </ul>
        <div class="warning-box">⚠️ Жалобы с истёкшим сроком и необоснованные требования возврата не рассматриваются.</div>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">6</div><div class="section-title">Ограничение ответственности</div></div>
      <div class="section-body">
        <ul class="list">
          <li>Исполнитель не несёт ответственности за сбои платформы Telegram</li>
          <li>Точность контента, созданного AI, не гарантируется — пользователь обязан проверить</li>
          <li>Ответственность за загружаемые материалы полностью лежит на пользователе</li>
          <li>В форс-мажорных обстоятельствах обязательства приостанавливаются</li>
        </ul>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">7</div><div class="section-title">Разрешение споров</div></div>
      <div class="section-body">
        <p>Споры разрешаются путём переговоров. В случае неудачи — в судах Республики Каракалпакстан в соответствии с законодательством Республики Узбекистан.</p>
        <p>Для обращений: команда <span class="highlight">/support</span> внутри бота или <span class="highlight">+998934897111</span></p>
      </div>
    </div>

    <div class="section">
      <div class="section-header"><div class="section-num">8</div><div class="section-title">Срок действия договора</div></div>
      <div class="section-body">
        <p>Договор заключается на неопределённый срок. Прекращение использования бота означает расторжение договора. Исполнитель вправе прекратить оказание услуг с уведомлением за 7 дней.</p>
      </div>
    </div>

    <div class="divider"></div>

    <div class="section yatt">
      <div class="section-header">
        <div class="section-num gold">★</div>
        <div class="section-title">Реквизиты исполнителя</div>
      </div>
      <div class="section-body">
        <div class="yatt-grid">
          <div class="yatt-item" style="grid-column: 1 / -1;">
            <div class="yatt-label">ПОЛНОЕ НАИМЕНОВАНИЕ</div>
            <div class="yatt-value">Каримов Шерали Зокиржон Угли</div>
          </div>
          <div class="yatt-item" style="grid-column: 1 / -1;">
            <div class="yatt-label">ИДЕНТИФИКАЦИОННЫЙ НОМЕР НАЛОГОПЛАТЕЛЬЩИКА (ИНН)</div>
            <div class="yatt-stir">502 030 566 000 29</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">ОРГАНИЗАЦИОННАЯ ФОРМА</div>
            <div class="yatt-value">Индивидуальный предприниматель (ИП)</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">ВИД ДЕЯТЕЛЬНОСТИ</div>
            <div class="yatt-value">Розничная торговля / Цифровые услуги</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">МЕСТО РЕГИСТРАЦИИ</div>
            <div class="yatt-value">Республика Каракалпакстан, Амударьинский район</div>
          </div>
          <div class="yatt-item">
            <div class="yatt-label">КОНТАКТ</div>
            <div class="yatt-value">+998 93 489 71 11</div>
          </div>
        </div>
      </div>
    </div>

  </div>

  <footer style="margin-top:48px">
    <p style="margin-bottom:12px">© 2025 Karimov Sherali — AI Quiz Bot</p>
    <p>
      <a href="./privacy.html">Maxfiylik siyosati</a> ·
      <a href="./terms.html">Foydalanish shartlari</a> ·
      <a href="https://t.me/quiz_import_bot">@quiz_import_bot</a>
    </p>
  </footer>

</div>
<script>
function setLang(lang, btn) {
  document.querySelectorAll('.content').forEach(c => c.classList.remove('active'));
  document.querySelectorAll('.lang-btn').forEach(b => b.classList.remove('active'));
  document.getElementById(lang).classList.add('active');
  btn.classList.add('active');
}
</script>
</body>
</html>

"""

REFERRAL_BONUS = 1000  # har ikki tomonga beriladigan so'm

# ============================================================
#  CLICK TO'LOV SOZLAMALARI
# ============================================================
CLICK_SERVICE_ID       = _os.environ.get("CLICK_SERVICE_ID", "")
CLICK_SECRET_KEY       = _os.environ.get("CLICK_SECRET_KEY", "")
CLICK_MERCHANT_ID      = _os.environ.get("CLICK_MERCHANT_ID", "")
CLICK_MERCHANT_USER_ID = _os.environ.get("CLICK_MERCHANT_USER_ID", "")
SERVER_PORT            = int(_os.environ.get("SERVER_PORT", "8080"))
CLICK_BASE_URL         = "https://my.click.uz/services/pay"

# Bot veb serveri asosiy URL — shartlar sahifalari uchun
WEB_BASE_URL = _os.environ.get("WEB_BASE_URL", "")  # masalan: https://yourapp.up.railway.app

def create_click_url(amount: int, merchant_trans_id: str) -> str:
    return (
        f"{CLICK_BASE_URL}"
        f"?service_id={CLICK_SERVICE_ID}"
        f"&merchant_id={CLICK_MERCHANT_ID}"
        f"&amount={amount}"
        f"&transaction_param={merchant_trans_id}"
        f"&return_url=https://t.me/quiz_import_bot"
    )

def verify_click_signature(data: dict, action: int) -> bool:
    """Click MD5 imzosini tekshiradi (sign_time bilan)."""
    import hashlib as _hl, hmac as _hmac
    try:
        if action == 0:
            raw = "{}{}{}{}{}{}{}".format(
                data.get("click_trans_id", ""),
                CLICK_SERVICE_ID,
                CLICK_SECRET_KEY,
                data.get("merchant_trans_id", ""),
                data.get("amount", ""),
                data.get("action", ""),
                data.get("sign_time", ""),
            )
        else:
            raw = "{}{}{}{}{}{}{}{}".format(
                data.get("click_trans_id", ""),
                CLICK_SERVICE_ID,
                CLICK_SECRET_KEY,
                data.get("merchant_trans_id", ""),
                data.get("merchant_prepare_id", ""),
                data.get("amount", ""),
                data.get("action", ""),
                data.get("sign_time", ""),
            )
        expected = _hl.md5(raw.encode("utf-8")).hexdigest()
        received = data.get("sign_string", "")
        ok = _hmac.compare_digest(expected, received)
        if not ok:
            log.warning(f"Click imzo XATO | exp={expected} | got={received}")
        return ok
    except Exception as e:
        log.error(f"Signature xato: {e}")
        return False

def db_create_click_invoice(user_id: int, amount: int) -> str:
    import uuid, time
    merchant_trans_id = f"quiz_{user_id}_{int(time.time())}_{uuid.uuid4().hex[:6]}"
    con = get_db()
    con.execute("""
        INSERT INTO click_invoices (user_id, amount, merchant_trans_id, status, created_at)
        VALUES (?, ?, ?, 'pending', datetime('now'))
    """, (user_id, amount, merchant_trans_id))
    con.commit()
    con.close()
    return merchant_trans_id

def db_get_click_invoice(merchant_trans_id: str) -> Optional[tuple]:
    con = get_db()
    row = con.execute("""
        SELECT id, user_id, amount, status FROM click_invoices
        WHERE merchant_trans_id = ?
    """, (merchant_trans_id,)).fetchone()
    con.close()
    return row

def db_confirm_click_invoice(merchant_trans_id: str) -> Optional[tuple]:
    con = get_db()
    row = con.execute("""
        SELECT user_id, amount FROM click_invoices
        WHERE merchant_trans_id = ? AND status = 'pending'
    """, (merchant_trans_id,)).fetchone()
    if row:
        con.execute("""
            UPDATE click_invoices SET status='paid', paid_at=datetime('now')
            WHERE merchant_trans_id = ?
        """, (merchant_trans_id,))
        con.commit()
    con.close()
    return row

def db_is_new_user(user_id: int) -> bool:
    """Foydalanuvchi avval kelganmi?"""
    con = get_db()
    row = con.execute("SELECT user_id FROM users WHERE user_id=?", (user_id,)).fetchone()
    con.close()
    return row is None

def db_is_agreed(user_id: int) -> bool:
    """Foydalanuvchi shartlarga rozimi?"""
    con = get_db()
    row = con.execute(
        "SELECT terms_agreed FROM users WHERE user_id=?", (user_id,)
    ).fetchone()
    con.close()
    return bool(row and row[0])

def db_set_agreed(user_id: int):
    """Rozilikni DB ga saqlash."""
    con = get_db()
    con.execute(
        "UPDATE users SET terms_agreed=1 WHERE user_id=?", (user_id,)
    )
    con.commit()
    con.close()


def db_set_invited_by(user_id: int, inviter_id: int):
    con = get_db()
    con.execute("UPDATE users SET invited_by=? WHERE user_id=? AND (invited_by IS NULL OR invited_by='')", (inviter_id, user_id))
    con.commit()
    con.close()


def db_get_invited_by(user_id: int) -> Optional[int]:
    con = get_db()
    row = con.execute("SELECT invited_by FROM users WHERE user_id=?", (user_id,)).fetchone()
    con.close()
    if not row or not row[0]:
        return None
    try:
        return int(row[0])
    except Exception:
        return None


def db_get_subscription_settings() -> dict:
    con = get_db()
    row = con.execute("SELECT enabled, chat_ref, invite_link, title FROM mandatory_subscription WHERE id=1").fetchone()
    con.close()
    if not row:
        return {"enabled": 0, "chat_ref": "", "invite_link": "", "title": ""}
    return {
        "enabled": int(row[0] or 0),
        "chat_ref": row[1] or "",
        "invite_link": row[2] or "",
        "title": row[3] or "",
    }


def db_set_subscription_enabled(enabled: bool):
    con = get_db()
    con.execute("UPDATE mandatory_subscription SET enabled=? WHERE id=1", (1 if enabled else 0,))
    con.commit()
    con.close()


def db_set_subscription_target(chat_ref: str, invite_link: str = '', title: str = ''):
    con = get_db()
    con.execute(
        "UPDATE mandatory_subscription SET chat_ref=?, invite_link=?, title=? WHERE id=1",
        ((chat_ref or '').strip(), (invite_link or '').strip(), (title or '').strip())
    )
    con.commit()
    con.close()

def db_save_referral(inviter_id: int, invited_id: int):
    """Referal munosabatini saqlash"""
    con = get_db()
    con.execute("""
        INSERT OR IGNORE INTO referrals (inviter_id, invited_id, bonus, created_at)
        VALUES (?, ?, ?, datetime('now'))
    """, (inviter_id, invited_id, REFERRAL_BONUS))
    con.execute(
        "UPDATE users SET invited_by=? WHERE user_id=?",
        (inviter_id, invited_id)
    )
    con.commit()
    con.close()

def db_get_referral_count(user_id: int) -> int:
    """Foydalanuvchi nechta odam taklif qilgani"""
    con = get_db()
    n = con.execute(
        "SELECT COUNT(*) FROM referrals WHERE inviter_id=?", (user_id,)
    ).fetchone()[0]
    con.close()
    return n

def db_get_referral_list(user_id: int, limit: int = 20) -> list:
    """Taklif qilinganlar ro'yxati"""
    con = get_db()
    rows = con.execute("""
        SELECT r.invited_id, u.first_name, u.last_name, u.username, r.created_at
        FROM referrals r
        LEFT JOIN users u ON u.user_id = r.invited_id
        WHERE r.inviter_id = ?
        ORDER BY r.id DESC
        LIMIT ?
    """, (user_id, limit)).fetchall()
    con.close()
    return rows

def db_already_referred(inviter_id: int, invited_id: int) -> bool:
    """Bu juft allaqachon referalda bormi?"""
    con = get_db()
    row = con.execute(
        "SELECT id FROM referrals WHERE inviter_id=? AND invited_id=?",
        (inviter_id, invited_id)
    ).fetchone()
    con.close()
    return row is not None


# ============================================================
#  HAMKORLIK TIZIMI
# ============================================================
PARTNER_JOIN_BONUS  = 50     # har bir jalb qilgan foydalanuvchi uchun (so'm)
PARTNER_PAY_PERCENT = 20     # to'lovdan foiz (%)
PARTNER_MIN_WITHDRAW = 15000 # minimal chiqarish (so'm)

def db_init_partner_tables(con):
    """Hamkorlik jadvallarini yaratish (db_init ichida chaqiriladi)"""
    con.executescript("""
        CREATE TABLE IF NOT EXISTS partner_applications (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            username    TEXT DEFAULT '',
            full_name   TEXT DEFAULT '',
            comment     TEXT DEFAULT '',
            status      TEXT DEFAULT 'pending',
            created_at  TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS partners (
            user_id     INTEGER PRIMARY KEY,
            partner_balance INTEGER DEFAULT 0,
            total_earned    INTEGER DEFAULT 0,
            total_refs      INTEGER DEFAULT 0,
            created_at  TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS partner_log (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            partner_id  INTEGER NOT NULL,
            amount      INTEGER NOT NULL,
            reason      TEXT DEFAULT '',
            created_at  TEXT DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS partner_withdrawals (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            partner_id  INTEGER NOT NULL,
            amount      INTEGER NOT NULL,
            card_num    TEXT DEFAULT '',
            status      TEXT DEFAULT 'pending',
            created_at  TEXT DEFAULT ''
        );
    """)

def db_is_partner(user_id: int) -> bool:
    con = get_db()
    row = con.execute("SELECT user_id FROM partners WHERE user_id=?", (user_id,)).fetchone()
    con.close()
    return row is not None

def db_get_partner_ids() -> list:
    """PARTNER_IDS env + DB dan"""
    env_ids = [int(x) for x in _os.environ.get("PARTNER_IDS", "").split(",") if x.strip().isdigit()]
    con = get_db()
    rows = con.execute("SELECT user_id FROM partners").fetchall()
    con.close()
    db_ids = [r[0] for r in rows]
    return list(set(env_ids + db_ids))

def db_add_partner(user_id: int):
    con = get_db()
    con.execute("""
        INSERT OR IGNORE INTO partners (user_id, partner_balance, total_earned, total_refs, created_at)
        VALUES (?, 0, 0, 0, datetime('now'))
    """, (user_id,))
    con.commit()
    con.close()

def db_get_partner_info(user_id: int) -> Optional[tuple]:
    """(user_id, partner_balance, total_earned, total_refs)"""
    con = get_db()
    row = con.execute(
        "SELECT user_id, partner_balance, total_earned, total_refs FROM partners WHERE user_id=?",
        (user_id,)
    ).fetchone()
    con.close()
    return row

def db_add_partner_balance(partner_id: int, amount: int, reason: str = ""):
    con = get_db()
    con.execute(
        "UPDATE partners SET partner_balance = partner_balance + ?, total_earned = total_earned + ? WHERE user_id=?",
        (amount, amount, partner_id)
    )
    con.execute(
        "INSERT INTO partner_log (partner_id, amount, reason, created_at) VALUES (?, ?, ?, datetime('now'))",
        (partner_id, amount, reason)
    )
    con.commit()
    con.close()

def db_partner_add_ref(partner_id: int):
    con = get_db()
    con.execute("UPDATE partners SET total_refs = total_refs + 1 WHERE user_id=?", (partner_id,))
    con.commit()
    con.close()

def db_partner_withdraw(partner_id: int, amount: int, card_num: str) -> bool:
    con = get_db()
    row = con.execute("SELECT partner_balance FROM partners WHERE user_id=?", (partner_id,)).fetchone()
    if not row or row[0] < amount:
        con.close()
        return False
    con.execute("UPDATE partners SET partner_balance = partner_balance - ? WHERE user_id=?", (amount, partner_id))
    con.execute(
        "INSERT INTO partner_withdrawals (partner_id, amount, card_num, status, created_at) VALUES (?, ?, ?, 'pending', datetime('now'))",
        (partner_id, amount, card_num)
    )
    con.commit()
    con.close()
    return True

def db_save_partner_application(user_id: int, username: str, full_name: str, comment: str):
    con = get_db()
    con.execute("""
        INSERT INTO partner_applications (user_id, username, full_name, comment, status, created_at)
        VALUES (?, ?, ?, ?, 'pending', datetime('now'))
    """, (user_id, username, full_name, comment))
    con.commit()
    con.close()

def db_has_pending_application(user_id: int) -> bool:
    con = get_db()
    row = con.execute(
        "SELECT id FROM partner_applications WHERE user_id=? AND status='pending'", (user_id,)
    ).fetchone()
    con.close()
    return row is not None

def db_get_referred_by(user_id: int) -> Optional[int]:
    """Bu foydalanuvchi qaysi PARTNER orqali kelgan?"""
    con = get_db()
    row = con.execute("SELECT invited_by FROM users WHERE user_id=?", (user_id,)).fetchone()
    con.close()
    if not row or not row[0]:
        return None
    partner_id = row[0]
    if db_is_partner(partner_id):
        return partner_id
    return None


# ============================================================
#  LOGGING
# ============================================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    handlers=[
        logging.FileHandler("bot.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)
log = logging.getLogger("AIQuizBot")

# ============================================================
#  GROQ AI — TEST TUZISH
# ============================================================
groq_client = Groq(api_key=GROQ_API_KEY)

LANGUAGES = {
    "uz": "O'zbek tilida",
    "ru": "Rus tilida",
    "en": "Ingliz tilida",
}

def build_prompt(fan: str, count: int, lang: str = "uz",
                 difficulty: str = "o'rta", topic: str = "") -> str:
    lang_text = LANGUAGES.get(lang, "O'zbek tilida")
    mavzu_text = f'Mavzu: "{topic}"' if topic else ""
    return f"""Siz {lang_text} test tuzuvchi mutaxassissiz.
"{fan}" fanidan {count} ta test savoli tuz.
{mavzu_text}
Qiyinlik darajasi: {difficulty}.

MUHIM QOIDALAR:
1. Faqat JSON formatda qaytar, boshqa hech narsa yozma
2. Har savolda 4 ta variant bo'lsin
3. Faqat bitta to'g'ri javob bo'lsin
4. Savollar mantiqli va aniq bo'lsin
5. Variantlar bir-biridan farqli bo'lsin
{f'6. Faqat "{topic}" mavzusidan savol tuz' if topic else ""}

JSON format (qat'iy shu ko'rinishda):
[
  {{
    "q": "savol matni",
    "opts": ["variant A", "variant B", "variant C", "variant D"],
    "ans": 0
  }}
]

ans — to'g'ri javob indeksi (0=A, 1=B, 2=C, 3=D)

Hozir {fan} fanidan{f' ({topic} mavzusidan)' if topic else ''} {count} ta savol yoz:"""


async def generate_questions(fan: str, count: int, lang: str = "uz",
                              difficulty: str = "o'rta", topic: str = "") -> list:
    """Groq AI orqali savollar generatsiya qilish — xatolarga chidamli"""
    prompt = build_prompt(fan, count, lang, difficulty, topic)
    log.info(f"AI so'rov: {fan} | mavzu: {topic or 'yo\'q'} | {count} ta | {lang} | {difficulty}")

    loop = asyncio.get_event_loop()

    def _call():
        return groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Siz faqat JSON formatda javob beradigan test tuzuvchisiz. "
                        "Hech qanday izoh, markdown, kod bloki yozma. "
                        "Faqat [ ... ] JSON massivi."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=4000,
        )

    resp = await loop.run_in_executor(None, _call)
    raw = resp.choices[0].message.content.strip()
    log.info(f"AI xom javob (dastlabki 200 belgi): {raw[:200]}")

    questions = _safe_parse_json(raw)

    if not questions:
        log.warning("Birinchi urinish muvaffaqiyatsiz, qayta so'rov yuborilmoqda...")
        # Qayta so'rov — yanada qattiqroq
        def _call2():
            return groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "Faqat sof JSON massivi qaytar. Hech narsa boshqa yo'q."
                    },
                    {"role": "user", "content": prompt},
                    {"role": "assistant", "content": "["},
                ],
            temperature=0.3,
            max_tokens=4000,
            )
        resp2 = await loop.run_in_executor(None, _call2)
        raw2 = "[" + resp2.choices[0].message.content.strip()
        log.info(f"AI 2-javob: {raw2[:200]}")
        questions = _safe_parse_json(raw2)

    # Tekshirish va tozalash
    valid = []
    for q in (questions or []):
        if (isinstance(q, dict) and
                "q" in q and "opts" in q and "ans" in q and
                len(q["opts"]) >= 2 and
                0 <= int(q.get("ans", 0)) < len(q["opts"])):
            opts = [str(o)[:100] for o in q["opts"]]
            ans = int(q["ans"])
            valid.append({
                "q": str(q["q"])[:255],
                "opts": opts,
                "ans": ans,
                "correct_text": opts[ans] if 0 <= ans < len(opts) else opts[0]
            })

    log.info(f"AI {len(valid)} ta savol yaratdi")
    return valid


def _safe_parse_json(text: str) -> Optional[list]:
    """JSON ni xavfsiz parse qilish — bir necha usul bilan"""
    if not text:
        return None

    # 1. Markdown code block olib tashlash
    text = re.sub(r'```(?:json)?', '', text)
    text = re.sub(r'```', '', text)
    text = text.strip().strip('`').strip()

    # 2. [ ... ] qismni ajratib olish
    match = re.search(r'\[.*\]', text, re.DOTALL)
    if match:
        text = match.group(0)

    # 3. To'g'ridan parse
    try:
        data = json.loads(text)
        if isinstance(data, list):
            return data
    except Exception:
        pass

    # 4. Har bir { } blokni alohida parse qilish
    try:
        items = []
        for m in re.finditer(r'\{[^{}]+\}', text, re.DOTALL):
            try:
                obj = json.loads(m.group(0))
                if isinstance(obj, dict):
                    items.append(obj)
            except Exception:
                # Yaroqsiz qatorlarni tuzatib ko'rish
                fixed = _fix_json_obj(m.group(0))
                try:
                    obj = json.loads(fixed)
                    if isinstance(obj, dict):
                        items.append(obj)
                except Exception:
                    pass
        if items:
            return items
    except Exception:
        pass

    return None


def _fix_json_obj(s: str) -> str:
    """Oddiy JSON xatolarini tuzatish"""
    # Oxirgi vergulni olib tashlash
    s = re.sub(r',\s*}', '}', s)
    s = re.sub(r',\s*\]', ']', s)
    # Yagona qo'shtirnoqni ikkilikka almashtirish
    s = re.sub(r"(?<!\\)'", '"', s)
    return s


# ============================================================
#  AKKAUNTLAR SAQLASH
# ============================================================
def load_phones() -> list:
    phones = list(PHONE_NUMBERS)
    if os.path.exists(ACCOUNTS_FILE):
        try:
            with open(ACCOUNTS_FILE) as f:
                for p in json.load(f):
                    if p not in phones:
                        phones.append(p)
        except Exception:
            pass
    return phones

def save_extra_phones(all_phones: list):
    extra = [p for p in all_phones if p not in PHONE_NUMBERS]
    with open(ACCOUNTS_FILE, "w") as f:
        json.dump(extra, f, ensure_ascii=False, indent=2)

# ============================================================
#  MA'LUMOT TUZILMALARI
# ============================================================
@dataclass
class QuizRequest:
    user_id: int
    chat_id: int
    questions: list
    fan_name: str
    variant_num: int
    time_choice: str
    order_choice: str
    total_variants: int = 1
    source: str = "file"   # "ai" yoki "file"
    progress_msg_id: Optional[int] = None

@dataclass
class UserState:
    step: str = "idle"
    # AI rejim
    fan_name: str = ""
    topic: str = ""          # mavzu (ixtiyoriy)
    q_count: int = 10
    lang: str = "uz"
    difficulty: str = "o'rta"
    questions: list = field(default_factory=list)
    total_questions: int = 0
    per_variant: int = 25
    time_choice: str = "30"
    order_choice: str = "order"

# ============================================================
#  GLOBAL
# ============================================================
user_states: dict = {}
admin_states: dict = {}
request_queue: deque = deque()
queue_lock = asyncio.Lock()
account_pool: list = []
account_busy: dict = {}
account_phones: dict = {}
all_clients: list = []   # barcha ulanган клиентлар (pool + notify)
bot_client: TelegramClient = None

# ============================================================
#  VAQT HISOBLASH
# ============================================================
SETUP_SECONDS = 20
SECONDS_PER_QUESTION = 2

def estimate_seconds(n: int) -> int:
    return SETUP_SECONDS + n * SECONDS_PER_QUESTION

def format_wait(s: int) -> str:
    if s < 60:   return f"{s} soniya"
    elif s < 3600:
        m, sec = divmod(s, 60)
        return f"{m} daq {sec} sek" if sec else f"{m} daqiqa"
    else:
        h, m2 = divmod(s, 3600)
        return f"{h} soat {m2//60} daq" if m2 else f"{h} soat"

def calc_wait(new_reqs: list) -> str:
    total_acc = len(account_pool)
    if not total_acc: return "?"
    slots = [0] * total_acc
    for req in list(request_queue):
        slots[slots.index(min(slots))] += estimate_seconds(len(req.questions))
    if not new_reqs:
        return format_wait(min(slots)) if any(slots) else "0 soniya"
    start = min(slots)
    total_new = sum(estimate_seconds(len(r.questions)) for r in new_reqs)
    return format_wait(int(start + total_new))

# ============================================================
#  POOL
# ============================================================
async def pool_add(client, phone):
    account_pool.append(client)
    account_busy[id(client)] = False
    account_phones[id(client)] = phone
    # Sessiyani DB ga saqlash
    sess_dir = _os.path.dirname(DB_FILE)
    session  = _os.path.join(sess_dir, f"userbot_{phone.replace('+','').replace(' ','')}")
    db_save_session(phone, session)

async def pool_remove(phone) -> bool:
    for c in account_pool:
        if account_phones.get(id(c)) == phone:
            if account_busy.get(id(c)): return False
            await c.disconnect()
            account_pool.remove(c)
            account_busy.pop(id(c), None)
            account_phones.pop(id(c), None)
            return True
    return False

async def get_free():
    """Bo'sh va ulangan akkaunt olish — uzilgan bo'lsa qayta ulaydi"""
    while True:
        for c in account_pool:
            if account_busy.get(id(c)):
                continue
            try:
                if not c.is_connected():
                    log.info(f"Akkaunt uzilgan, qayta ulanmoqda: {account_phones.get(id(c))}")
                    await c.connect()
                if await c.is_user_authorized():
                    account_busy[id(c)] = True
                    return c
            except Exception as e:
                log.error(f"Akkaunt tekshirishda xato: {e}")
        await asyncio.sleep(3)

def release(c): account_busy[id(c)] = False
def is_admin(uid): return uid in ADMIN_IDS

# ============================================================
#  QUIZ YARATISH (@QuizBot ga yuborish)
# ============================================================
async def send_poll(userbot, peer, q, opts, ans):
    answers = [PollAnswer(
        text=TextWithEntities(text=o[:100], entities=[]),
        option=bytes([i])
    ) for i, o in enumerate(opts)]
    poll = Poll(
        id=random.randint(1, 2**31),
        question=TextWithEntities(text=q[:255], entities=[]),
        answers=answers, quiz=True,
        public_voters=False, multiple_choice=False, closed=False,
    )
    await userbot(SendMediaRequest(
        peer=peer,
        media=InputMediaPoll(poll=poll, correct_answers=[bytes([ans])]),
        message="", random_id=random.randint(1, 2**63),
    ))

async def make_quiz(userbot: TelegramClient, req: QuizRequest) -> Optional[str]:
    try:
        qbot = await userbot.get_entity("@QuizBot")
        title = f"{req.fan_name} — Variant {req.variant_num}"

        # Oldingi sessiyani tugatish
        try:
            await userbot.send_message(qbot, "/cancel")
            await asyncio.sleep(3)
        except Exception:
            pass

        # Boshlash oldidan oxirgi xabar ID ni eslab qolamiz
        try:
            start_msgs = await userbot.get_messages(qbot, limit=1)
            start_msg_id = start_msgs[0].id if start_msgs else 0
        except Exception:
            start_msg_id = 0
        log.info(f"Start msg ID: {start_msg_id}")

        await userbot.send_message(qbot, "/newquiz"); await asyncio.sleep(4)
        await userbot.send_message(qbot, title);     await asyncio.sleep(3)
        await userbot.send_message(qbot, "/skip");   await asyncio.sleep(3)

        for i, q in enumerate(req.questions):
            try:
                if AD_EVERY > 0 and i > 0 and i % AD_EVERY == 0:
                    await userbot.send_message(qbot, AD_TEXT)
                    await asyncio.sleep(2)

                await send_poll(userbot, qbot, q["q"], q["opts"], stable_answer_index(q))
                log.info(f"  [{i+1}/{len(req.questions)}] OK")
                await asyncio.sleep(2)
            except Exception as e:
                log.error(f"  [{i+1}] xato: {e}")
                await asyncio.sleep(3)

        await userbot.send_message(qbot, "/done")
        await asyncio.sleep(6)

        # Vaqt — eng yangi tugmali xabarni topamiz
        msgs = await userbot.get_messages(qbot, limit=5)
        msg = next((m for m in msgs if m.reply_markup), None)
        if msg:
            tmap = {"15": "15", "30": "30", "60": "60", "0": "No limit"}
            target = tmap.get(req.time_choice, "30")
            clicked = False
            for row in msg.reply_markup.rows:
                for btn in row.buttons:
                    if target in btn.text:
                        await msg.click(text=btn.text); clicked = True; break
                if clicked: break
            if not clicked:
                await msg.click(text=msg.reply_markup.rows[0].buttons[0].text)
        await asyncio.sleep(4)

        # Tartib — eng yangi tugmali xabarni topamiz
        msgs = await userbot.get_messages(qbot, limit=5)
        msg = next((m for m in msgs if m.reply_markup), None)
        if msg:
            clicked = False
            for row in msg.reply_markup.rows:
                for btn in row.buttons:
                    if req.order_choice.lower() in btn.text.lower():
                        await msg.click(text=btn.text); clicked = True; break
                if clicked: break
            if not clicked:
                await msg.click(text=msg.reply_markup.rows[0].buttons[0].text)
        await asyncio.sleep(6)

        # Havola olish — faqat yangi xabarlardan (start_msg_id dan keyin)
        raw_url = None
        for attempt in range(5):
            await asyncio.sleep(4)
            msgs = await userbot.get_messages(qbot, limit=10)
            # Faqat start_msg_id dan keyin kelgan xabarlar
            new_msgs = [m for m in msgs if m.id > start_msg_id]
            log.info(f"Havola qidirilmoqda (urinish {attempt+1}), {len(new_msgs)} ta yangi xabar")

            for m in new_msgs:
                # Tugma ichidagi URL — eng ishonchli
                if m.reply_markup:
                    for row in m.reply_markup.rows:
                        for btn in row.buttons:
                            if hasattr(btn, "url") and btn.url and "t.me" in btn.url:
                                raw_url = btn.url; break
                        if raw_url: break
                if raw_url: break

                # Matn ichidagi URL
                if m.text:
                    urls = re.findall(r"https?://t\.me/\S+", m.text)
                    for url in urls:
                        if "start" in url or "startgroup" in url:
                            raw_url = url; break
                    if not raw_url and urls:
                        raw_url = urls[0]
                if raw_url: break

                # Entity ichidagi URL
                if m.entities:
                    for ent in m.entities:
                        if hasattr(ent, "url") and ent.url and "t.me" in ent.url:
                            raw_url = ent.url; break
                if raw_url: break

            if raw_url:
                log.info(f"Havola topildi: {raw_url}")
                break
            log.warning(f"Havola topilmadi, {attempt+1}-urinish")

        if not raw_url:
            log.error("Havola 5 urinishdan keyin ham topilmadi")
            return None

        fixed = re.sub(r"(https?://t\.me/)([^?/]+)",
                       lambda m: m.group(1) + "QuizBot", raw_url, count=1)
        fixed = fixed.replace("?startgroup=", "?start=")
        return fixed

    except Exception as e:
        log.error(f"make_quiz xato: {e}")
        return None


# ============================================================
#  ADMIN NOTIFY — global funksiya
# ============================================================
async def notify_admin(text: str):
    """Barcha adminlarga xabar yuborish"""
    for admin_id in ADMIN_IDS:
        try:
            await bot_client.send_message(admin_id, text)
        except Exception as e:
            log.error(f"Admin notify xato: {e}")


async def log_action(uid: int, username: str, full_name: str, action: str):
    """User harakatini log guruhiga yuborish"""
    if not LOG_GROUP_ID:
        return
    try:
        uname = f"@{username}" if username else f"#{uid}"
        text = (
            f"👤 {full_name} ({uname})\n"
            f"🆔 `{uid}`\n"
            f"📌 {action}"
        )
        await bot_client.send_message(LOG_GROUP_ID, text)
    except Exception as e:
        log.error(f"Log yuborishda xato: {e}")



# ============================================================
#  USER XABARLARINI TOZALASH VA PROGRESS
# ============================================================
async def safe_delete_messages(chat_id: int, msg_ids: list):
    """Keraksiz xabarlarni jim o'chirish."""
    for mid in list(dict.fromkeys([m for m in msg_ids if m])):
        try:
            await bot_client.delete_messages(chat_id, mid)
        except Exception as e:
            log.debug(f"Xabar o'chmadi {mid}: {e}")


def remember_cleanup(uid: int, *msg_ids):
    """UserState ichida keyin o'chiriladigan xabar ID larni saqlash."""
    st = user_states.get(uid)
    if not st:
        return
    ids = st.__dict__.setdefault("cleanup_msg_ids", [])
    for mid in msg_ids:
        if mid and mid not in ids:
            ids.append(mid)
    user_states[uid] = st


async def cleanup_user_flow_messages(uid: int, chat_id: int):
    """Preview / narx / oraliq xabarlarni o'chirish."""
    st = user_states.get(uid)
    ids = []
    if st:
        ids = st.__dict__.get("cleanup_msg_ids", []) or []
        st.__dict__["cleanup_msg_ids"] = []
        user_states[uid] = st
    if ids:
        await safe_delete_messages(chat_id, ids)


def progress_text(percent: int, done: int, total: int, remain_seconds: int) -> str:
    filled = max(0, min(10, round(percent / 10)))
    bar = "🟦" * filled + "⬜" * (10 - filled)
    return (
        "🚀 **Quiz yaratilmoqda...**\n\n"
        f"📊 Jarayon: **{percent}%**\n"
        f"{bar}\n\n"
        f"✅ Tayyorlangan: **{done} ta**\n"
        f"⏳ Qolgan: **{max(total - done, 0)} ta**\n"
        f"🕐 Taxminiy qolgan vaqt: **~{format_wait(max(remain_seconds, 0))}**\n\n"
        "💡 Botdan chiqib ketmang. Tayyor bo'lgach quiz link avtomatik yuboriladi."
    )



def _norm_answer_text(x: str) -> str:
    return re.sub(r"\s+", " ", str(x or "")).strip().lower()

def stable_answer_index(q: dict) -> int:
    """To'g'ri javobni matn bo'yicha topadi. Variantlar tartibi o'zgarsa ham indeks adashmaydi."""
    opts = q.get("opts") or []
    ans = int(q.get("ans", 0) or 0)
    correct_text = q.get("correct_text")
    if correct_text:
        target = _norm_answer_text(correct_text)
        for i, opt in enumerate(opts):
            if _norm_answer_text(opt) == target:
                return i
    if 0 <= ans < len(opts):
        return ans
    return 0

# ============================================================
#  NAVBAT ISHLOVCHISI
# ============================================================
async def queue_worker():
    log.info("Navbat ishlovchisi ishga tushdi")
    while True:
        try:
            if request_queue and account_pool:
                async with queue_lock:
                    if request_queue:  # Lock ichida qayta tekshiramiz
                        req = request_queue.popleft()
                        userbot = await get_free()
                        asyncio.create_task(run_request(userbot, req))
        except Exception as e:
            log.error(f"queue_worker xato: {e}")
        await asyncio.sleep(1)

async def run_request(userbot, req: QuizRequest):
    import time
    started = time.time()
    progress_msg = getattr(req, "progress_msg_id", None)
    progress_task = None

    async def progress_loop():
        last_text = ""
        total = max(len(req.questions), 1)
        est = max(estimate_seconds(total), 20)
        while True:
            elapsed = int(time.time() - started)
            # 95% gacha ko'rsatamiz, 100% faqat link tayyor bo'lganda chiqadi
            percent = min(95, max(5, int((elapsed / est) * 100)))
            done = min(total, max(0, int(total * percent / 100)))
            remain = max(est - elapsed, 1)
            txt = progress_text(percent, done, total, remain)
            try:
                nonlocal progress_msg
                if progress_msg is None:
                    progress_msg = await bot_client.send_message(req.chat_id, txt)
                elif txt != last_text:
                    if hasattr(progress_msg, "edit"):
                        await progress_msg.edit(txt)
                    else:
                        await bot_client.edit_message(req.chat_id, progress_msg, txt)
                last_text = txt
            except Exception as e:
                if "message was not modified" not in str(e).lower():
                    log.warning(f"Progress edit xato: {e}")
            await asyncio.sleep(4)

    try:
        progress_task = asyncio.create_task(progress_loop())
        try:
            await send_progress_voice(req.user_id)
        except Exception:
            pass

        url = await make_quiz(userbot, req)
        elapsed = int(time.time() - started)
        tl = {"15": "15s", "30": "30s", "60": "60s", "0": "Chegarasiz"}

        if progress_task:
            progress_task.cancel()
            try:
                await progress_task
            except asyncio.CancelledError:
                pass

        if progress_msg:
            try:
                done_txt = progress_text(100, len(req.questions), len(req.questions), 0)
                if hasattr(progress_msg, "edit"):
                    await progress_msg.edit(done_txt)
                    await asyncio.sleep(1)
                    await progress_msg.delete()
                else:
                    await bot_client.edit_message(req.chat_id, progress_msg, done_txt)
                    await asyncio.sleep(1)
                    await bot_client.delete_messages(req.chat_id, progress_msg)
            except Exception:
                pass

        if url:
            db_save_quiz(
                user_id    = req.user_id,
                fan_name   = req.fan_name,
                q_count    = len(req.questions),
                variant_num= req.variant_num,
                url        = url,
                time_choice= req.time_choice,
                order_type = req.order_choice,
                source     = getattr(req, 'source', 'file'),
            )
            src = "🤖 AI" if getattr(req, 'source', 'file') == 'ai' else "📂 Fayl"
            phone = account_phones.get(id(userbot), "?")
            await notify_admin(
                f"✅ **Quiz yaratildi**\n\n"
                f"👤 user: `{req.user_id}`\n"
                f"{src} | 📚 {req.fan_name} V{req.variant_num}\n"
                f"❓ {len(req.questions)} savol | 🕐 {format_wait(elapsed)}\n"
                f"📱 Akkaunt: `{phone}`\n"
                f"🔗 {url}"
            )
            await bot_client.send_message(
                req.chat_id,
                f"✅ **Quiz tayyor!**\n\n"
                f"📚 {req.fan_name} — Variant {req.variant_num}\n"
                f"❓ {len(req.questions)} savol\n"
                f"⏱ {tl.get(req.time_choice, req.time_choice)} | "
                f"🔀 {'Aralash' if req.order_choice=='shuffle' else 'Ketma-ket'}\n"
                f"🕐 {format_wait(elapsed)}\n\n"
                f"🔗 {url}"
            )
        else:
            refund = calc_file_price(len(req.questions))
            db_add_balance(req.user_id, refund, f"Qaytarildi: quiz V{req.variant_num} xato")
            bal_left = db_get_balance(req.user_id)
            await bot_client.send_message(
                req.chat_id,
                f"❌ **Quiz yaratishda xato!**\n\n"
                f"Havola olinmadi — @QuizBot javob bermadi.\n"
                f"💰 **{refund:,} so'm qaytarildi** | Balans: {bal_left:,} so'm\n\n"
                f"Qayta urinib ko'ring:"
            )
    except Exception as e:
        if progress_task:
            progress_task.cancel()
        try:
            if progress_msg:
                if hasattr(progress_msg, "delete"):
                    await progress_msg.delete()
                else:
                    await bot_client.delete_messages(req.chat_id, progress_msg)
        except Exception:
            pass
        try:
            refund = calc_file_price(len(req.questions))
            db_add_balance(req.user_id, refund, f"Qaytarildi: xato — {str(e)[:50]}")
            bal_left = db_get_balance(req.user_id)
            await bot_client.send_message(
                req.chat_id,
                f"❌ **Xato yuz berdi!**\n\n`{e}`\n\n"
                f"💰 **{refund:,} so'm qaytarildi** | Balans: {bal_left:,} so'm"
            )
        except Exception as e2:
            log.error(f"Qaytarish xatosi: {e2}")
        log.error(f"run_request xato: {e}")
    finally:
        release(userbot)

# ============================================================
#  MAIN
# ============================================================
async def main():
    global bot_client

    bot_client = TelegramClient(
        _os.path.join(_os.path.dirname(DB_FILE), "bot_session"),
        API_ID, API_HASH
    )
    await bot_client.start(bot_token=BOT_TOKEN)
    log.info("Bot ulandi!")

    # DB ishga tushirish
    db_init()
    log.info(f"DB tayyor: {DB_FILE} | Jami users: {db_count_users()} ta")

    for phone in load_phones():
        try:
            sess_dir = _os.path.dirname(DB_FILE)
            session  = _os.path.join(sess_dir, f"userbot_{phone.replace('+','').replace(' ','')}")

            # DB dan sessiyani tiklash (sessiya fayli yo'q bo'lsa)
            if not _os.path.exists(session + ".session"):
                if db_load_session(phone, session):
                    log.info(f"Sessiya DB dan tiklandi: {phone}")

            client = TelegramClient(session, API_ID, API_HASH)

            async def password_input():
                pwd = input(f"🔐 2FA paroli ({phone}): ")
                return pwd

            await client.start(phone=phone, password=password_input)

            # Sessiyani DB ga saqlash — keyingi restart uchun
            db_save_session(phone, session)

            all_clients.append(client)
            account_phones[id(client)] = phone

            if phone == NOTIFY_PHONE:
                log.info(f"Notify akkaunt ulandi: {phone}")
            else:
                account_pool.append(client)
                account_busy[id(client)] = False
                log.info(f"Quiz akkaunt ulandi: {phone}")
        except Exception as e:
            log.error(f"Ulanmadi {phone}: {e}")

    # ============================================================
    #  KNOPKALAR
    # ============================================================
    def main_menu(adm=False, uid=None, force_partner: bool = False):
        # Ixcham asosiy menu.
        # MUHIM: har safar DB dagi real hamkor statusiga qarab qayta yaratiladi.
        # force_partner=True admin tasdiqlagan zahoti userga yangi menu yuborish uchun.
        is_partner_now = bool(force_partner or (db_is_partner(uid) if uid else False))

        partner_btn = "🤝 Hamkor paneli" if is_partner_now else "🤝 Hamkor bo'lish"

        btns = [
            [
                Button.text("📂 Fayldan quiz yaratish", resize=True),
                Button.text("📋 Mening quizlarim", resize=True),
            ],
            [
                Button.text("❓ Yordam", resize=True),
                Button.text(partner_btn, resize=True),
            ],
            [
                Button.text("👤 Profil", resize=True),
            ],
        ]

        if adm:
            btns.append([Button.text("🔧 Admin panel", resize=True)])

        return btns

    async def force_refresh_main_menu(user_id: int, text: str = "🏠 Menu yangilandi"):
        # Telegram reply keyboard ba'zan cache bo'lib qoladi.
        # Shuning uchun avval eski keyboardni butunlay olib tashlaymiz, keyin yangisini yuboramiz.
        try:
            await bot_client.send_message(user_id, "🔄 Menu yangilanmoqda...", buttons=Button.clear())
            await asyncio.sleep(0.35)
        except Exception as e:
            log.warning(f"Menu clear xato: {e}")
        try:
            await bot_client.send_message(
                user_id,
                text,
                buttons=main_menu(is_admin(user_id), user_id, force_partner=db_is_partner(user_id))
            )
        except Exception as e:
            log.warning(f"Menu refresh xato: {e}")


    # ============================================================
    #  FAYL SHABLONI BO'YICHA QO'LLANMA MATNLARI
    # ============================================================
    def file_quiz_guide_text():
        return """
📂 **Fayldan quiz yaratish**

Faylingiz botga to'g'ri tushishi uchun avval testlaringizni AI orqali bot shabloniga moslab oling.

✅ **Nima qilish kerak?**
1. Test faylingizni ChatGPT, Claude, Gemini yoki Groq AI ga yuboring.
2. Pastdagi promptni ham birga yuboring.
3. AI qaytargan natijani `.txt` yoki `.docx` fayl qilib saqlang.
4. Shu faylni botga yuboring.

━━━━━━━━━━━━━━━
📋 **AI ga yuboriladigan tayyor prompt:**

```
Ushbu fayldagi BARCHA savollarni o'qib chiq.
Savollarni aslo qisqartirma, o'zgartirma yoki tashlab ketma.

Agar savollarda tayyor variantlar bo'lsa, ularni saqlab qol.
Agar variantlar bo'lmasa, har bir savol uchun 4 ta mantiqiy variant yarat.

Natijani FAQAT quyidagi shablonda qaytar:

Savol matni
=====
Variant 1
=====
#To'g'ri javob
=====
Variant 3
=====
Variant 4
+++++

MUHIM QOIDALAR:
- Har bir savolda faqat 1 ta to'g'ri javob bo'lsin.
- To'g'ri javob oldiga albatta # belgisi qo'yilsin.
- Variantlar orasida ===== belgisi bo'lsin.
- Har bir savol oxirida +++++ belgisi bo'lsin.
- Hech qanday izoh, kirish so'zi yoki xulosa yozma.
- Kod bloki ishlatma, oddiy matn ko'rinishida qaytar.
- Savollar sonini kamaytirma.
```

━━━━━━━━━━━━━━━
✅ **Bot qabul qiladigan oddiy namuna:**

```
O'zbekiston poytaxti qaysi shahar?
=====
Samarqand
=====
#Toshkent
=====
Buxoro
=====
Xiva
+++++
```

📌 **Narx:** Har 25 savolga 1 500 so'm
_(50 savol = 3 000, 100 savol = 6 000)_

Endi tayyorlangan DOCX, PDF yoki TXT faylni shu yerga yuboring.
"""

    def bad_template_guide_text(reason=""):
        reason_line = f"\n📌 Sabab: {reason}\n" if reason else ""
        return (
            "⚠️ **Fayl shablonga mos kelmadi**\n"
            f"{reason_line}\n"
            "Bot quiz yaratishi uchun fayl ichida savol, variantlar va to'g'ri javob `#` belgisi bilan aniq ko'rsatilgan bo'lishi kerak.\n\n"
            "✅ **Nima qiling?**\n"
            "1. Test faylingizni ChatGPT, Claude, Gemini yoki Groq AI ga yuboring.\n"
            "2. Quyidagi promptni ham birga yuboring.\n"
            "3. AI qaytargan matnni `.txt`, `.docx` yoki `.pdf` qilib botga qayta yuboring.\n\n"
            "📋 **AI uchun tayyor prompt:**\n\n"
            "```\n"
            "Ushbu fayldagi BARCHA savollarni o'qib chiq.\n"
            "Savollarni aslo qisqartirma, o'zgartirma yoki tashlab ketma.\n\n"
            "Agar savollarda tayyor variantlar bo'lsa, ularni saqlab qol.\n"
            "Agar variantlar bo'lmasa, har bir savol uchun 4 ta mantiqiy variant yarat.\n\n"
            "Natijani FAQAT quyidagi shablonda qaytar:\n\n"
            "Savol matni\n"
            "=====\n"
            "Variant 1\n"
            "=====\n"
            "#To'g'ri javob\n"
            "=====\n"
            "Variant 3\n"
            "=====\n"
            "Variant 4\n"
            "+++++\n\n"
            "MUHIM QOIDALAR:\n"
            "- Har bir savolda faqat 1 ta to'g'ri javob bo'lsin.\n"
            "- To'g'ri javob oldiga albatta # belgisi qo'yilsin.\n"
            "- Variantlar orasida ===== belgisi bo'lsin.\n"
            "- Har bir savol oxirida +++++ belgisi bo'lsin.\n"
            "- Hech qanday izoh yozma.\n"
            "- Kod bloki ishlatma, oddiy matn ko'rinishida qaytar.\n"
            "```"
        )

    # ============================================================
    #  OVOZLI QO'LLANMALAR — maxfiy kanal orqali yuborish
    # ============================================================
    VOICE_CHANNEL_ID = int(_os.environ.get("VOICE_CHANNEL_ID", "-1003984400731"))
    START_VOICE_MSG_ID = int(_os.environ.get("START_VOICE_MSG_ID", "2"))
    TEMPLATE_VOICE_MSG_ID = int(_os.environ.get("TEMPLATE_VOICE_MSG_ID", "3"))
    PROGRESS_VOICE_MSG_ID = int(_os.environ.get("PROGRESS_VOICE_MSG_ID", "4"))

    async def send_voice_from_channel(chat_id: int, msg_id: int):
        if not msg_id:
            return
        try:
            await bot_client.forward_messages(
                entity=chat_id,
                messages=msg_id,
                from_peer=VOICE_CHANNEL_ID
            )
        except Exception as e:
            log.warning(f"Voice kanal orqali yuborishda xato: {e}")

    async def send_start_voice(chat_id: int):
        await send_voice_from_channel(chat_id, START_VOICE_MSG_ID)

    async def send_template_voice(chat_id: int):
        await send_voice_from_channel(chat_id, TEMPLATE_VOICE_MSG_ID)

    async def send_progress_voice(chat_id: int):
        await send_voice_from_channel(chat_id, PROGRESS_VOICE_MSG_ID)

    def ai_settings_btns(state: UserState):
        topic_show = state.topic if state.topic else "Barcha mavzu"
        return [
            [Button.text("📝 Fan nomini o'zgartirish"),
             Button.text("📌 Mavzuni o'zgartirish")],
            [Button.text("🔢 5 ta"),  Button.text("🔢 10 ta"),
             Button.text("🔢 15 ta"), Button.text("🔢 20 ta"),
             Button.text("🔢 25 ta")],
            [Button.text("🟢 Oson"), Button.text("🟡 O'rta"), Button.text("🔴 Qiyin")],
            [Button.text("🇺🇿 O'zbek"), Button.text("🇷🇺 Rus"), Button.text("🇬🇧 Ingliz")],
            [Button.text(f"✅ Yaratish — {state.q_count} ta savol")],
            [Button.text("🔙 Bosh menyu")],
        ]

    def time_btns():
        return [
            [Button.text("⏱ 15s"), Button.text("⏱ 30s")],
            [Button.text("⏱ 60s"), Button.text("⏱ Chegarasiz")],
        ]

    def order_btns():
        return [[Button.text("📋 Ketma-ket"), Button.text("🔀 Aralash")]]

    def variant_btns(total):
        rows, row = [], []
        for n in [5, 10, 15, 20, 25, 30, 50]:
            if total // n >= 1:
                row.append(Button.text(f"{n} ta"))
                if len(row) == 4: rows.append(row); row = []
        if row: rows.append(row)
        rows.append([Button.text("Hammasi bitta quiz")])
        return rows

    def answer_btns(opts):
        letters = ["A","B","C","D","E","F"]
        rows, row = [], []
        for i, opt in enumerate(opts):
            row.append(Button.text(f"{letters[i] if i<6 else i+1}. {opt[:18]}"))
            if len(row) == 2: rows.append(row); row = []
        if row: rows.append(row)
        rows.append([Button.text("⏭ O'tkazib yuborish"), Button.text("🔙 Bosh menyu")])
        return rows

    # ============================================================
    #  HANDLERLAR
    # ============================================================

    @bot_client.on(events.NewMessage(pattern=r"/addpartner(?:\s+(\d+))?"))
    async def cmd_addpartner(event):
        if event.sender_id not in ADMIN_IDS:
            return
        match = re.match(r'/addpartner\s+(\d+)', event.raw_text.strip())
        if not match:
            await event.respond("❌ Ishlatish: `/addpartner USER_ID`")
            return
        target_uid = int(match.group(1))
        db_add_partner(target_uid)
        await event.respond(f"✅ `{target_uid}` hamkor sifatida qo'shildi!")
        try:
            bot_me3 = await bot_client.get_me()
            plink = f"https://t.me/{bot_me3.username}?start=ref_{target_uid}"

            # 1) Eski keyboardni majburan olib tashlaymiz
            await bot_client.send_message(
                target_uid,
                "🔄 Menu yangilanmoqda...",
                buttons=Button.clear()
            )
            await asyncio.sleep(0.35)

            # 2) Yangi keyboardni majburan yuboramiz: Hamkor bo'lish -> Hamkor paneli
            await bot_client.send_message(
                target_uid,
                f"🎉 **Tabriklaymiz! Siz hamkor bo'ldingiz!**\n\n"
                f"🔗 Sizning shaxsiy havolangiz:\n`{plink}`\n\n"
                f"📢 Bu havolani kanal, guruh va ijtimoiy tarmoqlarda ulashing!\n\n"
                f"💰 Har jalb: +{PARTNER_JOIN_BONUS} so'm\n"
                f"💳 Har to'lovdan: {PARTNER_PAY_PERCENT}%\n\n"
                f"✅ Pastdagi menyuda endi **Hamkor paneli** tugmasi chiqadi 👇",
                buttons=main_menu(is_admin(target_uid), target_uid, force_partner=True)
            )
        except Exception as e:
            await event.respond(f"⚠️ Foydalanuvchiga xabar yuborilmadi: {e}")

    @bot_client.on(events.NewMessage(pattern=r"/fixmenu"))
    async def cmd_fixmenu(event):
        uid = event.sender_id
        await event.respond("🔄 Menu yangilanmoqda...", buttons=Button.clear())
        await asyncio.sleep(0.35)
        await event.respond("🏠 Bosh menu", buttons=main_menu(is_admin(uid), uid, force_partner=db_is_partner(uid)))

    async def user_is_subscribed(user_id: int) -> bool:
        settings = db_get_subscription_settings()
        if not settings.get("enabled"):
            return True
        chat_ref = (settings.get("chat_ref") or "").strip()
        if not chat_ref:
            return True
        try:
            entity = await bot_client.get_entity(int(chat_ref) if re.fullmatch(r'-?\d+', chat_ref) else chat_ref)
            await bot_client(GetParticipantRequest(entity, user_id))
            return True
        except UserNotParticipantError:
            return False
        except Exception as e:
            log.warning(f"Majburiy a'zolik tekshiruvi xato: {e}")
            return False

    async def show_subscription_prompt(target_event, uid: int):
        settings = db_get_subscription_settings()
        title = settings.get("title") or settings.get("chat_ref") or "kanal/guruh"
        link = (settings.get("invite_link") or "").strip()
        buttons = []
        if link:
            buttons.append([Button.url("📢 Kanal/guruhga qo'shilish", link)])
        buttons.append([Button.inline("✅ A'zolikni tekshirish", b"check_subscription")])
        await target_event.respond(
            f"📢 **Majburiy a'zolik yoqilgan**\n\n"
            f"Botdan foydalanish va bonus olish uchun avval **{title}** ga a'zo bo'ling.\n\n"
            f"A'zo bo'lgach, pastdagi **Tekshirish** tugmasini bosing.",
            buttons=buttons
        )

    async def finalize_user_access(uid: int, first_name: str = "", full_name: str = "", uname: str = ""):
        inviter_id = db_get_invited_by(uid)
        bonus_msg = ""
        if inviter_id and inviter_id != uid and not db_already_referred(inviter_id, uid):
            db_save_referral(inviter_id, uid)
            db_add_balance(uid, REFERRAL_BONUS, f"Referal bonusi — {inviter_id} taklif qildi")
            db_add_balance(inviter_id, REFERRAL_BONUS, f"Referal bonusi — {uid} qo'shildi")
            ref_count = db_get_referral_count(inviter_id)
            bonus_msg += f"\n\n🎁 **Referal bonusi: +{REFERRAL_BONUS:,} so'm** balansga qo'shildi!"
            if db_is_partner(inviter_id):
                db_add_partner_balance(inviter_id, PARTNER_JOIN_BONUS, f"Yangi foydalanuvchi: {uid}")
                db_partner_add_ref(inviter_id)
            try:
                await bot_client.send_message(
                    inviter_id,
                    f"🎉 **Yangi referal!**\n\n"
                    f"👤 **{full_name or first_name or uid}** sizning havolangiz orqali qo'shildi!\n"
                    f"💰 +{REFERRAL_BONUS:,} so'm balansga qo'shildi\n"
                    f"👥 Jami referallar: **{ref_count} ta**"
                )
            except Exception:
                pass
            await notify_admin(
                f"🎁 **Referal**\n\n"
                f"👤 {full_name or uid} (`{uid}`) → `{inviter_id}` havolasidan keldi\n"
                f"💰 Ikkalasiga +{REFERRAL_BONUS:,} so'm"
            )
        return bonus_msg

    @bot_client.on(events.NewMessage(pattern="/start"))
    async def cmd_start(event):
        uid = event.sender_id
        is_new = db_is_new_user(uid)
        user_states[uid] = UserState()

        sender = await event.get_sender()
        first  = getattr(sender, 'first_name', '') or ''
        last   = getattr(sender, 'last_name',  '') or ''
        uname  = getattr(sender, 'username',   '') or ''
        full_name = f"{first} {last}".strip() or uname or str(uid)
        db_save_user(user_id=uid, first_name=first, last_name=last, username=uname)
        track_user(uid, full_name, "idle", "/start")

        # Referalni hozircha faqat eslab qolamiz.
        # Bonus faqat foydalanuvchi shartlarga rozilik berib,
        # majburiy a'zolik bo'lsa kanal/guruhga ham a'zo bo'lgandan keyin beriladi.
        ref_bonus_msg = ""
        raw = event.raw_text.strip()
        ref_match = re.match(r'^/start\s+ref_(\d+)$', raw)
        if ref_match and is_new:
            inviter_id = int(ref_match.group(1))
            if inviter_id != uid:
                db_set_invited_by(uid, inviter_id)


        # Yangi foydalanuvchi bo'lsa admin ga xabar
        if is_new:
            total = db_count_users()
            uname_str = f"@{uname}" if uname else f"`{uid}`"
            await notify_admin(
                f"👤 **Yangi foydalanuvchi**\n\n"
                f"Ism: **{full_name}**\n"
                f"ID: `{uid}` | {uname_str}\n"
                f"Jami: {total} ta"
            )

        # Shartlarga rozilik tekshirish
        if not db_is_agreed(uid):
            # Shartlar havolalari
            if WEB_BASE_URL:
                terms_url   = f"{WEB_BASE_URL}/terms"
                privacy_url = f"{WEB_BASE_URL}/privacy"
                oferta_url  = f"{WEB_BASE_URL}/oferta"
                doc_buttons = [
                    [Button.url("📋 Foydalanish shartlari", terms_url)],
                    [Button.url("🔒 Maxfiylik siyosati",    privacy_url)],
                    [Button.url("📄 Ommaviy oferta",         oferta_url)],
                ]
            else:
                doc_buttons = []
            agree_buttons = doc_buttons + [
                [Button.inline("✅ Roziman, botdan foydalanaman", b"agree_terms")],
            ]
            await event.respond(
                f"👋 **Salom, {first or 'do\'stim'}!**\n\n"
                f"🤖 **AI Quiz Bot** ga xush kelibsiz!\n\n"
                f"Botdan foydalanishdan oldin quyidagi hujjatlarni o'qib, "
                f"rozilik bildiring:\n\n"
                f"📋 Foydalanish shartlari\n"
                f"🔒 Maxfiylik siyosati\n"
                f"📄 Ommaviy oferta (YATT shartnomasi)\n\n"
                f"_Hujjatlarni o'qib chiqqach, \"✅ Roziman\" tugmasini bosing._",
                buttons=agree_buttons
            )
            return

        settings = db_get_subscription_settings()
        if settings.get("enabled") and not await user_is_subscribed(uid):
            await show_subscription_prompt(event, uid)
            return

        await event.respond(
            f"👋 **Salom! AI Quiz Bot**\n\n"
            f"🤖 AI yordamida istalgan fandan test tuzing!\n"
            f"📁 Fayl yuklang yoki matn kiriting{ref_bonus_msg}\n\n"
            f"Boshlash uchun tugmani bosing 👇",
            buttons=main_menu(is_admin(uid), uid)
        )
        await send_start_voice(uid)


    @bot_client.on(events.CallbackQuery(data=b"agree_terms"))
    async def on_agree_terms(event):
        """Foydalanuvchi shartlarga roziligi."""
        uid = event.sender_id
        db_set_agreed(uid)

        sender = await event.get_sender()
        first  = getattr(sender, "first_name", "") or ""
        last   = getattr(sender, "last_name",  "") or ""
        uname  = getattr(sender, "username",   "") or ""
        full_name = f"{first} {last}".strip() or uname or str(uid)
        db_save_user(user_id=uid, first_name=first, last_name=last, username=uname)

        # Adminga xabar
        uname_str = f"@{uname}" if uname else f"`{uid}`"
        await notify_admin(
            f"✅ **Yangi foydalanuvchi (rozilik berdi)**\n\n"
            f"Ism: **{full_name}**\n"
            f"ID: `{uid}` | {uname_str}"
        )

        settings = db_get_subscription_settings()
        if settings.get("enabled"):
            try:
                await event.edit(
                    f"✅ **Rozilik tasdiqlandi!**\n\n"
                    f"Keyingi qadam: majburiy kanal/guruhga a'zo bo'ling."
                )
            except Exception:
                pass
            await show_subscription_prompt(event, uid)
            return

        bonus_msg = await finalize_user_access(uid, first, full_name, uname)

        try:
            await event.edit(
                f"✅ **Rozilik tasdiqlandi!**\n\n"
                f"Xush kelibsiz, **{first or full_name}**! 🎉{bonus_msg}\n\n"
                f"Endi botdan to'liq foydalanishingiz mumkin."
            )
        except Exception:
            pass

        await bot_client.send_message(
            uid,
            f"🏠 **Asosiy menyu**\n\n"
            f"🤖 AI yordamida istalgan fandan test tuzing!\n"
            f"📁 Fayl yuklang yoki matn kiriting{bonus_msg}",
            buttons=main_menu(is_admin(uid), uid)
        )
        await send_start_voice(uid)

    @bot_client.on(events.CallbackQuery(data=b"check_subscription"))
    async def on_check_subscription(event):
        uid = event.sender_id
        if not db_is_agreed(uid):
            await event.answer("Avval shartlarga rozilik bering.", alert=True)
            return
        if await user_is_subscribed(uid):
            sender = await event.get_sender()
            first = getattr(sender, "first_name", "") or ""
            last = getattr(sender, "last_name", "") or ""
            uname = getattr(sender, "username", "") or ""
            full_name = f"{first} {last}".strip() or uname or str(uid)
            bonus_msg = await finalize_user_access(uid, first, full_name, uname)
            try:
                await event.edit(
                    f"✅ **A'zolik tasdiqlandi!**\n\n"
                    f"Endi botdan to'liq foydalanishingiz mumkin.{bonus_msg}"
                )
            except Exception:
                pass
            await bot_client.send_message(
                uid,
                f"🏠 **Asosiy menyu**\n\n"
                f"🤖 AI yordamida istalgan fandan test tuzing!\n"
                f"📁 Fayl yuklang yoki matn kiriting.",
                buttons=main_menu(is_admin(uid), uid)
            )
            await send_start_voice(uid)
        else:
            await event.answer("Hali a'zo bo'lmagansiz yoki bot tekshira olmadi.", alert=True)
            await show_subscription_prompt(event, uid)

    @bot_client.on(events.NewMessage(func=lambda e: e.file))
    async def on_file(event):
        uid = event.sender_id
        adm = is_admin(uid)
        log.info(f"Fayl keldi: user={uid}")

        # Log guruhiga xabar
        sender = await event.get_sender()
        first = getattr(sender, 'first_name', '') or ''
        last  = getattr(sender, 'last_name',  '') or ''
        uname = getattr(sender, 'username',   '') or ''
        full_name = f"{first} {last}".strip() or uname or str(uid)
        fname_attr = getattr(event.file, 'name', None) or getattr(event.file, 'mime_type', 'fayl')
        await log_action(uid, uname, full_name, f"📎 Fayl yubordi: {fname_attr}")

        # Faqat fayl kutilayotgan holatlarda davom etamiz
        # Boshqa admin holatlarda (masalan wait_phone) — ignore
        astate_step = admin_states.get(uid, {}).get("step", "")

        if not adm:
            if not db_is_agreed(uid):
                await event.respond("⚠️ Avval shartlarga rozilik bering.")
                return
            settings = db_get_subscription_settings()
            if settings.get("enabled") and not await user_is_subscribed(uid):
                await show_subscription_prompt(event, uid)
                return

        # Broadcast — fayl yuborganda ham ishlashi uchun
        if astate_step == "wait_broadcast" and is_admin(uid):
            admin_states.pop(uid, None)
            all_users = get_db().execute("SELECT user_id FROM users").fetchall()
            get_db().close()
            ok, fail = 0, 0
            prog = await event.respond(f"📤 Yuborilmoqda... 0/{len(all_users)}")
            for i, (target_id,) in enumerate(all_users):
                try:
                    await bot_client.forward_messages(target_id, event.message)
                    ok += 1
                except Exception as e:
                    log.warning(f"Broadcast xato user={target_id}: {e}")
                    fail += 1
                if (i + 1) % 20 == 0:
                    try:
                        await prog.edit(f"📤 Yuborilmoqda... {i+1}/{len(all_users)}")
                    except Exception:
                        pass
            await prog.edit(f"✅ Broadcast tugadi!\n✅ {ok} ta | ❌ {fail} ta")
            await event.respond("🔙", buttons=[[Button.text("🔙 Admin panel")]])
            return

        # User ID ga fayl yuborish
        if astate_step == "wait_user_msg_text" and is_admin(uid):
            target_id = admin_states[uid].get("target_id")
            admin_states.pop(uid, None)
            try:
                await bot_client.send_message(int(target_id), text)
                await event.respond(f"✅ User `{target_id}` ga yuborildi!", buttons=[[Button.text("🔙 Admin panel")]])
            except Exception as e:
                await event.respond(f"❌ Xato: {e}", buttons=[[Button.text("🔙 Admin panel")]])
            return

        if astate_step and astate_step not in ("wait_session_file", "wait_db_file"):
            log.info(f"Admin holati aktiv ({astate_step}), fayl ignore: user={uid}")
            return

        msg = await event.respond("📥 O'qilmoqda...")
        log.info(f"Fayl yuklanmoqda: user={uid}")

        try:
            import io

            buf = io.BytesIO()
            await event.download_media(file=buf)
            buf.seek(0)
            data = buf.read()
            log.info(f"Fayl yuklandi: {len(data)} bayt, user={uid}")

            if not data:
                await msg.edit("❌ Fayl bo'sh yoki yuklanmadi!")
                return

            # Kengaytma va MIME
            fname = ""
            mime  = ""
            try:
                fname = (event.file.name or "").lower()
                mime  = str(getattr(event.file, 'mime_type', '') or '')
            except Exception:
                pass

            log.info(f"Fayl: name={fname}, mime={mime}, user={uid}")

            supported_docx = fname.endswith(".docx") or "officedocument.wordprocessingml" in mime
            supported_pdf  = fname.endswith(".pdf") or "pdf" in mime
            supported_txt  = fname.endswith(".txt") or mime.startswith("text/") or mime in ("application/octet-stream", "")

            # Rasm/video/audio yoki boshqa fayllardan quiz tuzmaymiz.
            if not (supported_docx or supported_pdf or supported_txt):
                await msg.edit(
                    bad_template_guide_text(
                        "Bot faqat DOCX, PDF yoki TXT formatdagi matnli test faylini qabul qiladi. Rasm yuborilgan bo'lsa, avval rasm ichidagi savollarni AI/OCR orqali matnga aylantiring."
                    ),
                    buttons=[[Button.text("📂 Fayldan quiz yaratish")], [Button.text("🔙 Bosh menyu")]]
                )
                await send_template_voice(uid)
                return

            content = ""

            if fname.endswith(".docx") or "officedocument.wordprocessingml" in mime:
                try:
                    from docx import Document

                    def _read_docx_text(blob: bytes) -> str:
                        doc = Document(io.BytesIO(blob))
                        parts = []

                        for p in doc.paragraphs:
                            t = (p.text or "").strip()
                            if t:
                                parts.append(t)

                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    cell_text = "\n".join(
                                        (pp.text or "").strip()
                                        for pp in cell.paragraphs
                                        if (pp.text or "").strip()
                                    ).strip()
                                    if cell_text:
                                        parts.append(cell_text)

                        return "\n".join(parts).strip()

                    content = _read_docx_text(data)
                    log.info(f"DOCX o'qildi: {len(content)} belgi")
                except Exception as e:
                    log.error(f"DOCX xato: {e}")
                    await msg.edit(f"❌ DOCX o'qishda xato: {e}")
                    return

            elif fname.endswith(".pdf") or "pdf" in mime:
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(io.BytesIO(data))
                    content = "".join(p.extract_text() or "" for p in reader.pages)
                    log.info(f"PDF o'qildi: {len(content)} belgi")
                except Exception as e:
                    log.error(f"PDF xato: {e}")
                    await msg.edit(f"❌ PDF o'qishda xato: {e}")
                    return

            else:
                # TXT yoki boshqa
                content = data.decode("utf-8", errors="ignore")
                log.info(f"TXT o'qildi: {len(content)} belgi")

            if not content.strip():
                await msg.edit(
                    "❌ Fayldan matn o'qib bo'lmadi!\n\n"
                    "Qo'llab-quvvatlanadigan: **DOCX, PDF, TXT**"
                )
                return

            qs = _parse_questions(content)
            log.info(f"Parse natija: {len(qs)} savol, user={uid}")

            if qs:
                q_count = len(qs)
                price   = calc_file_price(q_count)
                bal     = db_get_balance(uid)
                blocks  = (q_count + 24) // 25

                # Savollarni RAM da saqlaymiz
                state = UserState(
                    step="file_preview",
                    questions=qs,
                    total_questions=q_count
                )
                user_states[uid] = state
                log.info(f"State saqlandi: step=file_preview, {q_count} savol, user={uid}")

                # ── 5 ta random savol QuizBot ga yuborib preview ko'rsatish ──
                await msg.edit(
                    f"📂 **{q_count} ta savol topildi!**\n\n"
                    f"⏳ Namuna sifatida 5 ta savol quiz qilinmoqda..."
                )

                preview_count = min(5, q_count)
                preview_qs = random.sample(qs, preview_count)

                # Preview quizni yuboramiz (bepul, haqiqiy akkaunt bilan)
                try:
                    if account_pool:
                        userbot = await get_free()
                        preview_req = QuizRequest(
                            user_id=uid, chat_id=event.chat_id,
                            questions=preview_qs,
                            fan_name="📋 Namuna ko'rish",
                            variant_num=0,
                            time_choice="30",
                            order_choice="order",
                            total_variants=1,
                            source="preview",
                        )
                        # Preview ni fon task sifatida yuboramiz
                        asyncio.create_task(_send_preview(userbot, preview_req, uid, event.chat_id,
                                                          q_count, price, bal, blocks))
                    else:
                        # Akkaunt yo'q — to'g'ridan narx ko'rsatamiz
                        await _show_file_price(event.chat_id, uid, q_count, price, bal, blocks)
                except Exception as pe:
                    log.error(f"Preview xato: {pe}")
                    await _show_file_price(event.chat_id, uid, q_count, price, bal, blocks)

            else:
                # Shablon topilmadi — foydalanuvchiga tushunarli qo'llanma va AI prompt beramiz.
                lines = [l.strip() for l in content.splitlines() if l.strip()]
                if not lines:
                    await msg.edit("❌ Faylda matn topilmadi!")
                    return

                user_states.pop(uid, None)
                reason = f"{len(lines)} ta qator topildi, lekin savol/variant/#to'g'ri javob formatini aniqlay olmadim."
                await msg.edit(
                    bad_template_guide_text(reason),
                    buttons=[[Button.text("📂 Fayldan quiz yaratish")], [Button.text("🔙 Bosh menyu")]]
                )
                await send_template_voice(uid)
                return

        except Exception as e:
            log.error(f"on_file xato: {e}", exc_info=True)
            try:
                await msg.edit(f"❌ Xato: {e}")
            except Exception:
                await event.respond(f"❌ Xato: {e}", buttons=main_menu(adm, uid))

    @bot_client.on(events.NewMessage(func=lambda e: not e.file and not e.text.startswith("/")))
    async def on_msg(event):
        uid = event.sender_id
        text = event.text.strip()
        adm = is_admin(uid)
        astate = admin_states.get(uid, {})

        # Har xabarda last_seen yangilash
        sender = await event.get_sender()
        first  = getattr(sender, 'first_name', '') or ''
        last   = getattr(sender, 'last_name',  '') or ''
        uname  = getattr(sender, 'username',   '') or ''
        full_name = f"{first} {last}".strip() or uname or str(uid)
        db_save_user(user_id=uid, first_name=first, last_name=last, username=uname)

        # Faol foydalanuvchini kuzatish
        state_now = user_states.get(uid, UserState())
        track_user(uid, full_name, state_now.step, text[:50])

        # User harakatini log guruhiga yuborish
        await log_action(uid, uname, full_name, f"✉️ {text[:150]}")

        # Admin oraliq holat
        if astate.get("step") == "wait_phone":
            await _admin_add_phone(event, uid, text); return
        if astate.get("step") == "wait_code":
            await _admin_enter_code(event, uid, text); return
        if astate.get("step") == "wait_password":
            await _admin_enter_pass(event, uid, text); return
        if astate.get("step") == "wait_remove":
            await _admin_do_remove(event, uid, text); return
        if astate.get("step") == "wait_bonus_user_id":
            await _admin_bonus_user_id(event, uid, text); return
        if astate.get("step") == "wait_bonus_amount":
            await _admin_bonus_amount(event, uid, text); return
        if astate.get("step") == "wait_subscription_target":
            if not adm:
                return
            if text == "🔙 Admin panel":
                admin_states.pop(uid, None)
                await _show_admin(event)
                return
            parts = [x.strip() for x in text.split("|")]
            if len(parts) < 2 or not parts[0] or not parts[1]:
                await event.respond(
                    "❌ Format noto'g'ri.\n\nYuboring:\n`chat_ref | invite_link | title`",
                    buttons=[[Button.text("🔙 Admin panel")]]
                )
                return
            chat_ref = parts[0]
            invite_link = parts[1]
            title = parts[2] if len(parts) >= 3 else ""
            db_set_subscription_target(chat_ref, invite_link, title)
            admin_states.pop(uid, None)
            await event.respond(
                f"✅ Kanal/guruh yangilandi!\n\nChat ref: `{chat_ref}`\nHavola: {invite_link}\nNomi: {title or '-'}",
                buttons=[[Button.text("📢 Majburiy a'zolik")], [Button.text("🔙 Admin panel")]]
            )
            return

        # ---- BROADCAST ----
        if astate.get("step") == "wait_broadcast":
            if not adm: return
            if text == "❌ Bekor qilish":
                admin_states.pop(uid, None)
                await _show_admin(event); return
            admin_states.pop(uid, None)
            all_users = get_db().execute("SELECT user_id FROM users").fetchall()
            get_db().close()
            ok, fail = 0, 0
            broadcast_text = event.message.text or text
            prog = await event.respond(f"📤 Yuborilmoqda... 0/{len(all_users)}")
            for i, (target_id,) in enumerate(all_users):
                try:
                    await bot_client.send_message(int(target_id), broadcast_text)
                    ok += 1
                except Exception as e:
                    log.warning(f"Broadcast xato user={target_id}: {e}")
                    fail += 1
                if (i + 1) % 20 == 0:
                    try:
                        await prog.edit(f"📤 Yuborilmoqda... {i+1}/{len(all_users)}")
                    except Exception:
                        pass
            await prog.edit(
                f"✅ Broadcast tugadi!\n\n"
                f"✅ Muvaffaqiyatli: {ok} ta\n"
                f"❌ Xato: {fail} ta"
            )
            await event.respond("🔙", buttons=[[Button.text("🔙 Admin panel")]])
            return

        # ---- USER ID GA XABAR ----
        if astate.get("step") == "wait_user_id_msg":
            if not adm: return
            if text == "❌ Bekor qilish":
                admin_states.pop(uid, None)
                await _show_admin(event); return
            try:
                target_id = int(text.strip())
                admin_states[uid] = {"step": "wait_user_msg_text", "target_id": target_id}
                await event.respond(
                    f"📨 User ID: `{target_id}`\n\nYubormoqchi bo'lgan xabarni yozing:",
                    buttons=[[Button.text("❌ Bekor qilish")]]
                )
            except ValueError:
                await event.respond("❌ Noto'g'ri ID! Faqat raqam yozing:", buttons=[[Button.text("❌ Bekor qilish")]])
            return

        if astate.get("step") == "wait_user_msg_text":
            if not adm: return
            if text == "❌ Bekor qilish":
                admin_states.pop(uid, None)
                await _show_admin(event); return
            target_id = astate.get("target_id")
            admin_states.pop(uid, None)
            try:
                await bot_client.send_message(int(target_id), text)
                await event.respond(f"✅ User `{target_id}` ga xabar yuborildi!", buttons=[[Button.text("🔙 Admin panel")]])
            except Exception as e:
                await event.respond(f"❌ Xato: {e}", buttons=[[Button.text("🔙 Admin panel")]])
            return

        state = user_states.get(uid, UserState())

        # Shartlarga rozilik tekshirish — admin va /start bundan mustasno
        if not is_admin(uid) and not db_is_agreed(uid) and text not in ("🔙 Bosh menyu",):
            if WEB_BASE_URL:
                agree_buttons = [
                    [Button.url("📋 Foydalanish shartlari", f"{WEB_BASE_URL}/terms")],
                    [Button.url("🔒 Maxfiylik siyosati",    f"{WEB_BASE_URL}/privacy")],
                    [Button.url("📄 Ommaviy oferta",         f"{WEB_BASE_URL}/oferta")],
                    [Button.inline("✅ Roziman, botdan foydalanaman", b"agree_terms")],
                ]
            else:
                agree_buttons = [
                    [Button.inline("✅ Roziman, botdan foydalanaman", b"agree_terms")],
                ]
            await event.respond(
                "⚠️ Botdan foydalanish uchun avval shartlarga rozilik bildiring:",
                buttons=agree_buttons
            )
            return

        if not is_admin(uid):
            settings = db_get_subscription_settings()
            if settings.get("enabled") and not await user_is_subscribed(uid):
                await show_subscription_prompt(event, uid)
                return

        # ---- BOSH MENYU KNOPKALARI ----
        if text == "🔙 Bosh menyu":
            user_states[uid] = UserState()
            admin_states.pop(uid, None)
            await event.respond("🏠 Bosh menyu", buttons=main_menu(adm, uid))
            return

        if text == "📋 Mening quizlarim":
            quizzes = db_get_user_quizzes(uid, limit=20)
            total_q = db_count_user_quizzes(uid)
            if not quizzes:
                await event.respond(
                    "📋 **Mening quizlarim**\n\n"
                    "Hali quiz yaratmagansiz.\n\n"
                    "🤖 AI yoki 📂 fayl orqali quiz tuzing!",
                    buttons=main_menu(adm)
                )
                return

            lines = [f"📋 **Mening quizlarim ({total_q} ta)**\n"]
            for q in quizzes:
                qid, fan, q_count, variant, url, source, created = q
                src_icon = "🤖" if source == "ai" else "📂"
                date = created[:10] if created else ""
                lines.append(
                    f"{src_icon} **{fan}** V{variant} — {q_count} savol\n"
                    f"   📅 {date} | [▶️ Ochish]({url})"
                )

            await event.respond(
                "\n\n".join(lines),
                buttons=[[Button.text("🔙 Bosh menyu")]],
                link_preview=False
            )
            return

        if text == "🎁 Referal":
            ref_count = db_get_referral_count(uid)
            ref_list  = db_get_referral_list(uid, limit=10)
            me = await event.get_sender()
            bot_me = await bot_client.get_me()
            bot_username = bot_me.username
            ref_link = f"https://t.me/{bot_username}?start=ref_{uid}"
            bal = db_get_balance(uid)

            lines = [
                f"🎁 **Referal dasturi**\n",
                f"👥 Siz taklif qilganlar: **{ref_count} ta**",
                f"💰 Har bir referal uchun: **{REFERRAL_BONUS:,} so'm** (ikkalangizga)\n",
                f"🔗 **Sizning havolangiz:**",
                f"`{ref_link}`\n",
                f"📌 Do'stingizga shu havolani yuboring. U ro'yxatdan o'tganda ikkalangizga **{REFERRAL_BONUS:,} so'm** beriladi!",
            ]

            if ref_list:
                lines.append(f"\n👤 **So'nggi referallar:**")
                for r in ref_list:
                    r_id, r_first, r_last, r_uname, r_date = r
                    r_name = f"{r_first or ''} {r_last or ''}".strip() or r_uname or str(r_id)
                    r_date_short = r_date[:10] if r_date else ""
                    lines.append(f"  • {r_name} — {r_date_short}")

            await event.respond(
                "\n".join(lines),
                buttons=[[Button.text("🔙 Bosh menyu")]],
                link_preview=False
            )
            return

        # ---- HAMKOR BO'LISH ----
        if text == "🤝 Hamkor bo'lish":
            if db_is_partner(uid):
                await event.respond("✅ Siz allaqachon hamkorsiz! Menu yangilanmoqda...", buttons=Button.clear())
                await asyncio.sleep(0.35)
                await event.respond("🤝 Hamkor paneli menyuga qo'shildi 👇", buttons=main_menu(adm, uid, force_partner=True))
                return
            await event.respond(
                "🤝 **HAMKORLIK DASTURI**\n\n"
                "━━━━━━━━━━━━━━━━━━━\n"
                "📢 Botimizni reklama qilib daromad toping!\n\n"
                "💰 **Daromad tizimi:**\n"
                f"  • Har jalb qilgan foydalanuvchi: **+{PARTNER_JOIN_BONUS} so'm**\n"
                f"  • Har to'lovdan: **{PARTNER_PAY_PERCENT}%** komissiya\n\n"
                "🔗 **Qanday ishlaydi:**\n"
                "  1. Ariza qoldiring\n"
                "  2. Admin tasdiqlaydi\n"
                "  3. Maxsus havola olasiz\n"
                "  4. Havolani ijtimoiy tarmoqlarda ulashing\n"
                "  5. Balansni kartaga chiqaring (min 15 000 so'm)\n\n"
                "━━━━━━━━━━━━━━━━━━━\n"
                "📌 Ariza qoldirishni xohlaysizmi?",
                buttons=[
                    [Button.text("📝 Ariza qoldirish")],
                    [Button.text("🔙 Bosh menyu")],
                ]
            )
            return

        if text == "📝 Ariza qoldirish":
            if db_is_partner(uid):
                await event.respond("✅ Siz allaqachon hamkorsiz! Menu yangilanmoqda...", buttons=Button.clear())
                await asyncio.sleep(0.35)
                await event.respond("🤝 Hamkor paneli menyuga qo'shildi 👇", buttons=main_menu(adm, uid, force_partner=True))
                return
            if db_has_pending_application(uid):
                await event.respond(
                    "⏳ **Arizangiz ko'rib chiqilmoqda.**\n\nAdmin tez orada javob beradi!",
                    buttons=[[Button.text("🔙 Bosh menyu")]]
                )
                return
            user_states[uid] = UserState(step="wait_partner_comment")
            await event.respond(
                "📝 **Ariza qoldirish**\n\n"
                "O'zingiz haqingizda qisqacha yozing:\n"
                "_(Qayerda reklama qilmoqchisiz, auditoriyangiz qancha va h.k.)_",
                buttons=[[Button.text("🔙 Bosh menyu")]]
            )
            return

        if state.step == "wait_partner_comment":
            comment = text
            sender2 = await event.get_sender()
            fn = getattr(sender2, 'first_name', '') or ''
            ln = getattr(sender2, 'last_name', '') or ''
            un = getattr(sender2, 'username', '') or ''
            fn2 = f"{fn} {ln}".strip() or un or str(uid)
            db_save_partner_application(uid, un, fn2, comment)
            user_states[uid] = UserState()
            # Adminga xabar
            uname_str = f"@{un}" if un else f"ID: {uid}"
            await notify_admin(
                f"🤝 **Yangi hamkorlik arizasi**\n\n"
                f"👤 {fn2} ({uname_str})\n"
                f"🆔 `{uid}`\n\n"
                f"💬 Ariza matni:\n{comment}\n\n"
                f"✅ Tasdiqlash uchun: `/addpartner {uid}`"
            )
            await event.respond(
                "✅ **Arizangiz qabul qilindi!**\n\n"
                "Admin tez orada ko'rib chiqadi va javob beradi.\n"
                "Kuting 🙏",
                buttons=[[Button.text("🔙 Bosh menyu")]]
            )
            return

        # ---- HAMKOR PANELI ----
        if text == "🤝 Hamkor paneli":
            if not db_is_partner(uid):
                await event.respond("⛔ Siz hamkor emassiz!", buttons=main_menu(adm, uid))
                return
            pinfo = db_get_partner_info(uid)
            if not pinfo:
                await event.respond("❌ Ma'lumot topilmadi!", buttons=main_menu(adm, uid))
                return
            _, pbal, total_earned, total_refs = pinfo
            bot_me2 = await bot_client.get_me()
            plink = f"https://t.me/{bot_me2.username}?start=ref_{uid}"
            await event.respond(
                f"🤝 **HAMKOR PANELI**\n\n"
                f"━━━━━━━━━━━━━━━━━━━\n"
                f"💰 **Hamkor balansi:** {pbal:,} so'm\n"
                f"📈 **Jami daromad:** {total_earned:,} so'm\n"
                f"👥 **Jalb qilganlar:** {total_refs} ta\n"
                f"━━━━━━━━━━━━━━━━━━━\n\n"
                f"🔗 **Sizning havolangiz:**\n`{plink}`\n\n"
                f"📢 Bu havolani kanal, guruh yoki ijtimoiy tarmoqlarda ulashing!\n\n"
                f"💸 Minimal chiqarish: {PARTNER_MIN_WITHDRAW:,} so'm",
                buttons=[
                    [Button.text("💸 Pul chiqarish")],
                    [Button.text("🔙 Bosh menyu")],
                ]
            )
            return

        if text == "💸 Pul chiqarish":
            if not db_is_partner(uid):
                return
            pinfo = db_get_partner_info(uid)
            pbal = pinfo[1] if pinfo else 0
            if pbal < PARTNER_MIN_WITHDRAW:
                await event.respond(
                    f"❌ **Balans yetarli emas!**\n\n"
                    f"💰 Balansda: {pbal:,} so'm\n"
                    f"📌 Minimal: {PARTNER_MIN_WITHDRAW:,} so'm\n\n"
                    f"Ko'proq foydalanuvchi jalb qiling!",
                    buttons=[
                        [Button.text("🤝 Hamkor paneli")],
                        [Button.text("🔙 Bosh menyu")],
                    ]
                )
                return
            user_states[uid] = UserState(step="wait_partner_card")
            await event.respond(
                f"💸 **Pul chiqarish**\n\n"
                f"💰 Mavjud: {pbal:,} so'm\n\n"
                f"Karta raqamingizni yuboring:\n_(16 raqam, masalan: 8600 1234 5678 9012)_",
                buttons=[[Button.text("🔙 Bosh menyu")]]
            )
            return

        if state.step == "wait_partner_card":
            if not db_is_partner(uid):
                return
            card_input = re.sub(r'\s', '', text)
            if not re.match(r'^\d{16}$', card_input):
                await event.respond("❌ Noto'g'ri karta raqami! 16 ta raqam kiriting:")
                return
            pinfo = db_get_partner_info(uid)
            pbal = pinfo[1] if pinfo else 0
            formatted_card = ' '.join([card_input[i:i+4] for i in range(0, 16, 4)])
            ok = db_partner_withdraw(uid, pbal, formatted_card)
            user_states[uid] = UserState()
            if ok:
                sender3 = await event.get_sender()
                un3 = getattr(sender3, 'username', '') or ''
                ustr = f"@{un3}" if un3 else f"ID: {uid}"
                await notify_admin(
                    f"💸 **Hamkor chiqarish so'rovi**\n\n"
                    f"👤 {ustr} (`{uid}`)\n"
                    f"💰 Summa: **{pbal:,} so'm**\n"
                    f"💳 Karta: `{formatted_card}`\n\n"
                    f"Iltimos o'tkazing!"
                )
                await event.respond(
                    f"✅ **So'rov yuborildi!**\n\n"
                    f"💰 {pbal:,} so'm → `{formatted_card}`\n\n"
                    f"Admin tez orada o'tkazadi.",
                    buttons=[[Button.text("🔙 Bosh menyu")]]
                )
            else:
                await event.respond("❌ Xato yuz berdi!", buttons=[[Button.text("🔙 Bosh menyu")]])
            return

        if text == "❓ Yordam":
            await event.respond(
                "📋 **YORDAM**\n\n"
                "**🤖 AI test tuzish** — 2 000 so'm\n"
                "Istalgan fan va mavzudan AI avtomatik test tuzadi\n\n"
                "**📂 Fayldan quiz yaratish** — har 25 savolga 2 000 so'm\n"
                "Tayyor testingizni yuklang, bot quizga aylantiradi\n\n"
                "━━━━━━━━━━━━━━━\n"
                "**📌 Fayl formatlari:** DOCX, PDF, TXT\n\n"
                "**Shablon 1:**\n```\n1.Savol\na.Variant\n#b.To'g'ri\nc.Variant\n```\n\n"
                "**Shablon 2:**\n```\nSavol\na.Variant\n#b.To'g'ri\nc.Variant\n```\n\n"
                "**Shablon 3:**\n```\nSavol\n=====\n#To'g'ri\nVariant\n+++++\n```\n\n"
                "**# belgisi** = to'g'ri javob",
                buttons=main_menu(adm)
            )
            return

        if text == "🔧 Admin panel":
            if not adm: await event.respond("⛔ Admin emassiz!"); return
            await _show_admin(event); return

        if text == "📂 Fayldan quiz yaratish":
            user_states[uid] = UserState(step="wait_file")
            await event.respond(
                file_quiz_guide_text(),
                buttons=[[Button.text("🔙 Bosh menyu")]]
            )
            await send_template_voice(uid)
            return

        if text == "✏️ Matn kiritish":
            user_states[uid] = UserState(step="wait_text")
            await event.respond("✏️ Savollarni yuboring:\n/yordam — shablonlar",
                buttons=[[Button.text("🔙 Bosh menyu")]]); return

        # ---- AI TEST TUZISH ----
        if text == "🤖 AI test tuzish":
            state = UserState(step="ai_ask_fan")
            user_states[uid] = state
            await event.respond(
                "🤖 **AI Test Tuzish**\n\nQaysi fandan test kerak?\n\n"
                "_(Misol: Matematika, Fizika, Tarix, Python...)_",
                buttons=[[Button.text("🔙 Bosh menyu")]]
            ); return

        if state.step == "ai_ask_fan":
            state.fan_name = text
            state.step = "ai_ask_topic"
            user_states[uid] = state
            await event.respond(
                f"📚 Fan: **{text}**\n\n"
                f"Qaysi mavzudan savol tuzish kerak?\n\n"
                f"_(Misol: Kasrlar, Fotosintez, II Jahon urushi...)_\n\n"
                f"Barcha mavzudan bo'lsa 👇 tugmani bosing:",
                buttons=[
                    [Button.text("📖 Barcha mavzudan")],
                    [Button.text("🔙 Bosh menyu")],
                ]
            ); return

        if state.step == "ai_ask_topic":
            if text == "📖 Barcha mavzudan":
                state.topic = ""
            else:
                state.topic = text
            state.step = "ai_settings"
            user_states[uid] = state
            topic_show = state.topic if state.topic else "Barcha mavzu"
            await event.respond(
                f"📝 Fan: **{state.fan_name}**\n"
                f"📌 Mavzu: **{topic_show}**\n"
                f"🔢 Savol: {state.q_count} ta | 🎯 {state.difficulty} | 🌐 {state.lang}\n\n"
                f"Sozlash yoki yaratish:",
                buttons=ai_settings_btns(state)
            ); return

        if state.step == "ai_settings":
            topic_show = state.topic if state.topic else "Barcha mavzu"

            def _show_settings():
                return (
                    f"📝 Fan: **{state.fan_name}**\n"
                    f"📌 Mavzu: **{topic_show}**\n"
                    f"🔢 {state.q_count} ta | 🎯 {state.difficulty} | 🌐 {state.lang}"
                )

            # Fan nomini o'zgartirish
            if text == "📝 Fan nomini o'zgartirish":
                state.step = "ai_ask_fan"
                user_states[uid] = state
                await event.respond("Yangi fan nomini yozing:",
                    buttons=[[Button.text("🔙 Bosh menyu")]]); return

            # Mavzuni o'zgartirish
            if text == "📌 Mavzuni o'zgartirish":
                state.step = "ai_ask_topic"
                user_states[uid] = state
                await event.respond(
                    f"📚 Fan: **{state.fan_name}**\n\nYangi mavzuni yozing:",
                    buttons=[
                        [Button.text("📖 Barcha mavzudan")],
                        [Button.text("🔙 Bosh menyu")],
                    ]
                ); return

            # Savol soni
            if re.match(r'^🔢 (\d+) ta$', text):
                state.q_count = int(re.match(r'^🔢 (\d+) ta$', text).group(1))
                user_states[uid] = state
                await event.respond(_show_settings(), buttons=ai_settings_btns(state)); return

            # Qiyinlik
            diff_map = {"🟢 Oson": "oson", "🟡 O'rta": "o'rta", "🔴 Qiyin": "qiyin"}
            if text in diff_map:
                state.difficulty = diff_map[text]
                user_states[uid] = state
                await event.respond(_show_settings(), buttons=ai_settings_btns(state)); return

            # Til
            lang_map = {"🇺🇿 O'zbek": "uz", "🇷🇺 Rus": "ru", "🇬🇧 Ingliz": "en"}
            if text in lang_map:
                state.lang = lang_map[text]
                user_states[uid] = state
                await event.respond(_show_settings(), buttons=ai_settings_btns(state)); return

            # YARATISH
            if text.startswith("✅ Yaratish"):
                # Balans tekshirish
                bal = db_get_balance(uid)
                if bal < AI_PRICE:
                    state.step = "wait_payment"
                    user_states[uid] = state
                    needed = AI_PRICE - bal
                    merchant_trans_id = db_create_click_invoice(uid, needed)
                    click_url = create_click_url(needed, merchant_trans_id)
                    await event.respond(
                        f"❌ **Balans yetarli emas!**\n\n"
                        f"💰 Balansda: {bal:,} so'm\n"
                        f"💳 Kerak: {AI_PRICE:,} so'm\n"
                        f"➖ Yetishmaydi: **{needed:,} so'm**\n\n"
                        f"📌 Sozlamalaringiz saqlanib qoldi — to'lovdan keyin avtomatik davom etadi!",
                        buttons=[
                            [Button.url(f"💳 CLICK orqali {needed:,} so'm to'lash", click_url)],
                        ]
                    )
                    return

                state.step = "ai_generating"
                user_states[uid] = state
                topic_label = f" | 📌 {state.topic}" if state.topic else ""
                await event.respond(
                    f"🤖 **AI test tuzmoqda...**\n\n"
                    f"📚 {state.fan_name}{topic_label}\n"
                    f"🔢 {state.q_count} ta | 🎯 {state.difficulty} | 🌐 {state.lang}\n\n"
                    f"⏳ Bir oz kuting..."
                )
                try:
                    qs = await generate_questions(
                        state.fan_name, state.q_count,
                        state.lang, state.difficulty, state.topic
                    )
                    if not qs:
                        await event.respond("❌ AI savol yarata olmadi! Qayta urining.",
                            buttons=main_menu(adm, uid)); return

                    state.questions = qs
                    state.total_questions = len(qs)
                    state.per_variant = len(qs)
                    state.step = "ask_time"
                    state.__dict__['source'] = "ai"   # manba belgisi
                    user_states[uid] = state

                    # Balansdan yechish faqat savollar muvaffaqiyatli yaratilganda
                    db_deduct_balance(uid, AI_PRICE, f"AI test: {state.fan_name}")
                    bal_left = db_get_balance(uid)

                    await event.respond(
                        f"✅ **{len(qs)} ta savol tayyor!**\n"
                        f"💰 -{AI_PRICE:,} so'm | Qoldi: {bal_left:,} so'm\n\n"
                        f"⏱ Vaqt:",
                        buttons=time_btns()
                    )
                except Exception as e:
                    log.error(f"AI xato: {e}")
                    await event.respond(
                        f"❌ AI xato: {e}\n\nGROQ_API_KEY ni tekshiring!",
                        buttons=main_menu(adm)
                    )
                return

        # ---- MATN HOLAT ----
        if state.step == "wait_text":
            qs = _parse_questions(text)
            if qs:
                state.questions = qs
                state.total_questions = len(qs)
                state.step = "ask_fan_name"
                user_states[uid] = state
                await event.respond(f"✅ **{len(qs)} ta savol!**\n\nFan nomini yozing:",
                    buttons=[[Button.text("🔙 Bosh menyu")]]); return
            else:
                lines = [l.strip() for l in text.splitlines() if l.strip()]
                state.step = "manual_start"
                state.__dict__['raw_lines'] = lines
                state.__dict__['manual_q_idx'] = 0
                state.__dict__['manual_q_text'] = ""
                state.__dict__['manual_opts'] = []
                user_states[uid] = state
                await event.respond(
                    f"⚠️ Shablon aniqlanmadi. {len(lines)} ta qator.\nJavoblarni siz ko'rsatasiz:",
                    buttons=[[Button.text("▶️ Davom etish")], [Button.text("🔙 Bosh menyu")]]
                ); return

        # ---- FILE CONFIRMED — namuna ko'rib bo'lgach ----
        if state.step == "file_confirmed":
            if text == "✅ Maqul, davom etamiz":
                remember_cleanup(uid, getattr(event, 'id', None))
                await cleanup_user_flow_messages(uid, event.chat_id)
                q_count = state.total_questions
                price   = state.__dict__.get("_price", calc_file_price(q_count))
                blocks  = (q_count + 24) // 25
                bal_now = db_get_balance(uid)
                if bal_now >= price:
                    # Balans yetarli — to'g'ridan kesib fan nomi so'raymiz
                    db_deduct_balance(uid, price, f"Fayl quiz: {q_count} ta savol")
                    bal_left = db_get_balance(uid)
                    state.step = "ask_fan_name"
                    user_states[uid] = state
                    await event.respond(
                        f"✅ **Xizmat haqqi to'landi!**\n\n"
                        f"📂 {q_count} ta savol\n"
                        f"💰 -{price:,} so'm | Qolgan balans: {bal_left:,} so'm\n\n"
                        f"📚 Fan nomini yozing:",
                        buttons=[[Button.text("🔙 Bosh menyu")]]
                    )
                else:
                    # Balans yetarli emas — to'liq ma'lumot + Click tugmasi
                    needed = price - bal_now
                    mtid = db_create_click_invoice(uid, needed)
                    curl = create_click_url(needed, mtid)
                    state.step = "wait_payment_file"
                    user_states[uid] = state
                    await event.respond(
                        f"📂 **{q_count} ta savol**\n\n"
                        f"💰 Xizmat haqqi: {blocks} × 1 500 = **{price:,} so'm**\n"
                        f"💼 Balansda: {bal_now:,} so'm\n"
                        f"➖ Yetishmaydi: **{needed:,} so'm**\n\n"
                        f"📌 Savollar saqlanib qoldi — to'lovdan keyin davom etadi!",
                        buttons=[
                            [Button.url(f"💳 CLICK orqali {needed:,} so'm to'lash", curl)],
                        ]
                    )
                return

            if text == "❌ Bekor qilish":
                await cleanup_user_flow_messages(uid, event.chat_id)
                user_states[uid] = UserState()
                # Admin username/linkni olish
                bot_me_info = await bot_client.get_me()
                admin_links = []
                for aid in ADMIN_IDS:
                    try:
                        admin_entity = await bot_client.get_entity(aid)
                        uname = getattr(admin_entity, "username", None)
                        if uname:
                            admin_links.append(f"@{uname}")
                        else:
                            admin_links.append(f"Admin (ID: {aid})")
                    except Exception:
                        admin_links.append(f"Admin (ID: {aid})")
                admin_str = " yoki ".join(admin_links) if admin_links else "Admin"
                await event.respond(
                    f"😕 **Namuna maqul kelmadi?**\n\n"
                    f"Muammo bo'lsa @ksh247 ga murojaat qiling — yordam beramiz!\n\n"
                    f"Yoki fayl formatini tekshirib qaytadan yuboring.",
                    buttons=main_menu(is_admin(uid), uid)
                )
                # Adminga xabar
                sender = await event.get_sender()
                fn = getattr(sender, "first_name", "") or ""
                un = getattr(sender, "username", "") or ""
                full = f"{fn}".strip() or un or str(uid)
                ustr = f"@{un}" if un else f"ID: {uid}"
                await notify_admin(
                    f"⚠️ **Namuna maqul kelmadi**\n\n"
                    f"👤 {full} ({ustr})\n"
                    f"🆔 `{uid}`\n\n"
                    f"Foydalanuvchi fayl namunasini rad etdi."
                )
                return

        # ---- MANUAL REJIM ----
        if state.step == "manual_start" and text == "▶️ Davom etish":
            state.step = "manual_detect"
            state.questions = []
            state.__dict__['manual_q_idx'] = 0
            user_states[uid] = state
            await _ask_manual(event, uid, state); return

        if state.step == "manual_answer":
            await _handle_manual(event, uid, state, text); return

        # ---- FAN NOMI ----
        if state.step == "ask_fan_name":
            state.fan_name = text
            state.step = "ask_split"
            user_states[uid] = state
            total = state.total_questions
            m = await event.respond(
                f"📚 **{text}** | ❓ {total} savol\n\nHar variantda necha ta?",
                buttons=variant_btns(total))
            remember_cleanup(uid, getattr(event, 'id', None), getattr(m, 'id', None))
            return

        # ---- VARIANT SONI ----
        if state.step == "ask_split":
            n = 0 if text == "Hammasi bitta quiz" else int(re.sub(r'\D', '', text) or 0)
            state.per_variant = state.total_questions if n == 0 else max(1, min(n, state.total_questions))
            pv, total = state.per_variant, state.total_questions
            nv = (total + pv - 1) // pv
            state.step = "ask_time"
            user_states[uid] = state
            m = await event.respond(
                f"✅ **{nv} ta variant** × {pv} savol\n\n⏱ Vaqt:",
                buttons=time_btns())
            remember_cleanup(uid, getattr(event, 'id', None), getattr(m, 'id', None))
            return

        # ---- VAQT ----
        if state.step == "ask_time":
            tm = {"⏱ 15s": "15", "⏱ 30s": "30", "⏱ 60s": "60", "⏱ Chegarasiz": "0"}
            if text not in tm:
                await event.respond("Tugmadan tanlang!", buttons=time_btns()); return
            state.time_choice = tm[text]
            state.step = "ask_order"
            user_states[uid] = state
            m = await event.respond("🔀 Tartib:", buttons=order_btns())
            remember_cleanup(uid, getattr(event, 'id', None), getattr(m, 'id', None))
            return

        # ---- TARTIB ----
        if state.step == "ask_order":
            if text in ("📋 Ketma-ket", "🔀 Aralash"):
                # To'g'ri javoblar har safar o'zgarib ketmasligi uchun QuizBotda ketma-ket tartib ishlatiladi.
                state.order_choice = "order"
            else:
                await event.respond("Tugmadan tanlang!", buttons=order_btns()); return

            total, pv = state.total_questions, state.per_variant
            nv = (total + pv - 1) // pv
            state.step = "idle"
            user_states[uid] = state
            qs = state.questions

            # Keraksiz oraliq xabarlarni tozalaymiz: fan, variant, vaqt, tartib, xulosa ko'rinmaydi
            remember_cleanup(uid, getattr(event, 'id', None))
            await cleanup_user_flow_messages(uid, event.chat_id)

            first_part_count = len(qs[0:min(pv, total)]) if qs else total
            initial_progress = await event.respond(
                progress_text(5, 0, max(first_part_count, 1), estimate_seconds(max(first_part_count, 1))),
                buttons=main_menu(adm, uid)
            )

            new_reqs = []
            for v in range(nv):
                part = qs[v*pv:min((v+1)*pv, total)]
                new_reqs.append(
                    QuizRequest(
                        user_id=uid, chat_id=event.chat_id,
                        questions=part,
                        fan_name=state.fan_name, variant_num=v+1,
                        time_choice=state.time_choice,
                        order_choice=state.order_choice, total_variants=nv,
                        source=getattr(state, 'source', 'file'),
                        progress_msg_id=getattr(initial_progress, 'id', None) if v == 0 else None,
                    )
                )

            async with queue_lock:
                for req in new_reqs:
                    request_queue.append(req)
            return

    # ============================================================
    #  MANUAL REJIM FUNKSIYALARI
    # ============================================================
    async def _ask_manual(event, uid, state):
        lines = state.__dict__.get('raw_lines', [])
        idx   = state.__dict__.get('manual_q_idx', 0)
        if idx >= len(lines):
            if not state.questions:
                await event.respond("❌ Savol qo'shilmadi!", buttons=main_menu(is_admin(uid), uid))
                user_states[uid] = UserState(); return
            state.step = "ask_fan_name"
            state.total_questions = len(state.questions)
            user_states[uid] = state
            await event.respond(
                f"✅ **{len(state.questions)} ta savol tayyor!**\n\nFan nomini yozing:",
                buttons=[[Button.text("🔙 Bosh menyu")]]); return

        q_text = lines[idx]
        opts = []
        i = idx + 1
        while i < len(lines) and len(opts) < 6:
            opts.append(lines[i]); i += 1

        if len(opts) < 2:
            state.__dict__['manual_q_idx'] = i
            user_states[uid] = state
            await _ask_manual(event, uid, state); return

        state.__dict__['manual_q_text'] = q_text
        state.__dict__['manual_opts']   = opts
        state.__dict__['manual_q_idx']  = i
        state.step = "manual_answer"
        user_states[uid] = state

        letters = ["A","B","C","D","E","F"]
        opts_txt = "\n".join(f"  {letters[j] if j<6 else j+1}. {o}" for j, o in enumerate(opts))
        done = len(state.questions)
        approx = len(lines) // (len(opts) + 1)

        await event.respond(
            f"📝 **Savol {done+1}** (~{approx} ta):\n\n❓ {q_text}\n\n{opts_txt}\n\n✅ To'g'ri javob:",
            buttons=answer_btns(opts)
        )

    async def _handle_manual(event, uid, state, text):
        if text == "⏭ O'tkazib yuborish":
            await _ask_manual(event, uid, state); return
        letters = ["A","B","C","D","E","F"]
        opts = state.__dict__.get('manual_opts', [])
        correct = None
        for i, opt in enumerate(opts):
            ltr = letters[i] if i < 6 else str(i+1)
            if text.startswith(f"{ltr}."):
                correct = i; break
        if correct is None:
            await event.respond("Tugmadan tanlang!", buttons=answer_btns(opts)); return
        state.questions.append({
            "q": state.__dict__.get('manual_q_text', ''),
            "opts": opts, "ans": correct,
            "correct_text": opts[correct] if 0 <= correct < len(opts) else (opts[0] if opts else "")
        })
        user_states[uid] = state
        await _ask_manual(event, uid, state)

    # ============================================================
    #  ADMIN PANEL
    # ============================================================
    # ============================================================
    #  ADMIN NOTIFY — muhim hodisalarda xabar
    # ============================================================
    # Faol foydalanuvchilar: {user_id: {name, step, last_action, time}}
    active_users: dict = {}

    def track_user(uid: int, name: str, step: str, action: str):
        """Faol foydalanuvchini kuzatish"""
        active_users[uid] = {
            "name":   name,
            "step":   step,
            "action": action,
            "time":   datetime.now().strftime("%H:%M:%S"),
        }

    async def _show_admin(event):
        busy = sum(1 for v in account_busy.values() if v)
        free = len(account_pool) - busy
        rows = [f"  {i+1}. {'🔴' if account_busy.get(id(c)) else '🟢'} `{account_phones.get(id(c),'?')}`"
                for i, c in enumerate(account_pool)]
        total_users = db_count_users()
        stats = db_payment_stats()
        await event.respond(
            f"🔧 **ADMIN PANEL**\n\n"
            f"👥 Foydalanuvchilar: **{total_users} ta**\n"
            f"👀 Faol hozir: **{len(active_users)} ta**\n"
            f"💰 Bugungi daromad: **{stats['today']:,} so'm**\n"
            f"💵 Jami daromad: **{stats['total']:,} so'm**\n"
            f"⏳ Kutilayotgan to'lov: **{stats['pending']} ta**\n\n"
            f"📱 Akkauntlar: **{len(account_pool)} ta**\n"
            f"  🟢 Bo'sh: {free} | 🔴 Band: {busy}\n"
            f"📋 Navbat: **{len(request_queue)} ta**\n\n" +
            ("\n".join(rows) if rows else "  (yo'q)"),
            buttons=[
                [Button.text("👥 Userlar ro'yxati"), Button.text("💳 To'lovlar")],
                [Button.text("👀 Faol foydalanuvchilar"), Button.text("📋 Navbat")],
                [Button.text("➕ Akkaunt qo'shish"), Button.text("➖ Akkaunt o'chirish")],
                [Button.text("💸 Userga pul yuborish"), Button.text("📊 Holat")],
                [Button.text("📤 Sessiya yuklash"), Button.text("⬇️ DB yuklash")],
                [Button.text("⬆️ DB yuklash (yangi)"), Button.text("🗑 Navbatni tozalash")],
                [Button.text("📢 Barchaga xabar"), Button.text("📨 ID ga xabar")],
                [Button.text("📢 Majburiy a'zolik")],
                [Button.text("🔙 Bosh menyu")],
            ]
        )

    @bot_client.on(events.NewMessage(pattern="/cancel"))
    async def cmd_cancel(event):
        uid = event.sender_id
        if uid in admin_states:
            admin_states.pop(uid)
            await event.respond("❌ Bekor qilindi.", buttons=main_menu(is_admin(uid), uid))
        else:
            await event.respond("Hech narsa bekor qilinmadi.")

    @bot_client.on(events.NewMessage(pattern="/admin"))
    async def cmd_admin(event):
        if not is_admin(event.sender_id):
            await event.respond("⛔ Admin emassiz!"); return
        await _show_admin(event)

    # ============================================================
    #  DB YUKLAB OLISH VA YUKLASH
    # ============================================================

    @bot_client.on(events.NewMessage(pattern="/dbyuklash"))
    async def cmd_db_download(event):
        """DB faylini foydalanuvchiga yuborish"""
        if not is_admin(event.sender_id):
            await event.respond("⛔ Admin emassiz!"); return

        if not _os.path.exists(DB_FILE):
            await event.respond("❌ DB fayl topilmadi!")
            return

        size = _os.path.getsize(DB_FILE)
        await event.respond(
            f"📦 **Ma'lumotlar bazasi**\n\n"
            f"📁 `{_os.path.basename(DB_FILE)}`\n"
            f"💾 Hajm: {size / 1024:.1f} KB\n\n"
            f"⬇️ Yuklanmoqda..."
        )
        await bot_client.send_file(
            event.chat_id,
            DB_FILE,
            caption=(
                f"🗄 **bot.db** — {size / 1024:.1f} KB\n"
                f"📅 {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
                f"Bu faylni tahrir qilib /newdbyuklash orqali qayta yuboring."
            ),
            force_document=True
        )
        log.info(f"DB yuklandi: admin={event.sender_id}, size={size}")

    @bot_client.on(events.NewMessage(pattern="/newdbyuklash"))
    async def cmd_db_upload_prompt(event):
        """Yangi DB yuklash uchun ko'rsatma"""
        if not is_admin(event.sender_id):
            await event.respond("⛔ Admin emassiz!"); return

        admin_states[event.sender_id] = {"step": "wait_db_file"}
        await event.respond(
            "📤 **Yangi DB yuklash**\n\n"
            "⚠️ **Diqqat!** Joriy ma'lumotlar bazasi almashtiriladi!\n\n"
            "1. /dbyuklash orqali eski DB ni yuklang\n"
            "2. SQLite editor bilan tahrirlang\n"
            "3. Tahrirlangan `.db` faylni shu yerga yuboring\n\n"
            "Bot fayl qabul qilgach avtomatik qayta ishga tushadi.\n\n"
            "/cancel — bekor qilish",
            buttons=[[Button.text("🔙 Bosh menyu")]]
        )

    @bot_client.on(events.NewMessage(
        func=lambda e: e.file and
        admin_states.get(e.sender_id, {}).get("step") == "wait_db_file"
    ))
    async def cmd_db_receive(event):
        """Yangi DB faylini qabul qilish va almashtirish"""
        uid = event.sender_id
        if not is_admin(uid):
            return

        fname = getattr(event.file, 'name', '') or ''
        if not fname.lower().endswith('.db') and not fname.lower().endswith('.sqlite'):
            await event.respond(
                "❌ Faqat `.db` yoki `.sqlite` fayl yuboring!\n"
                "Qayta urinib ko'ring yoki /cancel."
            )
            return

        try:
            msg = await event.respond("📥 Yangi DB yuklanmoqda...")

            import io, shutil

            # Yangi DB ni olish
            buf = io.BytesIO()
            await event.download_media(file=buf)
            buf.seek(0)
            new_data = buf.read()

            if len(new_data) < 100:
                await msg.edit("❌ Fayl juda kichik yoki buzilgan!")
                return

            # SQLite fayl ekanligini tekshirish
            if not new_data.startswith(b'SQLite format 3'):
                await msg.edit("❌ Bu SQLite fayl emas!")
                return

            # Eski DB ni zaxiralash
            backup_path = DB_FILE + ".backup"
            if _os.path.exists(DB_FILE):
                shutil.copy2(DB_FILE, backup_path)
                log.info(f"DB zaxira: {backup_path}")

            # Yangi DB ni yozish
            with open(DB_FILE, 'wb') as f:
                f.write(new_data)

            admin_states.pop(uid, None)

            # Statistika
            import sqlite3 as _sq
            con = _sq.connect(DB_FILE)
            users_n  = con.execute("SELECT COUNT(*) FROM users").fetchone()[0]
            quizzes_n = con.execute("SELECT COUNT(*) FROM quizzes").fetchone()[0]
            pays_n   = con.execute("SELECT COUNT(*) FROM payments WHERE status='confirmed'").fetchone()[0]
            con.close()

            await msg.edit(
                f"✅ **DB muvaffaqiyatli yangilandi!**\n\n"
                f"👥 Foydalanuvchilar: {users_n} ta\n"
                f"🎯 Quizlar: {quizzes_n} ta\n"
                f"💰 To'lovlar: {pays_n} ta\n\n"
                f"💾 Zaxira: `{_os.path.basename(backup_path)}`\n\n"
                f"✅ Bot yangi DB bilan ishlaydi!"
            )
            log.info(f"DB yangilandi: admin={uid}, size={len(new_data)}")

        except Exception as e:
            log.error(f"DB yuklash xato: {e}")
            admin_states.pop(uid, None)
            await event.respond(f"❌ Xato: {e}")

    @bot_client.on(events.NewMessage(
        func=lambda e: not e.file and not e.text.startswith("/")
        and e.sender_id in ADMIN_IDS
        and e.text.strip() in ["➕ Akkaunt qo'shish","➖ Akkaunt o'chirish",
                                "📊 Holat","📋 Navbat","🗑 Navbatni tozalash",
                                "🔙 Admin panel", "👥 Userlar ro'yxati",
                                "💳 To'lovlar", "📤 Sessiya yuklash",
                                "⬇️ DB yuklash", "⬆️ DB yuklash (yangi)",
                                "👀 Faol foydalanuvchilar", "💸 Userga pul yuborish",
                                "💸 Yana yuborish", "📢 Barchaga xabar", "📨 ID ga xabar", "📢 Majburiy a'zolik", "✅ A'zolikni yoqish", "⛔️ A'zolikni o'chirish", "✏️ Kanal/guruhni almashtirish"]
    ))
    async def admin_btns(event):
        uid = event.sender_id
        adm = is_admin(uid)
        text = event.text.strip()

        if text == "📢 Majburiy a'zolik":
            s = db_get_subscription_settings()
            status = "✅ Yoqilgan" if s.get("enabled") else "⛔️ O'chirilgan"
            target = s.get("title") or s.get("chat_ref") or "sozlanmagan"
            link = s.get("invite_link") or "yo'q"
            chat_ref_val = s.get("chat_ref") or "-"
            await event.respond(
                f"📢 **Majburiy a'zolik sozlamalari**\n\n"
                f"Holat: {status}\n"
                f"Kanal/guruh: **{target}**\n"
                f"Chat ref: `{chat_ref_val}`\n"
                f"Havola: {link}\n\n"
                f"Foydalanuvchi bonus olishi uchun:\n"
                f"1) Shartlarga rozilik bildiradi\n"
                f"2) Kanal/guruhga a'zo bo'ladi\n"
                f"3) Shundan keyin referal/hamkor bonus beriladi",
                buttons=[
                    [Button.text("✅ A'zolikni yoqish"), Button.text("⛔️ A'zolikni o'chirish")],
                    [Button.text("✏️ Kanal/guruhni almashtirish")],
                    [Button.text("🔙 Admin panel")],
                ]
            )
            return

        if text == "✅ A'zolikni yoqish":
            s = db_get_subscription_settings()
            if not s.get("chat_ref"):
                await event.respond("❌ Avval kanal/guruhni sozlang.", buttons=[[Button.text("✏️ Kanal/guruhni almashtirish")],[Button.text("🔙 Admin panel")]])
                return
            db_set_subscription_enabled(True)
            await event.respond("✅ Majburiy a'zolik yoqildi.", buttons=[[Button.text("📢 Majburiy a'zolik")],[Button.text("🔙 Admin panel")]])
            return

        if text == "⛔️ A'zolikni o'chirish":
            db_set_subscription_enabled(False)
            await event.respond("⛔️ Majburiy a'zolik o'chirildi.", buttons=[[Button.text("📢 Majburiy a'zolik")],[Button.text("🔙 Admin panel")]])
            return

        if text == "✏️ Kanal/guruhni almashtirish":
            admin_states[uid] = {"step": "wait_subscription_target"}
            await event.respond(
                "✏️ **Kanal/guruhni almashtirish**\n\n"
                "Bitta xabarda quyidagi formatda yuboring:\n"
                "`chat_ref | invite_link | title`\n\n"
                "Misol 1:\n"
                "`@quiz_import_news | https://t.me/quiz_import_news | Yangiliklar kanali`\n\n"
                "Misol 2:\n"
                "`-1001234567890 | https://t.me/+abcxyz | VIP guruh`\n\n"
                "`title` ixtiyoriy.",
                buttons=[[Button.text("🔙 Admin panel")]]
            )
            return

        if text == "👀 Faol foydalanuvchilar":
            if not active_users:
                await event.respond(
                    "👀 **Faol foydalanuvchilar**\n\nHozir hech kim faol emas.",
                    buttons=[[Button.text("🔙 Admin panel")]]
                )
                return
            step_names = {
                "idle": "🏠 Bosh menyu",
                "ai_ask_fan": "🤖 Fan nomi yozmoqda",
                "ai_ask_topic": "🤖 Mavzu yozmoqda",
                "ai_settings": "🤖 AI sozlamalar",
                "ai_generating": "🤖 AI generatsiya kutmoqda",
                "wait_file": "📂 Fayl kutmoqda",
                "wait_text": "✏️ Matn yozmoqda",
                "wait_payment": "💳 To'lov kutmoqda (AI)",
                "wait_payment_file": "💳 To'lov kutmoqda (fayl)",
                "ask_fan_name": "📚 Fan nomi kiritmoqda",
                "ask_split": "🔢 Variant soni tanlayapti",
                "ask_time": "⏱ Vaqt tanlayapti",
                "ask_order": "🔀 Tartib tanlayapti",
                "manual_start": "✋ Manual rejim boshladi",
                "manual_detect": "✋ Manual savol ko'rib chiqmoqda",
                "manual_answer": "✋ Javob ko'rsatmoqda",
            }
            lines = [f"👀 **Faol foydalanuvchilar: {len(active_users)} ta**\n"]
            for u_id, info in list(active_users.items()):
                step_label = step_names.get(info['step'], info['step'])
                lines.append(
                    f"• **{info['name']}** (`{u_id}`)\n"
                    f"  {step_label}\n"
                    f"  📝 {info['action']}\n"
                    f"  🕐 {info['time']}"
                )
            await event.respond(
                "\n\n".join(lines),
                buttons=[[Button.text("🔄 Yangilash"), Button.text("🔙 Admin panel")]]
            )
            return

        if text == "🔙 Admin panel":
            admin_states.pop(uid, None)
            await _show_admin(event)

        elif text == "📤 Sessiya yuklash":
            admin_states[uid] = {"step": "wait_session_file"}
            # Mavjud akkauntlarni ko'rsatish
            existing = [account_phones.get(id(c), "?") for c in account_pool]
            notify_p = account_phones.get(id(notify_client_holder.get("client")), "") if notify_client_holder.get("client") else ""
            lines = ["📤 **Sessiya fayli yuklash**\n"]
            lines.append("Hozirgi akkauntlar:")
            for p in existing:
                lines.append(f"  🟢 {p}")
            if notify_p:
                lines.append(f"  🔔 {notify_p} (notify)")
            if not existing and not notify_p:
                lines.append("  (yo'q)")
            lines.append("\n`.session` faylini yuboring:")
            lines.append("_(Misol: userbot_998901234567.session)_\n")
            lines.append("/cancel — bekor")
            await event.respond(
                "\n".join(lines),
                buttons=[[Button.text("🔙 Admin panel")]]
            )

        elif text == "⬇️ DB yuklash":
            # /dbyuklash bilan bir xil
            if not _os.path.exists(DB_FILE):
                await event.respond("❌ DB fayl topilmadi!",
                    buttons=[[Button.text("🔙 Admin panel")]]); return
            size = _os.path.getsize(DB_FILE)
            await bot_client.send_file(
                event.chat_id,
                DB_FILE,
                caption=(
                    f"🗄 **bot.db** — {size/1024:.1f} KB\n"
                    f"📅 {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
                    f"Tahrir qilib /newdbyuklash orqali qayta yuboring."
                ),
                force_document=True
            )

        elif text == "⬆️ DB yuklash (yangi)":
            admin_states[uid] = {"step": "wait_db_file"}
            await event.respond(
                "📤 **Yangi DB yuklash**\n\n"
                "⚠️ Joriy ma'lumotlar bazasi almashtiriladi!\n\n"
                "Tahrirlangan `.db` faylni yuboring:\n\n"
                "/cancel — bekor",
                buttons=[[Button.text("🔙 Admin panel")]]
            )

        elif text == "👥 Userlar ro'yxati":
            total = db_count_users()
            users = db_get_users(limit=20)
            if not users:
                await event.respond("👥 Hali foydalanuvchi yo'q.",
                    buttons=[[Button.text("🔙 Admin panel")]]); return
            lines = [f"👥 **Foydalanuvchilar: {total} ta**\n"]
            for u in users:
                user_id, first, last, uname, created, last_seen = u
                name = f"{first} {last}".strip() or "Nomsiz"
                uname_str = f"@{uname}" if uname else "username yo'q"
                lines.append(
                    f"• {name} ({uname_str})\n"
                    f"  ID: `{user_id}` | {last_seen[:10]}"
                )
            if total > 20:
                lines.append(f"\n_...va yana {total-20} ta_")
            await event.respond(
                "\n".join(lines),
                buttons=[[Button.text("🔙 Admin panel")]]
            )

        elif text == "➕ Akkaunt qo'shish":
            admin_states[uid] = {"step": "wait_phone"}
            await event.respond("📱 Telefon raqam:\n_(+998901234567)_",
                buttons=[[Button.text("🔙 Admin panel")]])

        elif text == "➖ Akkaunt o'chirish":
            if not account_pool:
                await event.respond("❌ Akkaunt yo'q!"); return
            btns = [[Button.text(account_phones.get(id(c), "?"))] for c in account_pool]
            btns.append([Button.text("🔙 Admin panel")])
            admin_states[uid] = {"step": "wait_remove"}
            await event.respond("Qaysi raqamni o'chirish?", buttons=btns)

        elif text == "💸 Yana yuborish":
            admin_states[uid] = {"step": "wait_bonus_user_id"}
            await event.respond(
                "💸 **Userga bonus yuborish**\n\nUser ID ni yozing:",
                buttons=[[Button.text("🔙 Admin panel")]]
            )

        elif text == "💸 Userga pul yuborish":
            admin_states[uid] = {"step": "wait_bonus_user_id"}
            await event.respond(
                "💸 **Userga bonus yuborish**\n\n"
                "User ID ni yozing:\n"
                "_(Misol: 7693087447)_\n\n"
                "/cancel — bekor",
                buttons=[[Button.text("🔙 Admin panel")]]
            )

        elif text == "📊 Holat":
            lines = ["📊 **HOLAT**\n"]
            for i, c in enumerate(account_pool):
                ph = account_phones.get(id(c), "?")
                bs = "🔴 Band" if account_busy.get(id(c)) else "🟢 Bo'sh"
                au = "✅" if await c.is_user_authorized() else "❌"
                lines.append(f"{i+1}. `{ph}` {bs} {au}")
            lines.append(f"\n📋 Navbat: {len(request_queue)} ta")
            await event.respond("\n".join(lines), buttons=[[Button.text("🔙 Admin panel")]])

        elif text == "📋 Navbat":
            if not request_queue:
                await event.respond("📋 Bo'sh!", buttons=[[Button.text("🔙 Admin panel")]]); return
            lines = [f"📋 **{len(request_queue)} ta**\n"]
            for i, req in enumerate(list(request_queue)[:15]):
                secs = estimate_seconds(len(req.questions))
                lines.append(f"{i+1}. `{req.user_id}` {req.fan_name} V{req.variant_num} "
                             f"({len(req.questions)}ta) ~{format_wait(secs)}")
            lines.append(f"\n⏳ Umumiy: ~{calc_wait([])}")
            await event.respond("\n".join(lines), buttons=[[Button.text("🔙 Admin panel")]])

        elif text == "🗑 Navbatni tozalash":
            async with queue_lock:
                n = len(request_queue); request_queue.clear()
            await event.respond(f"🗑 {n} ta bekor qilindi.",
                buttons=[[Button.text("🔙 Admin panel")]])

        elif text == "📢 Barchaga xabar":
            if not adm: return
            admin_states[uid] = {"step": "wait_broadcast"}
            await event.respond(
                "📢 **Barchaga xabar yuborish**\n\n"
                "Yubormoqchi bo'lgan xabarni yozing:\n"
                "_(matn, rasm, video — hammasi bo'ladi)_",
                buttons=[[Button.text("❌ Bekor qilish")]]
            )

        elif text == "📨 ID ga xabar":
            if not adm: return
            admin_states[uid] = {"step": "wait_user_id_msg"}
            await event.respond(
                "📨 **User ID ga xabar yuborish**\n\n"
                "Avval user ID ni yozing:",
                buttons=[[Button.text("❌ Bekor qilish")]]
            )

        elif text == "💳 To'lovlar":
            stats = db_payment_stats()
            # Oxirgi 10 to'lov
            con = get_db()
            last_pays = con.execute(
                """SELECT p.id, p.user_id, p.amount, p.card_num, p.status, p.paid_at
                   FROM payments p ORDER BY p.id DESC LIMIT 10"""
            ).fetchall()
            con.close()

            # Kartalar holati
            card_lines = []
            for card in HUMO_CARDS:
                assigned = card_assignments.get(card)
                if assigned:
                    card_lines.append(f"  🔴 `{card[-9:]}` → user `{assigned}`")
                else:
                    card_lines.append(f"  🟢 `{card[-9:]}` bo'sh")

            pay_lines = []
            for p in last_pays:
                pid, uid2, amt, card, status, paid = p
                icon = {"confirmed": "✅", "pending": "⏳", "expired": "❌"}.get(status, "❓")
                pay_lines.append(
                    f"{icon} #{pid} | `{uid2}` | {amt:,} so'm | ...{card[-9:]}"
                )

            lines = [
                f"💳 **TO'LOV STATISTIKASI**\n",
                f"💵 Jami daromad: **{stats['total']:,} so'm**",
                f"📅 Bugun: **{stats['today']:,} so'm**",
                f"⏳ Kutilayotgan: **{stats['pending']} ta**\n",
                f"**Kartalar:**",
                *card_lines,
                f"\n**Oxirgi to'lovlar:**",
                *(pay_lines if pay_lines else ["  (yo'q)"]),
            ]
            await event.respond(
                "\n".join(lines),
                buttons=[[Button.text("🔙 Admin panel")]]
            )

    # ============================================================
    #  ADMIN: AKKAUNT QO'SHISH
    # ============================================================
    async def _admin_add_phone(event, uid, phone):
        phone = phone.strip()
        if not re.match(r'^\+\d{10,15}$', phone):
            await event.respond("❌ Format: +998901234567"); return
        if phone in [account_phones.get(id(c)) for c in account_pool]:
            await event.respond(f"⚠️ `{phone}` allaqachon bor!")
            admin_states.pop(uid, None); return
        await event.respond(f"📲 `{phone}` ga kod...")
        try:
            sess_dir = _os.path.dirname(DB_FILE)
            session = _os.path.join(sess_dir, f"userbot_{phone.replace('+','').replace(' ','')}")
            client = TelegramClient(session, API_ID, API_HASH)
            await client.connect()
            result = await client.send_code_request(phone)
            admin_states[uid] = {"step": "wait_code", "phone": phone,
                                  "client": client, "hash": result.phone_code_hash}
            await event.respond("✅ Kod yuborildi!\n\nKodni yuboring:\n_(12345)_",
                buttons=[[Button.text("🔙 Admin panel")]])
        except Exception as e:
            await event.respond(f"❌ {e}")
            admin_states.pop(uid, None)

    async def _admin_enter_code(event, uid, code):
        astate = admin_states.get(uid, {})
        client, phone, ph_hash = astate.get("client"), astate.get("phone"), astate.get("hash")
        if not client:
            await event.respond("❌ Sessiya tugadi."); admin_states.pop(uid, None); return
        try:
            await client.sign_in(phone=phone, code=code.strip().replace(" ",""),
                                  phone_code_hash=ph_hash)
            await pool_add(client, phone)
            save_extra_phones([account_phones.get(id(c)) for c in account_pool
                               if account_phones.get(id(c))])
            admin_states.pop(uid, None)
            await event.respond(f"✅ `{phone}` qo'shildi! Jami: **{len(account_pool)} ta**",
                buttons=[[Button.text("🔙 Admin panel")]])
        except SessionPasswordNeededError:
            admin_states[uid]["step"] = "wait_password"
            await event.respond("🔐 Parol kerak:", buttons=[[Button.text("🔙 Admin panel")]])
        except PhoneCodeInvalidError:
            await event.respond("❌ Kod noto'g'ri! Qayta:")
        except Exception as e:
            await event.respond(f"❌ {e}")
            try: await client.disconnect()
            except: pass
            admin_states.pop(uid, None)

    async def _admin_enter_pass(event, uid, password):
        astate = admin_states.get(uid, {})
        client, phone = astate.get("client"), astate.get("phone")
        if not client: admin_states.pop(uid, None); return
        try:
            await client.sign_in(password=password)
            await pool_add(client, phone)
            save_extra_phones([account_phones.get(id(c)) for c in account_pool
                               if account_phones.get(id(c))])
            admin_states.pop(uid, None)
            await event.respond(f"✅ `{phone}` qo'shildi! Jami: **{len(account_pool)} ta**",
                buttons=[[Button.text("🔙 Admin panel")]])
        except Exception as e:
            await event.respond(f"❌ Parol xato: {e}")
            try: await client.disconnect()
            except: pass
            admin_states.pop(uid, None)

    async def _admin_do_remove(event, uid, phone):
        phone = phone.strip()
        if await pool_remove(phone):
            save_extra_phones([account_phones.get(id(c)) for c in account_pool
                               if account_phones.get(id(c))])
            await event.respond(f"✅ `{phone}` o'chirildi! Qoldi: **{len(account_pool)} ta**",
                buttons=[[Button.text("🔙 Admin panel")]])
        else:
            await event.respond(f"❌ `{phone}` topilmadi yoki band!",
                buttons=[[Button.text("🔙 Admin panel")]])
        admin_states.pop(uid, None)

    async def _admin_bonus_user_id(event, uid, text):
        """Admin user ID kiritdi"""
        try:
            target_id = int(text.strip())
        except ValueError:
            await event.respond(
                "❌ Noto'g'ri format! Faqat raqam yozing:\n_(Misol: 7693087447)_",
                buttons=[[Button.text("🔙 Admin panel")]]
            )
            return
        user = db_get_user(target_id)
        if not user:
            await event.respond(
                f"❌ `{target_id}` ID li foydalanuvchi topilmadi!\n\n"
                f"User botga /start bosgan bo'lishi kerak.",
                buttons=[[Button.text("🔙 Admin panel")]]
            )
            return
        first  = user[1] or ""
        last   = user[2] or ""
        uname  = user[3] or ""
        bal    = user[4] or 0
        name   = f"{first} {last}".strip() or uname or str(target_id)
        uname_str = f"@{uname}" if uname else ""
        admin_states[uid] = {
            "step": "wait_bonus_amount",
            "target_id": target_id,
            "target_name": name,
        }
        await event.respond(
            f"👤 **Foydalanuvchi topildi:**\n\n"
            f"Ism: **{name}** {uname_str}\n"
            f"ID: `{target_id}`\n"
            f"💰 Hozirgi balans: **{bal:,} so'm**\n\n"
            f"Qancha so'm yuborasiz?\n_(Manfiy son ham bo'lishi mumkin, masalan: -1000)_",
            buttons=[[Button.text("🔙 Admin panel")]]
        )

    async def _admin_bonus_amount(event, uid, text):
        """Admin miqdor kiritdi"""
        astate = admin_states.get(uid, {})
        target_id   = astate.get("target_id")
        target_name = astate.get("target_name", str(target_id))
        try:
            amount = int(text.strip().replace(" ", "").replace(",", ""))
        except ValueError:
            await event.respond(
                "❌ Noto'g'ri miqdor! Faqat raqam yozing:\n_(Misol: 5000 yoki -1000)_",
                buttons=[[Button.text("🔙 Admin panel")]]
            )
            return
        if amount == 0:
            await event.respond("❌ 0 yuborib bo'lmaydi!", buttons=[[Button.text("🔙 Admin panel")]])
            return
        # Balansi yetarlimi (ayirish uchun)
        if amount < 0:
            bal = db_get_balance(target_id)
            if bal + amount < 0:
                await event.respond(
                    f"❌ Balans yetarli emas!\n"
                    f"Hozirgi balans: **{bal:,} so'm**\n"
                    f"Ayirilmoqchi: **{abs(amount):,} so'm**",
                    buttons=[[Button.text("🔙 Admin panel")]]
                )
                return
        db_add_balance(target_id, amount, f"Admin bonusi — {uid}")
        new_bal = db_get_balance(target_id)
        admin_states.pop(uid, None)
        icon = "💸" if amount > 0 else "➖"
        # Foydalanuvchiga xabar
        try:
            if amount > 0:
                await bot_client.send_message(
                    target_id,
                    f"🎁 **Sizga bonus yuborildi!**\n\n"
                    f"💰 +{amount:,} so'm\n"
                    f"💼 Yangi balans: **{new_bal:,} so'm**"
                )
            else:
                await bot_client.send_message(
                    target_id,
                    f"ℹ️ **Balans o'zgartirildi**\n\n"
                    f"💰 {amount:,} so'm\n"
                    f"💼 Yangi balans: **{new_bal:,} so'm**"
                )
        except Exception as e:
            log.warning(f"Foydalanuvchiga xabar yuborilmadi: {e}")
        await event.respond(
            f"{icon} **Muvaffaqiyatli!**\n\n"
            f"👤 {target_name} (`{target_id}`)\n"
            f"💰 {'+' if amount > 0 else ''}{amount:,} so'm\n"
            f"💼 Yangi balans: **{new_bal:,} so'm**",
            buttons=[[Button.text("💸 Yana yuborish"), Button.text("🔙 Admin panel")]]
        )
        log.info(f"Admin bonus: {uid} → {target_id}, {amount} so'm")

    # ============================================================
    #  PARSER (ichki)
    # ============================================================
    def _parse_questions(text: str) -> list:
        text = (text or "").replace("\ufeff", "").strip()

        def is_separator(s: str) -> bool:
            """====, ++++, ----- kabi satrlar"""
            return bool(re.match(r'^[=+\-_*]{3,}$', (s or '').strip()))

        def clean_line(s: str) -> str:
            return (s or '').replace('\xa0', ' ').strip()

        if re.search(r'={4,}', text) and re.search(r'\+{4,}', text):
            qs = []
            for block in re.split(r'\+{4,}', text):
                block = block.strip()
                if not block:
                    continue

                parts = re.split(r'={4,}', block, maxsplit=1)
                if len(parts) < 2:
                    continue

                q_text = clean_line(parts[0])
                opts_raw = [
                    clean_line(l) for l in parts[1].splitlines()
                    if clean_line(l) and not is_separator(clean_line(l))
                ]

                if not q_text or not opts_raw:
                    continue

                options = []
                correct = 0
                for idx, opt in enumerate(opts_raw):
                    if opt.startswith('#'):
                        correct = idx
                        options.append(opt[1:].strip())
                    else:
                        options.append(opt)

                if len(options) >= 2:
                    qs.append({"q": q_text, "opts": options, "ans": correct, "correct_text": options[correct] if 0 <= correct < len(options) else options[0]})

            if qs:
                return qs

        qs = []
        lines = [clean_line(l) for l in text.splitlines()]
        i = 0
        while i < len(lines):
            line = lines[i]
            if not line:
                i += 1
                continue

            if is_separator(line):
                i += 1
                continue

            is_q = bool(re.match(r'^\d+[\.\)]\s*.+', line)) or \
                   not re.match(r'^[a-zA-Z#][\.\)]\s*', line)
            if is_q:
                q_text = re.sub(r'^\d+[\.\)]\s*', '', line).strip()
                if not q_text:
                    i += 1
                    continue

                options, correct, opt_idx = [], 0, 0
                i += 1
                while i < len(lines):
                    vline = lines[i]
                    if not vline:
                        i += 1
                        break

                    if is_separator(vline):
                        i += 1
                        continue

                    if re.match(r'^\d+[\.\)]\s*.+', vline) and options:
                        break

                    m = re.match(r'^(#?)([a-zA-Z]?)[\.\)]\s*(.*)', vline)
                    if m:
                        is_correct = bool(m.group(1))
                        opt_text = m.group(3).strip()
                        if opt_text and not is_separator(opt_text):
                            if is_correct:
                                correct = opt_idx
                            options.append(opt_text)
                            opt_idx += 1
                        i += 1
                    elif vline.startswith("#"):
                        clean = re.sub(r'^#[a-dA-D]?[\.\)]\s*', '', vline[1:]).strip() or vline[1:].strip()
                        if clean and not is_separator(clean):
                            correct = opt_idx
                            options.append(clean)
                            opt_idx += 1
                        i += 1
                    else:
                        i += 1

                if q_text and len(options) >= 2:
                    qs.append({"q": q_text, "opts": options, "ans": correct})
            else:
                i += 1

        return qs

    # ============================================================
    #  TO'LOV HANDLERLARI
    # ============================================================

    @bot_client.on(events.NewMessage(pattern="💳 To'lov qilish"))
    async def cmd_pay(event):
        uid = event.sender_id
        bal = db_get_balance(uid)
        user_states[uid] = UserState(step="wait_click_amount")
        await event.respond(
            f"💳 **CLICK orqali to'lov**\n\n"
            f"💰 Hozirgi balans: **{bal:,} so'm**\n\n"
            f"Qancha to'lamoqchisiz? (so'mda yozing)\n"
            f"👇👇👇👇👇",
            buttons=[[Button.text("🔙 Bosh menyu")]]
        )

    @bot_client.on(events.NewMessage(func=lambda e: not e.file
                                     and e.text.strip() == "💰 Balansni ko'rish"))
    async def cmd_balance(event):
        uid = event.sender_id
        bal = db_get_balance(uid)
        tests = bal // AI_PRICE
        if bal < AI_PRICE:
            needed = AI_PRICE - bal
            merchant_trans_id = db_create_click_invoice(uid, needed)
            click_url = create_click_url(needed, merchant_trans_id)
            await event.respond(
                f"💰 **Balans: {bal:,} so'm**\n"
                f"🤖 AI test uchun kerak: {AI_PRICE:,} so'm\n"
                f"❌ Yetishmaydi: **{needed:,} so'm**\n\n"
                f"📌 To'lovdan keyin /start bosing — davom etadi.",
                buttons=[
                    [Button.url(f"💳 CLICK orqali {needed:,} so'm to'lash", click_url)],
                ]
            )
        else:
            await event.respond(
                f"💰 **Balans: {bal:,} so'm**\n"
                f"🤖 Tuzish mumkin: **{tests} ta** AI test\n\n"
                f"✅ Balans yetarli!",
                buttons=[
                    [Button.text("🤖 AI test tuzish")],
                    [Button.text("💳 To'lov qilish")],
                    [Button.text("🔙 Bosh menyu")],
                ]
            )

    # pay_click_start olib tashlandi — cmd_pay o'zi bajaradi

    @bot_client.on(events.NewMessage(
        func=lambda e: not e.file and
        user_states.get(e.sender_id, UserState()).step == "wait_click_amount"
    ))
    async def pay_click_amount(event):
        uid = event.sender_id
        text = event.text.strip()

        if text == "🔙 Bosh menyu":
            user_states[uid] = UserState()
            await event.respond("🏠 Bosh menyu", buttons=main_menu(is_admin(uid), uid))
            return

        # Summani parse qilish
        try:
            amount = int(re.sub(r'[^\d]', '', text))
        except Exception:
            await event.respond("Masalan: 10000")
            return

        if amount < 1000:
            await event.respond("❌ Minimal to'lov summasi: **1 000 so'm**")
            return

        if amount > 10_000_000:
            await event.respond("❌ Maksimal to'lov summasi: **10 000 000 so'm**")
            return

        user_states[uid] = UserState()

        # CLICK invoice yaratish
        merchant_trans_id = db_create_click_invoice(uid, amount)
        click_url = create_click_url(amount, merchant_trans_id)

        await event.respond(
            f"💳 **CLICK orqali to'lov**\n\n"
            f"💰 Summa: **{amount:,} so'm**\n\n"
            f"👇 Tugmani bosing → CLICK ilovasi ochiladi → to'lang\n"
            f"✅ To'lov o'tgach balans **avtomatik** yangilanadi",
            buttons=[
                [Button.url(f"💳 CLICK da {amount:,} so'm to'lash", click_url)],
            ]
        )
        log.info(f"CLICK invoice: user={uid}, amount={amount}, trans_id={merchant_trans_id}")

    # ============================================================
    #  @HUMOCARDBOT XABAR TINGLOVCHI
    # ============================================================
    # notify_client global — /notify_ulash orqali o'rnatilishi mumkin
    notify_client_holder = {"client": None}

    def setup_notify_listener(client):
        """Notify client ga @humocardbot handler o'rnatish"""
        @client.on(events.NewMessage(from_users="humocardbot"))
        async def on_humo_notify(event):
            text = event.text or ""
            log.info(f"humocardbot xabari: {text[:150]}")
            amount = _parse_amount(text)
            card   = _parse_card(text)
            log.info(f"Parse natijasi: summa={amount}, karta={card}")
            if not amount:
                log.warning(f"Summa aniqlanmadi: {text[:80]}")
                return
            if not card:
                log.warning(f"Karta aniqlanmadi: {text[:80]}")
                return
            log.info(f"To'lov aniqlandi: karta={card}, summa={amount}")
            con = get_db()
            row = con.execute(
                """SELECT id, user_id, amount FROM payments
                   WHERE card_num=? AND status='pending'
                   AND expires_at > datetime('now')
                   ORDER BY id DESC LIMIT 1""",
                (card,)
            ).fetchone()
            con.close()
            if not row:
                log.warning(f"Mos to'lov topilmadi: karta={card}")
                return
            pay_id, user_id, expected_amount = row
            if amount < expected_amount:
                await bot_client.send_message(
                    user_id,
                    f"⚠️ **Noto'g'ri summa!**\n\n"
                    f"Kerak: **{expected_amount:,} so'm**\n"
                    f"Tushgan: **{amount:,} so'm**\n\n"
                    f"Farq: {expected_amount - amount:,} so'm qo'shimcha yuboring!"
                )
                return
            db_confirm_payment(pay_id)
            db_add_balance(user_id, amount, f"To'lov #{pay_id} tasdiqlandi")
            release_card(card)
            bal = db_get_balance(user_id)
            tests = bal // AI_PRICE
            prev_state = user_states.get(user_id)
            has_pending_ai   = prev_state and prev_state.step == "wait_payment" and prev_state.fan_name
            has_pending_file = prev_state and prev_state.step in ("wait_payment_file", "file_confirmed") and prev_state.questions
            if has_pending_file:
                q_count = prev_state.total_questions
                price   = calc_file_price(q_count)
                if bal >= price:
                    db_deduct_balance(user_id, price, f"Fayl quiz: {q_count} ta savol")
                    bal_left = db_get_balance(user_id)
                    prev_state.step = "ask_fan_name"
                    user_states[user_id] = prev_state
                    await bot_client.send_message(
                        user_id,
                        f"✅ **To'lov tasdiqlandi! +{amount:,} so'm**\n\n"
                        f"📂 {q_count} ta savol tayyor\n"
                        f"💰 -{price:,} so'm | Balans: {bal_left:,} so'm\n\n"
                        f"Fan nomini yozing:",
                        buttons=[[Button.text("🔙 Bosh menyu")]]
                    )
                else:
                    await bot_client.send_message(
                        user_id,
                        f"✅ +{amount:,} so'm | Balans: {bal:,} so'm\n"
                        f"⚠️ Hali yetarli emas. Kerak: {price:,} so'm",
                        buttons=[[Button.text(f"💳 {price-bal:,} so'm to'lash"),
                                  Button.text("🔙 Bosh menyu")]]
                    )
            elif has_pending_ai:
                await bot_client.send_message(
                    user_id,
                    f"✅ **To'lov tasdiqlandi! +{amount:,} so'm**\n\n"
                    f"💼 Balans: **{bal:,} so'm**\n\n"
                    f"🤖 Oldingi sozlamalar:\n"
                    f"📚 {prev_state.fan_name}"
                    f"{f' | 📌 {prev_state.topic}' if prev_state.topic else ''}\n"
                    f"🔢 {prev_state.q_count} ta | 🎯 {prev_state.difficulty}\n\n"
                    f"⏳ AI test tuzilmoqda..."
                )
                try:
                    qs = await generate_questions(
                        prev_state.fan_name, prev_state.q_count,
                        prev_state.lang, prev_state.difficulty, prev_state.topic
                    )
                    if not qs:
                        await bot_client.send_message(user_id, "❌ AI savol yarata olmadi!")
                        return
                    db_deduct_balance(user_id, AI_PRICE, f"AI test: {prev_state.fan_name}")
                    bal_left = db_get_balance(user_id)
                    prev_state.questions = qs
                    prev_state.total_questions = len(qs)
                    prev_state.per_variant = len(qs)
                    prev_state.step = "ask_time"
                    user_states[user_id] = prev_state
                    await bot_client.send_message(
                        user_id,
                        f"✅ **{len(qs)} ta savol tayyor!**\n"
                        f"💰 Balans: {bal_left:,} so'm\n\n⏱ Vaqt:",
                        buttons=[[Button.text("⏱ 15s"), Button.text("⏱ 30s")],
                                 [Button.text("⏱ 60s"), Button.text("⏱ Chegarasiz")]]
                    )
                except Exception as e:
                    log.error(f"AI xato (to'lovdan keyin): {e}")
                    await bot_client.send_message(user_id, f"❌ AI xato: {e}")
            else:
                await bot_client.send_message(
                    user_id,
                    f"✅ **To'lov tasdiqlandi!**\n\n"
                    f"💰 +{amount:,} so'm\n"
                    f"💼 Balans: **{bal:,} so'm**\n"
                    f"🤖 {tests} ta AI test mumkin 🎉",
                    buttons=[[Button.text("🤖 AI test tuzish", resize=True),
                              Button.text("🔙 Bosh menyu", resize=True)]]
                )
            log.info(f"✅ To'lov tasdiqlandi: user={user_id}, +{amount} so'm")
        log.info(f"✅ @humocardbot tinglash aktiv: {client}")

    # Mavjud ulangan notify clientni sozlash
    for c in all_clients:
        if account_phones.get(id(c)) == NOTIFY_PHONE:
            notify_client_holder["client"] = c
            setup_notify_listener(c)
            log.info(f"✅ Notify aktiv: {NOTIFY_PHONE}")
            break
    else:
        log.warning(f"⚠️ Notify akkaunt topilmadi: {NOTIFY_PHONE} — /notify_ulash buyrug'ini ishlating")

    # ============================================================
    #  /session_yuklash — sessiya faylini bot orqali yuklash
    # ============================================================
    @bot_client.on(events.NewMessage(pattern="/session_yuklash"))
    async def cmd_session_upload(event):
        if not is_admin(event.sender_id): return
        admin_states[event.sender_id] = {"step": "wait_session_file"}
        await event.respond(
            "📤 **Sessiya fayli yuklash**\n\n"
            "`.session` faylini yuboring\n"
            "_(Misol: userbot_998934897111.session)_\n\n"
            "Fayl DB ga saqlanadi va bot uni ishlatadi.\n\n"
            "/cancel — bekor"
        )

    @bot_client.on(events.NewMessage(
        func=lambda e: e.file and
        admin_states.get(e.sender_id, {}).get("step") == "wait_session_file"
        and e.sender_id in ADMIN_IDS
    ))
    async def cmd_session_receive(event):
        uid = event.sender_id
        if not is_admin(uid): return

        fname = getattr(event.file, 'name', '') or ''
        if not fname.lower().endswith('.session'):
            await event.respond("❌ Faqat `.session` fayl yuboring!")
            return

        try:
            import io, base64
            buf = io.BytesIO()
            await event.download_media(file=buf)
            buf.seek(0)
            data = buf.read()

            if len(data) < 10:
                await event.respond("❌ Fayl bo'sh!")
                return

            # Telefon raqamini fayl nomidan ajratish
            # userbot_998934897111.session → +998934897111
            name = fname.replace('.session', '')
            digits = name.replace('userbot_', '').strip()
            if digits.startswith('998') and len(digits) >= 12:
                phone = '+' + digits
            elif digits.startswith('+'):
                phone = digits
            else:
                phone = '+' + digits

            # Sessiya faylini diskka yozish
            sess_dir = _os.path.dirname(DB_FILE)
            if sess_dir:
                _os.makedirs(sess_dir, exist_ok=True)
            session_path = _os.path.join(sess_dir, f"userbot_{phone.replace('+','').replace(' ','')}")
            with open(session_path + ".session", "wb") as f:
                f.write(data)

            # DB ga ham saqlash
            encoded = base64.b64encode(data).decode()
            con = get_db()
            con.execute("""
                INSERT INTO sessions (phone, session_data, updated_at)
                VALUES (?, ?, datetime('now'))
                ON CONFLICT(phone) DO UPDATE SET
                    session_data = excluded.session_data,
                    updated_at   = excluded.updated_at
            """, (phone, encoded))
            con.commit()
            con.close()

            admin_states.pop(uid, None)

            # Sessiya faylini yuklab, darhol ulanib ko'ramiz
            is_notify = (phone == NOTIFY_PHONE)
            try:
                client = TelegramClient(session_path, API_ID, API_HASH)
                await client.connect()
                if await client.is_user_authorized():
                    me = await client.get_me()
                    all_clients.append(client)
                    account_phones[id(client)] = phone
                    if is_notify:
                        notify_client_holder["client"] = client
                        setup_notify_listener(client)
                        await event.respond(
                            f"✅ **Notify akkaunt ulandi!**\n\n"
                            f"📱 `{phone}` (@{me.username or me.first_name})\n"
                            f"🔔 @humocardbot endi tinglanadi!",
                            buttons=[[Button.text("🔙 Admin panel")]]
                        )
                    else:
                        # Quiz pool ga qo'shamiz
                        already = any(account_phones.get(id(c)) == phone for c in account_pool)
                        if not already:
                            account_pool.append(client)
                            account_busy[id(client)] = False
                        await event.respond(
                            f"✅ **Akkaunt ulandi!**\n\n"
                            f"📱 `{phone}` (@{me.username or me.first_name})\n"
                            f"🎯 Quiz pool ga qo'shildi!\n"
                            f"Jami akkaunt: {len(account_pool)} ta",
                            buttons=[[Button.text("🔙 Admin panel")]]
                        )
                    log.info(f"Sessiya yuklandi va ulandi: {phone}")
                else:
                    await client.disconnect()
                    await event.respond(
                        f"💾 Sessiya saqlandi, lekin avtorizatsiya eski.\n"
                        f"📱 `{phone}`\n\n"
                        f"{'Endi /notify_ulash ni bosing' if is_notify else '➕ Akkaunt qo\'shish orqali qayta ulang'}",
                        buttons=[[Button.text("🔙 Admin panel")]]
                    )
            except Exception as conn_err:
                log.error(f"Sessiya ulanish xato: {conn_err}")
                await event.respond(
                    f"💾 Sessiya saqlandi!\n📱 `{phone}`\n\n"
                    f"Ulanishda xato: {conn_err}\n"
                    f"{'Qayta /notify_ulash bosing' if is_notify else 'Admin panel → ➕ Akkaunt qo\'shish'}",
                    buttons=[[Button.text("🔙 Admin panel")]]
                )
            log.info(f"Sessiya yuklandi: {phone}, {len(data)} bayt")

        except Exception as e:
            log.error(f"session_receive xato: {e}")
            admin_states.pop(uid, None)
            await event.respond(f"❌ Xato: {e}")

    # ============================================================
    #  /notify_ulash — bot orqali notify akkauntni ulash
    # ============================================================
    @bot_client.on(events.NewMessage(pattern="/notify_ulash"))
    async def cmd_notify_connect(event):
        if not is_admin(event.sender_id): return

        if notify_client_holder["client"]:
            phone = account_phones.get(id(notify_client_holder["client"]), "?")
            await event.respond(f"✅ Notify akkaunt allaqachon ulangan: `{phone}`")
            return
        if not NOTIFY_PHONE:
            await event.respond("❌ NOTIFY_PHONE environment variable o'rnatilmagan!")
            return

        # Agar avvalgi urinish hali aktiv bo'lsa — faqat kod so'raymiz
        existing = admin_states.get(event.sender_id, {})
        if existing.get("step") == "wait_notify_code" and existing.get("client"):
            await event.respond(
                f"⏳ Oldingi kod hali aktiv!\n\n"
                f"`{NOTIFY_PHONE}` ga kelgan kodni yuboring:\n"
                f"_(yoki /notify_yangi_kod — yangi kod olish)_"
            )
            return

        await event.respond(f"📲 `{NOTIFY_PHONE}` ga kod yuborilmoqda...")

        try:
            sess_dir = _os.path.dirname(DB_FILE)
            session  = _os.path.join(sess_dir, f"userbot_{NOTIFY_PHONE.replace('+','').replace(' ','')}")

            # Eski sessiya bo'lsa — avval undan urinib ko'ramiz
            if db_load_session(NOTIFY_PHONE, session):
                client = TelegramClient(session, API_ID, API_HASH)
                await client.connect()
                if await client.is_user_authorized():
                    all_clients.append(client)
                    account_phones[id(client)] = NOTIFY_PHONE
                    notify_client_holder["client"] = client
                    setup_notify_listener(client)
                    await event.respond(
                        f"✅ **Sessiya tiklandi! Kod shart emas.**\n\n"
                        f"📱 `{NOTIFY_PHONE}`\n"
                        f"🔔 @humocardbot xabarlari qabul qilinadi!"
                    )
                    return
                await client.disconnect()

            client = TelegramClient(session, API_ID, API_HASH)
            await client.connect()
            result = await client.send_code_request(NOTIFY_PHONE)

            admin_states[event.sender_id] = {
                "step":   "wait_notify_code",
                "phone":  NOTIFY_PHONE,
                "client": client,
                "hash":   result.phone_code_hash,
            }
            await event.respond(
                f"✅ **Kod yuborildi!**\n\n"
                f"📱 `{NOTIFY_PHONE}` ga kelgan kodni yuboring\n"
                f"_(bo'shliqsiz: 12345)_\n\n"
                f"⚠️ Kodni **2 daqiqa ichida** yuboring!\n"
                f"/cancel — bekor qilish"
            )
        except Exception as e:
            await event.respond(f"❌ Xato: {e}\n\nQayta urinish: /notify_ulash")
            admin_states.pop(event.sender_id, None)

    @bot_client.on(events.NewMessage(
        func=lambda e: not e.file and not e.text.startswith("/")
        and admin_states.get(e.sender_id, {}).get("step") == "wait_notify_code"
        and e.sender_id in ADMIN_IDS
    ))
    async def on_notify_code(event):
        uid    = event.sender_id
        code   = event.text.strip().replace(" ", "")
        astate = admin_states.get(uid, {})
        client  = astate.get("client")
        phone   = astate.get("phone")
        ph_hash = astate.get("hash")
        if not client:
            admin_states.pop(uid, None); return
        try:
            await client.sign_in(phone=phone, code=code, phone_code_hash=ph_hash)
            all_clients.append(client)
            account_phones[id(client)] = phone
            notify_client_holder["client"] = client
            setup_notify_listener(client)
            sess_dir = _os.path.dirname(DB_FILE)
            session  = _os.path.join(sess_dir, f"userbot_{phone.replace('+','').replace(' ','')}")
            db_save_session(phone, session)
            admin_states.pop(uid, None)
            await event.respond(
                f"✅ **Notify akkaunt ulandi!**\n\n"
                f"📱 `{phone}`\n"
                f"🔔 @humocardbot xabarlari endi qabul qilinadi!"
            )
            log.info(f"Notify akkaunt ulandi: {phone}")
        except SessionPasswordNeededError:
            admin_states[uid]["step"] = "wait_notify_pass"
            await event.respond("🔐 2FA parol kerak. Parolni yuboring:")
        except PhoneCodeInvalidError:
            await event.respond("❌ Kod noto'g'ri! Qayta yuboring:")
        except Exception as e:
            err = str(e).lower()
            # Kod eskirgan bo'lsa — avtomatik yangi kod yuboramiz
            if any(x in err for x in ["expired", "signinrequest", "phone_code_expired", "code_expired"]):
                try:
                    result = await client.send_code_request(phone)
                    admin_states[uid]["hash"] = result.phone_code_hash
                    await event.respond(
                        f"⚠️ **Kod eskirdi — yangi kod yuborildi!**\n\n"
                        f"`{phone}` ga kelgan yangi kodni yuboring:"
                    )
                except Exception as e2:
                    await event.respond(f"❌ Yangi kod yuborishda xato: {e2}\n\nQayta: /notify_ulash")
                    try: await client.disconnect()
                    except: pass
                    admin_states.pop(uid, None)
            else:
                await event.respond(f"❌ Xato: {e}\n\nQayta urinish: /notify_ulash")
                try: await client.disconnect()
                except: pass
                admin_states.pop(uid, None)

    @bot_client.on(events.NewMessage(
        func=lambda e: not e.file and not e.text.startswith("/")
        and admin_states.get(e.sender_id, {}).get("step") == "wait_notify_pass"
        and e.sender_id in ADMIN_IDS
    ))
    async def on_notify_pass(event):
        uid    = event.sender_id
        astate = admin_states.get(uid, {})
        client = astate.get("client")
        phone  = astate.get("phone")
        if not client:
            admin_states.pop(uid, None); return
        try:
            await client.sign_in(password=event.text.strip())
            all_clients.append(client)
            account_phones[id(client)] = phone
            notify_client_holder["client"] = client
            setup_notify_listener(client)
            # Sessiyani DB ga saqlash
            sess_dir = _os.path.dirname(DB_FILE)
            session  = _os.path.join(sess_dir, f"userbot_{phone.replace('+','').replace(' ','')}")
            db_save_session(phone, session)
            admin_states.pop(uid, None)
            await event.respond(
                f"✅ **Notify akkaunt ulandi!**\n📱 `{phone}`\n"
                f"🔔 @humocardbot endi tinglanadi!"
            )
        except Exception as e:
            await event.respond(f"❌ Parol xato: {e}")
            try: await client.disconnect()
            except: pass
            admin_states.pop(uid, None)

            amount = _parse_amount(text)
            card   = _parse_card(text)

            log.info(f"Parse natijasi: summa={amount}, karta={card}")

            if not amount:
                log.warning(f"Summa aniqlanmadi: {text[:80]}")
                return
            if not card:
                log.warning(f"Karta aniqlanmadi: {text[:80]}")
                return

    # ============================================================
    #  YORDAMCHI: XABARDAN SUMMA VA KARTA AJRATIB OLISH
    # ============================================================
    def _parse_amount(text: str) -> Optional[int]:
        """
        Xabardan summani topish.
        Humo formatlari:
          ➕ 2.000,00 UZS
          ➕ **2.000,00 UZS**
          +2000 UZS
          2 000,00 UZS
        """
        # Markdown bold va belgilarni tozalash
        clean = text.replace('*', '').replace('_', '').replace('`', '')

        patterns = [
            # "2.000,00 UZS" — nuqta minglik, vergul kasr (Humo asosiy format)
            r'[➕\+]?\s*([\d]{1,3}(?:\.[\d]{3})*),\d{2}\s*UZS',
            # "2 000,00 UZS"
            r'[➕\+]?\s*([\d]{1,3}(?:\s[\d]{3})*),\d{2}\s*UZS',
            # "2000 UZS" yoki "2000,00 UZS"
            r'[➕\+]?\s*(\d+)(?:,\d+)?\s*UZS',
            # Umumiy fallback
            r'(\d[\d\.\s]+\d)\s*UZS',
        ]

        for pat in patterns:
            m = re.search(pat, clean, re.IGNORECASE)
            if m:
                raw = m.group(1)
                # Nuqta va bo'shliqlarni olib tashlaymiz (minglik ajratgich)
                raw = raw.replace('.', '').replace(' ', '').replace('\xa0', '')
                try:
                    return int(raw)
                except Exception:
                    continue
        return None

    def _parse_card(text: str) -> Optional[str]:
        """
        Xabardan karta oxirgi 4 raqamini topib, DB dan to'liq raqamni olish.
        Humo format: HUMOCARD *8906
        """
        # Oxirgi 4 raqamni topish: "*8906", "* 8906", "**8906"
        m = re.search(r'\*+\s*(\d{4})\b', text)
        if not m:
            return None
        last4 = m.group(1)

        # DB dagi kartalar ichidan oxirgi 4 raqami mos keladiganni topish
        for card in HUMO_CARDS:
            if card.replace(' ', '').endswith(last4):
                return card
        return None


    # ============================================================
    #  FAYL PREVIEW — 5 ta random savol yuborish
    # ============================================================
    async def _show_file_price(chat_id: int, uid: int, q_count: int,
                               price: int, bal: int, blocks: int):
        state = user_states.get(uid)
        if state:
            state.step = "wait_payment_file"
            user_states[uid] = state
        if bal >= price:
            db_deduct_balance(uid, price, f"Fayl quiz: {q_count} ta savol")
            bal_left = db_get_balance(uid)
            if state:
                state.step = "ask_fan_name"
                user_states[uid] = state
            await bot_client.send_message(
                chat_id,
                f"\u2705 **Namuna yuborildi!**\n\n"
                f"\U0001F4C2 {q_count} ta savol topildi\n"
                f"\U0001F4B0 -{price:,} so'm | Balans: {bal_left:,} so'm\n\n"
                f"Quiz nomini yozing: \nmasalan: yakuniy savollari",
                buttons=[[Button.text("\U0001F519 Bosh menyu")]]
            )
        else:
            needed = price - bal
            mtid = db_create_click_invoice(uid, needed)
            curl = create_click_url(needed, mtid)
            await bot_client.send_message(
                chat_id,
                f"\U0001F4C2 **{q_count} ta savol topildi!**\n\n"
                f"\U0001F4B0 Xizmat narxi: {blocks} \xd7 1 500 = **{price:,} so'm**\n"
                f"\U0001F4BC Balansda: {bal:,} so'm\n"
                f"\u2796 Yetishmaydi: **{needed:,} so'm**\n\n"
                f"\U0001F4CC Savollar saqlanib qoldi \u2014 to'lovdan keyin davom etadi!",
                buttons=[
                    [Button.url(f"💳 CLICK orqali {needed:,} so'm to'lash", curl)],
                ]
            )

    async def _send_preview(userbot, req: QuizRequest, uid: int, chat_id: int,
                            q_count: int, price: int, bal: int, blocks: int):
        try:
            await bot_client.send_message(
                chat_id,
                f"\U0001F50D **Namuna (5 ta random savol)**\n"
                f"Savollaringiz to'g'ri o'qildimi? Ko'rib chiqing \U0001F447"
            )
            url = await make_quiz(userbot, req)
            if url:
                await bot_client.send_message(
                    chat_id,
                    f"\u2705 **Namuna tayyor!**\n\n"
                    f"\u25b6\ufe0f Quizni sinab ko'ring: {url}\n\n"
                    f"Maqul bo'lsa \u2014 to'lov qilib, butun to'plamni oling \U0001F447",
                    buttons=[
                        [Button.text("✅ Maqul, davom etamiz")],
                        [Button.text("❌ Bekor qilish")],
                    ]
                )
                state = user_states.get(uid)
                if state:
                    state.step = "file_confirmed"
                    state.__dict__["_price"]  = price
                    state.__dict__["_bal"]    = bal
                    state.__dict__["_blocks"] = blocks
                    user_states[uid] = state
            else:
                await _show_file_price(chat_id, uid, q_count, price, bal, blocks)
        except Exception as e:
            log.error(f"_send_preview xato: {e}")
            await _show_file_price(chat_id, uid, q_count, price, bal, blocks)
        finally:
            release(userbot)

    # ============================================================
    #  MUDDATI O'TGAN TO'LOVLARNI BEKOR QILISH (fon task)
    # ============================================================
    async def expire_checker():
        while True:
            await asyncio.sleep(60)
            expired = db_expire_old()
            if expired:
                log.info(f"⏰ {expired} ta to'lov muddati o'tdi, bekor qilindi")

    # ============================================================
    #  CLICK WEBHOOK SERVER (aiohttp)
    # ============================================================
    from aiohttp import web as aio_web

    async def click_prepare(request):
        try:
            data = dict(await request.post())
            log.info(f"CLICK Prepare: {data}")
            merchant_trans_id = data.get("merchant_trans_id", "")
            amount = float(data.get("amount", 0))
            if not verify_click_signature(data, action=0):
                return aio_web.json_response({"error": -1, "error_note": "SIGN CHECK FAILED"})
            invoice = db_get_click_invoice(merchant_trans_id)
            if not invoice:
                return aio_web.json_response({"error": -5, "error_note": "Invoice topilmadi"})
            inv_id, user_id, inv_amount, status = invoice
            if status == "paid":
                return aio_web.json_response({"error": -4, "error_note": "Allaqachon to'langan"})
            if abs(amount - inv_amount) > 1:
                return aio_web.json_response({"error": -2, "error_note": "Summa mos kelmaydi"})
            return aio_web.json_response({
                "click_trans_id": data.get("click_trans_id"),
                "merchant_trans_id": merchant_trans_id,
                "merchant_prepare_id": inv_id,
                "error": 0, "error_note": "Success"
            })
        except Exception as e:
            log.error(f"CLICK Prepare xato: {e}")
            return aio_web.json_response({"error": -9, "error_note": str(e)})

    async def click_complete(request):
        try:
            data = dict(await request.post())
            log.info(f"CLICK Complete: {data}")
            merchant_trans_id = data.get("merchant_trans_id", "")
            error = int(data.get("error", 0))
            if not verify_click_signature(data, action=1):
                return aio_web.json_response({"error": -1, "error_note": "SIGN CHECK FAILED"})
            if error < 0:
                return aio_web.json_response({"error": 0, "error_note": "Success"})
            result = db_confirm_click_invoice(merchant_trans_id)
            if not result:
                return aio_web.json_response({"error": -4, "error_note": "Invoice topilmadi"})
            user_id, amount = result
            db_add_balance(user_id, amount, f"CLICK: {merchant_trans_id}")
            # Hamkor komissiyasi
            partner_id = db_get_referred_by(user_id)
            if partner_id:
                commission = int(amount * PARTNER_PAY_PERCENT / 100)
                if commission > 0:
                    db_add_partner_balance(partner_id, commission, f"CLICK komissiya: {user_id}")
            # Foydalanuvchiga xabar — pending state bo'lsa avtomatik davom et
            bal = db_get_balance(user_id)
            prev_state = user_states.get(user_id)
            has_pending_ai   = prev_state and prev_state.step == "wait_payment" and prev_state.fan_name
            has_pending_file = prev_state and prev_state.step in ("wait_payment_file", "file_confirmed") and prev_state.questions

            try:
                if has_pending_file:
                    q_count = prev_state.total_questions
                    price   = calc_file_price(q_count)
                    if bal >= price:
                        db_deduct_balance(user_id, price, f"Fayl quiz: {q_count} savol")
                        bal_left = db_get_balance(user_id)
                        prev_state.step = "ask_fan_name"
                        user_states[user_id] = prev_state
                        await bot_client.send_message(
                            user_id,
                            f"✅ **To'lov qabul qilindi! +{amount:,} so'm**\n\n"
                            f"📂 {q_count} ta savol tayyor\n"
                            f"💰 -{price:,} so'm | Balans: {bal_left:,} so'm\n\n"
                            f"Fan nomini yozing:",
                            buttons=[[Button.text("🔙 Bosh menyu")]]
                        )
                    else:
                        still_need = price - bal
                        merchant_t = db_create_click_invoice(user_id, still_need)
                        curl = create_click_url(still_need, merchant_t)
                        await bot_client.send_message(
                            user_id,
                            f"✅ +{amount:,} so'm | Balans: {bal:,} so'm\n"
                            f"⚠️ Hali yetarli emas. Kerak: {price:,} so'm\n"
                            f"➖ Yana: **{still_need:,} so'm**",
                            buttons=[[Button.url(f"💳 CLICK {still_need:,} so'm", curl)]]
                        )
                elif has_pending_ai:
                    await bot_client.send_message(
                        user_id,
                        f"✅ **To'lov qabul qilindi! +{amount:,} so'm**\n\n"
                        f"💼 Balans: **{bal:,} so'm**\n\n"
                        f"🤖 Oldingi sozlamalar saqlanib qoldi:\n"
                        f"📚 {prev_state.fan_name}"
                        f"{f' | 📌 {prev_state.topic}' if prev_state.topic else ''}\n"
                        f"🔢 {prev_state.q_count} ta | 🎯 {prev_state.difficulty}\n\n"
                        f"⏳ AI test tuzilmoqda..."
                    )
                    try:
                        qs = await generate_questions(
                            prev_state.fan_name, prev_state.q_count,
                            prev_state.lang, prev_state.difficulty, prev_state.topic
                        )
                        if not qs:
                            await bot_client.send_message(user_id, "❌ AI savol yarata olmadi!")
                        else:
                            db_deduct_balance(user_id, AI_PRICE, f"AI test: {prev_state.fan_name}")
                            bal_left = db_get_balance(user_id)
                            prev_state.questions = qs
                            prev_state.total_questions = len(qs)
                            prev_state.per_variant = len(qs)
                            prev_state.step = "ask_time"
                            user_states[user_id] = prev_state
                            await bot_client.send_message(
                                user_id,
                                f"✅ **{len(qs)} ta savol tayyor!**\n"
                                f"💰 Balans: {bal_left:,} so'm\n\n⏱ Vaqt:",
                                buttons=[[Button.text("⏱ 15s"), Button.text("⏱ 30s")],
                                         [Button.text("⏱ 60s"), Button.text("⏱ Chegarasiz")]]
                            )
                    except Exception as ae:
                        log.error(f"AI (click to'lovdan keyin): {ae}")
                        await bot_client.send_message(user_id, f"❌ AI xato: {ae}")
                else:
                    tests = bal // AI_PRICE
                    await bot_client.send_message(
                        user_id,
                        f"✅ **To'lov qabul qilindi!**\n\n"
                        f"💰 +{amount:,} so'm\n"
                        f"💼 Balans: **{bal:,} so'm**\n"
                        f"🤖 {tests} ta AI test mumkin 🎉",
                        buttons=main_menu(is_admin(user_id), user_id)
                    )
            except Exception as e:
                log.error(f"Click complete xabar xato: {e}")
            await notify_admin(
                f"💳 **CLICK to'lov**\n\n"
                f"👤 `{user_id}`\n"
                f"💰 **{amount:,} so'm**\n"
                f"🔖 `{merchant_trans_id}`"
            )
            return aio_web.json_response({"error": 0, "error_note": "Success"})
        except Exception as e:
            log.error(f"CLICK Complete xato: {e}")
            return aio_web.json_response({"error": -9, "error_note": str(e)})

    app = aio_web.Application()
    app.router.add_post("/click/prepare", click_prepare)
    app.router.add_post("/click/complete", click_complete)
    app.router.add_get("/", lambda r: aio_web.json_response({"status": "ok", "bot": "AI Quiz Bot"}))
    app.router.add_get("/health", lambda r: aio_web.json_response({"status": "ok"}))
    app.router.add_get("/privacy", lambda r: aio_web.Response(text=PRIVACY_HTML, content_type="text/html"))
    app.router.add_get("/terms",   lambda r: aio_web.Response(text=TERMS_HTML,   content_type="text/html"))
    app.router.add_get("/oferta",  lambda r: aio_web.Response(text=OFERTA_HTML,  content_type="text/html"))

    runner = aio_web.AppRunner(app)
    await runner.setup()
    await aio_web.TCPSite(runner, "0.0.0.0", SERVER_PORT).start()
    log.info(f"✅ Webhook server: port {SERVER_PORT}")

    asyncio.create_task(expire_checker())
    asyncio.create_task(queue_worker())
    log.info(f"✅ AI Quiz Bot tayyor! {len(account_pool)} ta akkaunt.")
    await bot_client.run_until_disconnected()


if __name__ == "__main__":
    asyncio.run(main())
